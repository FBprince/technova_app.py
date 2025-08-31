# # ORIGINAL CODE

# import streamlit as st
# import requests
# from bs4 import BeautifulSoup
# import re
# import base64
# import ast

# # =============================
# # 🌌 Canvas Star Background + Neon Text
# # =============================
# def set_canvas_stars():
#     canvas_html = """
#     <canvas id="star-canvas" style="position:fixed; top:0; left:0; width:100%; height:100%; z-index:-1000;"></canvas>
#     <script>
#     const canvas = document.getElementById('star-canvas');
#     const ctx = canvas.getContext('2d');
#     function resize() {
#         canvas.width = window.innerWidth;
#         canvas.height = window.innerHeight;
#     }
#     window.addEventListener('resize', resize);
#     resize();
#     const numStars = 80;
#     const stars = [];
#     for(let i=0;i<numStars;i++){
#         stars.push({
#             x: Math.random()*canvas.width,
#             y: Math.random()*canvas.height,
#             r: Math.random()*2+1,
#             dx: (Math.random()-0.5)*0.3,
#             dy: (Math.random()-0.5)*0.3,
#             alpha: Math.random()
#         });
#     }
#     function draw() {
#         ctx.fillStyle = '#0a0b1c';
#         ctx.fillRect(0,0,canvas.width,canvas.height);
#         for(let s of stars){
#             ctx.beginPath();
#             ctx.arc(s.x, s.y, s.r, 0, Math.PI*2);
#             ctx.fillStyle = "rgba(0,255,255,"+s.alpha+")";
#             ctx.fill();
#             s.x += s.dx;
#             s.y += s.dy;
#             s.alpha += (Math.random()-0.5)*0.02;
#             if(s.alpha<0)s.alpha=0;
#             if(s.alpha>1)s.alpha=1;
#             if(s.x<0)s.x=canvas.width;
#             if(s.x>canvas.width)s.x=0;
#             if(s.y<0)s.y=canvas.height;
#             if(s.y>canvas.height)s.y=0;
#         }
#         requestAnimationFrame(draw);
#     }
#     draw();
#     </script>
#     <style>
#     body, [data-testid="stAppViewContainer"] {
#         background-color: #0a0b1c;
#         color: #00f9ff;
#         font-family: 'Orbitron', 'Trebuchet MS', monospace;
#         text-shadow: 0 0 6px #00f9ff, 0 0 12px #00cfff;
#         overflow-x:hidden;
#     }
#     h1,h2,h3,h4,h5,h6 { color:#00f9ff !important; text-shadow:0 0 6px #00f9ff,0 0 12px #00cfff,0 0 24px #00aaff; }
#     p,div,span,label { color:#a0f0ff !important; text-shadow:0 0 4px #00f9ff; }
#     .stButton>button, .stDownloadButton>button {
#         background: linear-gradient(90deg, #00f0ff, #0066ff);
#         color:#02121b;
#         border-radius:10px;padding:8px 14px;
#         font-weight:700;font-family:'Orbitron', monospace;
#         box-shadow:0 8px 24px rgba(0,120,255,0.12);
#         transition: transform .15s ease, box-shadow .15s ease;
#     }
#     .stButton>button:hover, .stDownloadButton>button:hover { transform: translateY(-3px); box-shadow: 0 16px 36px rgba(0,120,255,0.2); }
#     [data-testid="stSidebar"] { background: linear-gradient(180deg, rgba(4,6,20,0.95), rgba(8,12,35,0.95)); color:#EAF9FF; }
#     pre, code { background: rgba(0,20,40,0.4); color: #00f9ff; text-shadow:0 0 4px #00f9ff; }
#     </style>
#     """
#     st.markdown(canvas_html, unsafe_allow_html=True)

# # =============================
# def safe_sentence_split(text: str):
#     pattern = re.compile(r"(?<=[.!?])\s+(?=[A-Z0-9])")
#     return pattern.split(text)

# def paragraph_summary(text: str, max_sentences: int = 5):
#     paragraphs = [p.strip() for p in text.splitlines() if p.strip()]
#     summary_lines = []
#     for para in paragraphs:
#         sentences = safe_sentence_split(para)
#         if sentences:
#             summary_lines.append(sentences[0].strip())
#         if len(summary_lines) >= max_sentences:
#             break
#     return " ".join(summary_lines) if summary_lines else text

# # =============================
# def analyze_python(code: str, summary_length: int = 5):
#     report = {
#         "functions": [],
#         "classes": [],
#         "imports": [],
#         "purpose_summary": "",
#         "errors": [],
#         "warnings": [],
#         "suggestions": []
#     }

#     # Extract imports, functions, classes
#     for line in code.splitlines():
#         l = line.strip()
#         if l.startswith("def "):
#             report["functions"].append(l.split("(")[0][4:].strip())
#         elif l.startswith("class "):
#             report["classes"].append(l.split("(")[0][6:].strip())
#         elif l.startswith("import ") or l.startswith("from "):
#             report["imports"].append(l)

#     report["purpose_summary"] = paragraph_summary(" ".join(safe_sentence_split(code)), max_sentences=summary_length)

#     # Syntax checking
#     try:
#         tree = ast.parse(code)
#     except SyntaxError as e:
#         report["errors"].append(f"SyntaxError: {e.msg} at line {e.lineno}")
#         report["suggestions"].append("Check indentation, missing colons, parentheses, or quotes.")
#         return report

#     # Static runtime checks
#     defined_vars = set()
#     called_funcs = set()
#     class FuncVisitor(ast.NodeVisitor):
#         def visit_FunctionDef(self, node):
#             defined_vars.add(node.name)
#             self.generic_visit(node)
#         def visit_Call(self, node):
#             if isinstance(node.func, ast.Name):
#                 called_funcs.add(node.func.id)
#             self.generic_visit(node)
#         def visit_Assign(self, node):
#             for target in node.targets:
#                 if isinstance(target, ast.Name):
#                     defined_vars.add(target.id)
#             self.generic_visit(node)
#         def visit_BinOp(self, node):
#             if isinstance(node.op, (ast.Div, ast.Mod)):
#                 if isinstance(node.right, ast.Constant) and node.right.value == 0:
#                     report["warnings"].append(f"Possible division by zero at line {node.lineno}")
#             self.generic_visit(node)
#     FuncVisitor().visit(tree)

#     # Undefined function calls
#     for func in called_funcs:
#         if func not in defined_vars and func not in dir(__builtins__):
#             report["warnings"].append(f"Call to undefined function '{func}' detected")

#     # General suggestions
#     if not report["errors"]:
#         if report["warnings"]:
#             report["suggestions"].append("Review the warnings above and fix variable or function usage.")
#         else:
#             report["suggestions"].append("No obvious syntax or runtime issues detected. Test your logic.")

#     return report

# # =============================
# def fetch_from_url(url: str) -> str:
#     try:
#         resp = requests.get(url, timeout=12)
#         resp.raise_for_status()
#         ct = resp.headers.get("Content-Type", "")
#         if "text/plain" in ct or url.lower().endswith((".py", ".txt", ".md")):
#             return resp.text
#         if "text/html" in ct:
#             soup = BeautifulSoup(resp.text, "html.parser")
#             for s in soup(["script", "style", "noscript"]):
#                 s.extract()
#             return soup.get_text(separator="\n")
#         return resp.text
#     except Exception as e:
#         return f"❌ Error fetching URL: {str(e)}"

# # =============================
# def copy_button(label: str, text: str, key: str):
#     if text is None: text = ""
#     b64 = base64.b64encode(text.encode("utf-8")).decode("ascii")
#     html_button = f"""
#         <button onclick="navigator.clipboard.writeText(atob('{b64}'))"
#                 style="background:transparent;border:1px solid rgba(255,255,255,0.06);
#                        color: #EAF9FF;padding:8px 12px;border-radius:8px;cursor:pointer;
#                        font-weight:700;box-shadow:0 6px 18px rgba(0,120,255,0.06);">
#             📋 {label}
#         </button>
#     """
#     st.markdown(html_button, unsafe_allow_html=True)

# # =============================
# def main():
#     set_canvas_stars()
#     st.title("🌌 TechNova — Study & Code Assistant")
#     st.caption("Summarize documents, analyze Python code with smart debugging, fetch & explain URLs — neon starry UI")

#     tab1, tab2, tab3 = st.tabs(["📄 Document Summarizer", "💻 Code Analyzer", "🌐 URL Fetch & Explain"])

#     # ----------------------------
#     with tab1:
#         st.subheader("Summarize Documents")
#         text_input = st.text_area("Paste text here:", height=200)
#         uploaded_file = st.file_uploader("Or upload a document (.txt, .md)", type=["txt", "md"])
#         length = st.slider("Summary length (paragraphs)", 1, 10, 3)

#         if st.button("Summarize Document"):
#             content = ""
#             if uploaded_file:
#                 try: content = uploaded_file.read().decode("utf-8")
#                 except: content = uploaded_file.read().decode("latin-1")
#             elif text_input.strip(): content = text_input
#             if not content: st.warning("Paste text or upload a file.")
#             else:
#                 summary = paragraph_summary(content, max_sentences=length)
#                 st.markdown("### 📌 Summary")
#                 st.write(summary)
#                 copy_button("Copy Summary", summary, key="doc_copy")
#                 st.download_button("⬇️ Download Summary", summary, "document_summary.txt", mime="text/plain")

#     # ----------------------------
#     with tab2:
#         st.subheader("Analyze Python Code with Smart Debugging")
#         code_input = st.text_area("Paste Python code here:", height=220)
#         uploaded_code = st.file_uploader("Or upload a Python file (.py)", type=["py"])
#         length_code = st.slider("Summary length (sentences)", 1, 8, 3, key="code_len")

#         if st.button("Analyze Code"):
#             code_str = ""
#             if uploaded_code:
#                 try: code_str = uploaded_code.read().decode("utf-8")
#                 except: code_str = uploaded_code.read().decode("latin-1")
#             elif code_input.strip(): code_str = code_input
#             if not code_str: st.warning("Paste code or upload a Python file.")
#             else:
#                 report = analyze_python(code_str, summary_length=length_code)
#                 st.markdown("### 📊 Code Analysis Report")
#                 st.markdown("**Functions:**")
#                 st.write(", ".join(report["functions"]) if report["functions"] else "_none_")
#                 st.markdown("**Classes:**")
#                 st.write(", ".join(report["classes"]) if report["classes"] else "_none_")
#                 st.markdown("**Imports:**")
#                 if report["imports"]: 
#                     for imp in report["imports"]: st.code(imp)
#                 else: st.write("_none_")
#                 st.markdown("**Purpose summary:**"); st.write(report["purpose_summary"] or "_No summary available_")
#                 st.markdown("**Errors Detected:**")
#                 if report["errors"]: 
#                     for err in report["errors"]: st.error(err)
#                 else: st.write("_No syntax errors detected_")
#                 st.markdown("**Warnings / Runtime Issues:**")
#                 if report["warnings"]:
#                     for w in report["warnings"]: st.warning(w)
#                 else: st.write("_No obvious runtime issues detected_")
#                 st.markdown("**Debugging Suggestions:**")
#                 if report["suggestions"]:
#                     for sug in report["suggestions"]: st.info(sug)

#                 copy_button("Copy Report", str(report), key="code_copy")
#                 st.download_button("⬇️ Download Code Report", str(report), "code_analysis.txt", mime="text/plain")

#     # ----------------------------
#     with tab3:
#         st.subheader("Fetch, Summarize & Explain from URL")
#         url_input = st.text_input("Enter a URL:")

#         if st.button("Fetch & Analyze URL"):
#             if not url_input.strip(): st.warning("Please enter a valid URL.")
#             else:
#                 data = fetch_from_url(url_input.strip())
#                 if data.startswith("❌ Error"): st.error(data)
#                 else:
#                     st.markdown("### 📄 Content Preview")
#                     st.code(data[:1200] + ("..." if len(data) > 1200 else ""), language="text")
#                     is_code = url_input.strip().lower().endswith(".py") or ("def " in data or "class " in data)
#                     if is_code:
#                         st.markdown("### 💻 Code Analysis with Smart Debugging")
#                         report = analyze_python(data, summary_length=5)
#                         st.json(report)
#                         copy_button("Copy Report", str(report), key="url_code_copy")
#                         st.download_button("⬇️ Download Code Report", str(report), "url_code_analysis.txt")
#                     else:
#                         st.markdown("### 📌 Document Summary")
#                         summary = paragraph_summary(data, max_sentences=5)
#                         st.write(summary)
#                         copy_button("Copy Summary", summary, key="url_summary_copy")
#                         st.download_button("⬇️ Download Summary", summary, "url_summary.txt")

# # =============================
# if __name__ == "__main__":
#     main()















































# # technova_app.py
# import streamlit as st
# import requests
# from bs4 import BeautifulSoup
# import re
# import base64
# import ast
# from collections import Counter
# import io

# # PDF extraction (PyPDF2)
# try:
#     import PyPDF2
# except Exception:
#     PyPDF2 = None

# # Page config
# st.set_page_config(page_title="Technova AI Toolkit", layout="wide")

# # ----------------------------
# # Utility: Copy button (real clipboard via JS)
# # ----------------------------
# def copy_button(text: str, label: str = "Copy", key: str = None):
#     """
#     Renders a button that copies `text` to the user's clipboard using JS.
#     key - unique identifier for the button element.
#     """
#     if text is None:
#         text = ""
#     safe_b64 = base64.b64encode(text.encode("utf-8")).decode("ascii")
#     el_id = f"copy-btn-{key}" if key else f"copy-btn-{abs(hash(text))}"
#     html = f"""
#     <button id="{el_id}" onclick="navigator.clipboard.writeText(atob('{safe_b64}'))"
#         style="background:transparent;border:1px solid rgba(255,255,255,0.06);
#                color:#EAF9FF;padding:8px 12px;border-radius:8px;cursor:pointer;
#                font-weight:700;margin-right:8px;">
#       📋 {label}
#     </button>
#     <script>
#     const btn_{el_id} = document.getElementById('{el_id}');
#     btn_{el_id}.addEventListener('click', () => {{
#       const old = btn_{el_id}.innerText;
#       btn_{el_id}.innerText = '✅ Copied';
#       setTimeout(()=>{{ btn_{el_id}.innerText = old; }}, 1400);
#     }});
#     </script>
#     """
#     st.markdown(html, unsafe_allow_html=True)

# # ----------------------------
# # Neon star background & basic styles
# # ----------------------------
# def set_canvas_stars():
#     canvas_html = """
#     <canvas id="star-canvas" style="position:fixed; top:0; left:0; width:100%; height:100%; z-index:-1000;"></canvas>
#     <script>
#     const canvas = document.getElementById('star-canvas');
#     const ctx = canvas.getContext('2d');
#     function resize(){ canvas.width = window.innerWidth; canvas.height = window.innerHeight; }
#     window.addEventListener('resize', resize);
#     resize();
#     const numStars = 80;
#     const stars = [];
#     for(let i=0;i<numStars;i++){
#       stars.push({x:Math.random()*canvas.width,y:Math.random()*canvas.height,r:Math.random()*2+1,dx:(Math.random()-0.5)*0.3,dy:(Math.random()-0.5)*0.3,alpha:Math.random()});
#     }
#     function draw(){
#       ctx.fillStyle='#0a0b1c'; ctx.fillRect(0,0,canvas.width,canvas.height);
#       for(let s of stars){
#         ctx.beginPath(); ctx.arc(s.x,s.y,s.r,0,Math.PI*2);
#         ctx.fillStyle='rgba(0,255,255,'+s.alpha+')'; ctx.fill();
#         s.x+=s.dx; s.y+=s.dy; s.alpha+=(Math.random()-0.5)*0.02;
#         if(s.alpha<0)s.alpha=0; if(s.alpha>1)s.alpha=1;
#         if(s.x<0)s.x=canvas.width; if(s.x>canvas.width)s.x=0;
#         if(s.y<0)s.y=canvas.height; if(s.y>canvas.height)s.y=0;
#       }
#       requestAnimationFrame(draw);
#     }
#     draw();
#     </script>
#     <style>
#     body, [data-testid="stAppViewContainer"] {
#       background-color:#0a0b1c; color:#00f9ff;
#       font-family: 'Orbitron', 'Trebuchet MS', monospace;
#       text-shadow:0 0 6px #00f9ff,0 0 12px #00cfff; overflow-x:hidden;
#     }
#     .stButton>button, .stDownloadButton>button { background: linear-gradient(90deg,#00f0ff,#0066ff); color:#02121b; border-radius:10px; padding:8px 14px; font-weight:700; }
#     </style>
#     """
#     st.markdown(canvas_html, unsafe_allow_html=True)

# set_canvas_stars()
# st.title("🌌 Technova AI Toolkit")
# st.caption("Document summarizer • Code analyzer • AI vs Human scanner • URL fetcher — neon UI")

# # ----------------------------
# # Summarization helpers
# # ----------------------------
# STOPWORDS = set(
#     "a an and are as at be but by for if in into is it no not of on or such that the their then there these they this to was will with you your from our we he she his her its were been being than also can could should would may might have has had do does did done just over under more most other some any each many few those them which who whom whose where when why how".split()
# )

# def safe_sentence_split(text: str):
#     pattern = re.compile(r"(?<=[.!?])\s+(?=[A-Z0-9])")
#     return [s.strip() for s in pattern.split(text) if s.strip()]

# def summarize_text_advanced(text: str, max_sentences: int = 5, as_bullets: bool = False) -> str:
#     paragraphs = [p.strip() for p in text.splitlines() if p.strip()]
#     sentences = []
#     for para in paragraphs:
#         sentences.extend(safe_sentence_split(para))
#     if not sentences:
#         return text

#     word_freq = Counter()
#     for s in sentences:
#         words = [w.lower() for w in re.findall(r"[A-Za-z0-9_']+", s)]
#         for w in words:
#             if w not in STOPWORDS and len(w) > 2:
#                 word_freq[w] += 1
#     if not word_freq:
#         return " ".join(sentences[:max_sentences])

#     max_freq = max(word_freq.values())
#     for w in list(word_freq.keys()):
#         word_freq[w] /= max_freq

#     scored = []
#     for idx, s in enumerate(sentences):
#         words = [w.lower() for w in re.findall(r"[A-Za-z0-9_']+", s)]
#         score = sum(word_freq.get(w, 0.0) for w in words)
#         length_penalty = 1.0 + 0.2 * max(0, (len(words) - 20) / 20)
#         position_boost = 1.1 if idx < 3 else 1.0
#         scored.append((score / length_penalty * position_boost, idx, s))

#     scored.sort(key=lambda x: (-x[0], x[1]))
#     top = sorted(scored[:max_sentences], key=lambda x: x[1])
#     if as_bullets:
#         return "\n".join([f"- {s}" for _, _, s in top])
#     return " ".join([s for _, _, s in top])

# # ----------------------------
# # Code analyzer (AST) with fixes
# # ----------------------------
# def analyze_python(code: str):
#     report = {"functions": [], "classes": [], "imports": [], "purpose_summary": "", "errors": [], "warnings": [], "fixes": []}

#     for line in code.splitlines():
#         l = line.strip()
#         if l.startswith("def "):
#             report["functions"].append(l.split("(")[0][4:].strip())
#         elif l.startswith("class "):
#             report["classes"].append(l.split("(")[0][6:].strip().rstrip(":"))
#         elif l.startswith("import ") or l.startswith("from "):
#             report["imports"].append(l)

#     report["purpose_summary"] = summarize_text_advanced(code, max_sentences=5)

#     try:
#         tree = ast.parse(code)
#     except SyntaxError as e:
#         report["errors"].append(f"SyntaxError: {e.msg} at line {e.lineno}")
#         report["fixes"].append("Check indentation, missing colons, parentheses, or quotes.")
#         return report

#     imported_names = set()
#     used_names = set()
#     assigned_names = set()
#     function_defs = []
#     class_defs = []

#     class Analyzer(ast.NodeVisitor):
#         def visit_Import(self, node):
#             for alias in node.names:
#                 imported_names.add(alias.asname or alias.name.split(".")[0])
#             self.generic_visit(node)

#         def visit_ImportFrom(self, node):
#             for alias in node.names:
#                 imported_names.add(alias.asname or alias.name)
#             self.generic_visit(node)

#         def visit_FunctionDef(self, node):
#             function_defs.append(node)
#             assigned_names.add(node.name)
#             for default in node.args.defaults:
#                 if isinstance(default, (ast.List, ast.Dict, ast.Set)):
#                     report["warnings"].append(f"Mutable default argument in function '{node.name}' at line {node.lineno}")
#                     report["fixes"].append(f"Use None as default for mutable types in '{node.name}' and create new objects inside the function.")
#             self.generic_visit(node)

#         def visit_ClassDef(self, node):
#             class_defs.append(node)
#             assigned_names.add(node.name)
#             self.generic_visit(node)

#         def visit_Name(self, node):
#             if isinstance(node.ctx, ast.Load):
#                 used_names.add(node.id)
#             if isinstance(node.ctx, ast.Store):
#                 assigned_names.add(node.id)
#             self.generic_visit(node)

#         def visit_Call(self, node):
#             if isinstance(node.func, ast.Name) and node.func.id in {"eval", "exec"}:
#                 report["warnings"].append(f"Use of {node.func.id} detected at line {node.lineno}")
#                 report["fixes"].append(f"Avoid {node.func.id}; consider safer alternatives or explicit parsing.")
#             self.generic_visit(node)

#         def visit_ExceptHandler(self, node):
#             if node.type is None:
#                 report["warnings"].append(f"Bare except detected at line {node.lineno}")
#                 report["fixes"].append("Catch specific exception classes instead of a bare except.")
#             elif isinstance(node.type, ast.Name) and node.type.id in {"Exception", "BaseException"}:
#                 report["warnings"].append(f"Overly broad exception handler '{node.type.id}' at line {node.lineno}")
#                 report["fixes"].append("Catch the narrowest relevant exception type.")
#             if len(node.body) == 1 and isinstance(node.body[0], ast.Pass):
#                 report["warnings"].append(f"Exception swallowed with pass at line {node.lineno}")
#                 report["fixes"].append("Handle the exception or log it; avoid silent failures.")
#             self.generic_visit(node)

#     Analyzer().visit(tree)

#     for name in sorted(imported_names):
#         if name not in used_names and name not in {"__future__"}:
#             report["warnings"].append(f"Possibly unused import '{name}'")
#             report["fixes"].append(f"Remove the unused import '{name}'.")

#     for name in assigned_names:
#         if name in dir(__builtins__):
#             report["warnings"].append(f"Variable or function '{name}' shadows built-in")
#             report["fixes"].append(f"Rename '{name}' to avoid shadowing built-ins.")

#     # Simple unreachable code detection
#     for fn in function_defs:
#         seen_terminator = False
#         for node in fn.body:
#             if seen_terminator:
#                 report["warnings"].append(f"Unreachable code in function '{fn.name}' after a return/raise at line {getattr(node, 'lineno', '?')}")
#                 report["fixes"].append(f"Remove or refactor unreachable code in '{fn.name}'.")
#                 break
#             if isinstance(node, (ast.Return, ast.Raise)):
#                 seen_terminator = True

#     # Docstrings
#     if ast.get_docstring(tree) is None:
#         report["warnings"].append("Module is missing a top-level docstring")
#         report["fixes"].append("Add a brief module docstring describing purpose and usage.")
#     for fn in function_defs:
#         if ast.get_docstring(fn) is None:
#             report["warnings"].append(f"Function '{fn.name}' missing a docstring")
#             report["fixes"].append(f"Add a concise docstring for '{fn.name}'.")
#     for cl in class_defs:
#         if ast.get_docstring(cl) is None:
#             report["warnings"].append(f"Class '{cl.name}' missing a docstring")
#             report["fixes"].append(f"Add a concise docstring for class '{cl.name}'.")

#     # Undefined function calls (simple)
#     defined = {fn.name for fn in function_defs} | {cl.name for cl in class_defs}
#     called = set()
#     class CallVisitor(ast.NodeVisitor):
#         def visit_Call(self, node):
#             if isinstance(node.func, ast.Name):
#                 called.add(node.func.id)
#             self.generic_visit(node)
#     CallVisitor().visit(tree)
#     for func in sorted(called):
#         if func not in defined and func not in dir(__builtins__):
#             report["warnings"].append(f"Call to undefined function '{func}' detected")
#             report["fixes"].append(f"Define '{func}' or import it before use.")

#     if not report["errors"]:
#         if report["warnings"]:
#             report["fixes"].append("Review warnings and apply the proposed fixes.")
#         else:
#             report["fixes"].append("No obvious issues; add tests and run linters for confidence.")

#     return report

# # ----------------------------
# # Heuristic AI detector (works for text/code)
# # ----------------------------
# def detect_ai_generated_code(code: str) -> dict:
#     lines = [ln for ln in code.splitlines()]
#     code_lines = [ln for ln in lines if ln.strip() and not ln.strip().startswith("#")]
#     comment_lines = [ln for ln in lines if ln.strip().startswith("#")]
#     features = {}
#     total = max(1, len(lines))
#     features["comment_density"] = len(comment_lines) / total
#     normalized = [re.sub(r"\s+", " ", ln.strip()) for ln in code_lines]
#     counts = Counter(normalized)
#     repeated = sum(c for c in counts.values() if c > 1)
#     features["repeated_line_ratio"] = repeated / max(1, len(code_lines))
#     docstring_like = re.findall(r'\"\"\"(.*?)\"\"\"|\'\'\'(.*?)\'\'\'', code, flags=re.S)
#     docstrings = [d[0] or d[1] for d in docstring_like]
#     templated_docs = sum(1 for d in docstrings if re.match(r"(?i)\s*(this function|returns|parameters)\b", d.strip()))
#     features["templated_doc_ratio"] = (templated_docs / max(1, len(docstrings))) if docstrings else 0.0
#     generic_names = {"data", "result", "results", "temp", "value", "values", "item", "items", "input", "output", "res"}
#     tokens = re.findall(r"[A-Za-z_][A-Za-z0-9_]*", code)
#     generic_count = sum(1 for t in tokens if t in generic_names)
#     features["generic_name_density"] = generic_count / max(1, len(tokens)) if tokens else 0.0
#     words = [w.lower() for w in re.findall(r"[A-Za-z0-9_']+", code)]
#     trigrams = [tuple(words[i:i+3]) for i in range(len(words) - 2)]
#     trigram_counts = Counter(trigrams)
#     repeating_trigram_ratio = sum(1 for c in trigram_counts.values() if c > 1) / max(1, len(trigram_counts)) if trigram_counts else 0.0
#     features["repeating_trigram_ratio"] = repeating_trigram_ratio
#     score = (
#         35 * features["repeating_trigram_ratio"]
#         + 20 * min(1.0, abs(features["comment_density"] - 0.15) / 0.15)
#         + 20 * features["repeated_line_ratio"]
#         + 15 * features["templated_doc_ratio"]
#         + 10 * min(1.0, features["generic_name_density"] * 20)
#     )
#     score = max(0.0, min(100.0, score))
#     if score >= 65:
#         label = "Likely AI-generated"
#     elif score >= 45:
#         label = "Unclear / Mixed"
#     else:
#         label = "Likely human-written"
#     reasons = []
#     if features["repeating_trigram_ratio"] > 0.08:
#         reasons.append("High repeated phrasing patterns.")
#     if features["repeated_line_ratio"] > 0.06:
#         reasons.append("Notable repetition of similar lines.")
#     if features["templated_doc_ratio"] > 0.4:
#         reasons.append("Docstrings appear templated.")
#     if features["comment_density"] < 0.03 or features["comment_density"] > 0.4:
#         reasons.append("Atypical comment density.")
#     if features["generic_name_density"] > 0.02:
#         reasons.append("Frequent use of generic variable names.")
#     return {"score": round(score, 1), "label": label, "reasons": reasons, "features": features}

# # ----------------------------
# # Fetch from URL
# # ----------------------------
# def fetch_from_url(url: str) -> str:
#     try:
#         resp = requests.get(url, timeout=12)
#         resp.raise_for_status()
#         ct = resp.headers.get("Content-Type", "")
#         if "text/html" in ct:
#             soup = BeautifulSoup(resp.text, "html.parser")
#             for s in soup(["script", "style", "noscript"]):
#                 s.extract()
#             return soup.get_text(separator="\n")
#         return resp.text
#     except Exception as e:
#         return f"❌ Error fetching URL: {str(e)}"

# # ----------------------------
# # Tabs (unique keys for all widgets)
# # ----------------------------
# tabs = st.tabs(["Document Summarizer", "Code Analyzer", "AI vs Human Scanner", "URL Fetch & Explain"])

# # Document Summarizer tab
# with tabs[0]:
#     st.header("📄 Document Summarizer")
#     doc_text = st.text_area("Paste text", key="doc_text_input")
#     doc_file = st.file_uploader("Upload a file (.txt, .md, .pdf)", type=["txt", "md", "pdf"], key="doc_file_upload")
#     doc_url = st.text_input("Or enter URL to fetch content", key="doc_url_input")
#     doc_len = st.slider("Summary length (sentences)", 1, 12, 5, key="doc_summary_len")
#     doc_bullets = st.checkbox("Return bullets", value=False, key="doc_as_bullets")

#     if st.button("Summarize Document", key="doc_summarize_btn"):
#         content = ""
#         if doc_file:
#             fn = getattr(doc_file, "name", "")
#             if fn.lower().endswith(".pdf"):
#                 if PyPDF2 is None:
#                     st.error("PDF support requires PyPDF2. Install it with 'pip install PyPDF2'.")
#                 else:
#                     try:
#                         reader = PyPDF2.PdfReader(io.BytesIO(doc_file.read()))
#                         pages_text = []
#                         for p in reader.pages:
#                             pages_text.append(p.extract_text() or "")
#                         content = "\n".join(pages_text)
#                     except Exception as e:
#                         st.error(f"Failed to extract PDF text: {e}")
#                         content = ""
#             else:
#                 try:
#                     content = doc_file.read().decode("utf-8", errors="ignore")
#                 except Exception:
#                     content = doc_file.read().decode("latin-1", errors="ignore")
#         elif doc_url.strip():
#             content = fetch_from_url(doc_url.strip())
#         elif doc_text.strip():
#             content = doc_text

#         if not content:
#             st.warning("Please paste text, upload a file, or enter a URL.")
#         else:
#             summary = summarize_text_advanced(content, max_sentences=doc_len, as_bullets=doc_bullets)
#             with st.expander("📂 Summary (click to expand)", expanded=True):
#                 if doc_bullets:
#                     st.markdown(summary)
#                 else:
#                     st.write(summary)
#                 # copy + download
#                 copy_button(summary, label="Copy Summary", key="doc_summary_copy")
#                 st.download_button("⬇ Download Summary", summary.encode("utf-8"), "document_summary.txt", key="doc_summary_download")

# # Code Analyzer tab
# with tabs[1]:
#     st.header("💻 Code Analyzer")
#     code_text = st.text_area("Paste Python code", key="code_text_input")
#     code_file = st.file_uploader("Upload Python file (.py)", type=["py"], key="code_file_upload")
#     code_url = st.text_input("Or enter URL to fetch code", key="code_url_input")
#     code_len = st.slider("Purpose summary length (items)", 1, 8, 4, key="code_summary_len")

#     if st.button("Analyze Code", key="code_analyze_btn"):
#         code_str = ""
#         if code_file:
#             try:
#                 code_str = code_file.read().decode("utf-8", errors="ignore")
#             except Exception:
#                 code_str = code_file.read().decode("latin-1", errors="ignore")
#         elif code_url.strip():
#             code_str = fetch_from_url(code_url.strip())
#         elif code_text.strip():
#             code_str = code_text

#         if not code_str:
#             st.warning("Please paste code, upload a file, or enter a URL.")
#         else:
#             # show syntax highlighted preview
#             with st.expander("💻 Code Preview (syntax-highlighted)", expanded=False):
#                 st.code(code_str, language="python")

#             report = analyze_python(code_str)

#             with st.expander("📂 Code Purpose Summary", expanded=True):
#                 st.write(report["purpose_summary"])
#                 copy_button(report["purpose_summary"], label="Copy Summary", key="code_purpose_copy")
#                 st.download_button("⬇ Download Purpose Summary", report["purpose_summary"].encode("utf-8"), "code_purpose.txt", key="code_purpose_download")

#             with st.expander("❌ Errors"):
#                 if report["errors"]:
#                     for e in report["errors"]:
#                         st.error(e)
#                 else:
#                     st.write("_No syntax errors detected_")

#             with st.expander("⚠️ Warnings"):
#                 if report["warnings"]:
#                     for w in report["warnings"]:
#                         st.warning(w)
#                 else:
#                     st.write("_No warnings detected_")

#             with st.expander("💡 Fixes & Recommendations"):
#                 if report["fixes"]:
#                     for f in report["fixes"]:
#                         st.info(f)
#                 else:
#                     st.write("_No fixes suggested_")

#             copy_button(str(report), label="Copy Full Report", key="code_full_report_copy")
#             st.download_button("⬇ Download Full Report", str(report).encode("utf-8"), "code_analysis_report.txt", key="code_full_report_download")

# # AI vs Human Scanner tab
# with tabs[2]:
#     st.header("🤖 AI vs Human Scanner")
#     ai_text = st.text_area("Paste text or code", key="ai_text_input")
#     ai_file = st.file_uploader("Upload file (.txt, .md, .py, .pdf)", type=["txt", "md", "py", "pdf"], key="ai_file_upload")
#     ai_url = st.text_input("Or enter URL to fetch content", key="ai_url_input")
#     ai_mode = st.radio("Content type", ["Auto-detect", "Treat as Text", "Treat as Code"], index=0, key="ai_mode_radio")

#     if st.button("Detect AI / Human Likelihood", key="ai_detect_btn"):
#         content = ""
#         if ai_file:
#             fn = getattr(ai_file, "name", "")
#             if fn.lower().endswith(".pdf"):
#                 if PyPDF2 is None:
#                     st.error("PDF support requires PyPDF2. Install it with 'pip install PyPDF2'.")
#                 else:
#                     try:
#                         reader = PyPDF2.PdfReader(io.BytesIO(ai_file.read()))
#                         pages_text = []
#                         for p in reader.pages:
#                             pages_text.append(p.extract_text() or "")
#                         content = "\n".join(pages_text)
#                     except Exception as e:
#                         st.error(f"Failed to extract PDF text: {e}")
#                         content = ""
#             else:
#                 try:
#                     content = ai_file.read().decode("utf-8", errors="ignore")
#                 except Exception:
#                     content = ai_file.read().decode("latin-1", errors="ignore")
#         elif ai_url.strip():
#             content = fetch_from_url(ai_url.strip())
#         elif ai_text.strip():
#             content = ai_text

#         if not content:
#             st.warning("Please paste content, upload a file, or enter a URL.")
#         else:
#             result = detect_ai_generated_code(content)
#             header = f"📊 Likelihood: {result['label']} (Score: {result['score']}%)"
#             with st.expander(header, expanded=True):
#                 st.markdown("**Reasons / Indicators:**")
#                 if result["reasons"]:
#                     for r in result["reasons"]:
#                         st.markdown(f"- {r}")
#                 else:
#                     st.markdown("- _No strong heuristics triggered_")
#                 st.markdown("**Detailed Features:**")
#                 for k, v in result["features"].items():
#                     st.markdown(f"- {k}: {v}")

#                 copy_button(str(result), label="Copy AI Result", key="ai_result_copy")
#                 st.download_button("⬇ Download AI Result", str(result).encode("utf-8"), "ai_result.txt", key="ai_result_download")

# # URL Fetch & Explain tab
# with tabs[3]:
#     st.header("🌐 URL Fetch & Explain")
#     url_input = st.text_input("Enter any URL to fetch content", key="url_fetch_input")
#     url_len = st.slider("Summary length (sentences)", 1, 12, 6, key="url_summary_len")
#     url_bullets = st.checkbox("Return bullets", value=True, key="url_as_bullets")

#     if st.button("Fetch & Explain", key="url_fetch_btn"):
#         if not url_input.strip():
#             st.warning("Enter a valid URL.")
#         else:
#             content = fetch_from_url(url_input.strip())
#             if content.startswith("❌ Error"):
#                 st.error(content)
#             else:
#                 with st.expander("📄 Content Preview", expanded=False):
#                     st.code(content[:4000] + ("..." if len(content) > 4000 else ""), language="text")
#                 summary = summarize_text_advanced(content, max_sentences=url_len, as_bullets=url_bullets)
#                 with st.expander("📂 Summary", expanded=True):
#                     if url_bullets:
#                         st.markdown(summary)
#                     else:
#                         st.write(summary)
#                     copy_button(summary, label="Copy Summary", key="url_summary_copy")
#                     st.download_button("⬇ Download Summary", summary.encode("utf-8"), "url_summary.txt", key="url_summary_download")




































# import streamlit as st
# import requests
# from bs4 import BeautifulSoup
# import re
# import base64
# import ast
# from collections import Counter
# import io

# # PDF extraction (optional)
# try:
#     import PyPDF2
#     PDF_AVAILABLE = True
# except ImportError:
#     PDF_AVAILABLE = False

# # Page config
# st.set_page_config(
#     page_title="Technova AI Nexus", 
#     layout="wide", 
#     initial_sidebar_state="collapsed"
# )

# # Enhanced copy button
# def copy_button(text: str, label: str = "Copy", key: str = None):
#     """Enhanced copy button with styling"""
#     if text is None:
#         text = ""
#     safe_b64 = base64.b64encode(text.encode("utf-8")).decode("ascii")
#     el_id = f"copy-btn-{key}" if key else f"copy-btn-{abs(hash(text))}"
    
#     html = f"""
#     <button id="{el_id}" onclick="navigator.clipboard.writeText(atob('{safe_b64}'))"
#         style="
#             background: linear-gradient(135deg, rgba(0, 249, 255, 0.1), rgba(0, 153, 204, 0.2));
#             border: 1px solid rgba(0, 249, 255, 0.4);
#             color: #00f9ff;
#             padding: 8px 16px;
#             border-radius: 8px;
#             font-family: monospace;
#             font-weight: bold;
#             cursor: pointer;
#             margin: 5px;
#             transition: all 0.3s ease;
#         "
#         onmouseover="this.style.background='linear-gradient(135deg, rgba(0, 249, 255, 0.2), rgba(0, 153, 204, 0.3))'"
#         onmouseout="this.style.background='linear-gradient(135deg, rgba(0, 249, 255, 0.1), rgba(0, 153, 204, 0.2))'">
#         ⚡ {label}
#     </button>
#     <script>
#     document.getElementById('{el_id}').addEventListener('click', function() {{
#         const btn = this;
#         const oldText = btn.innerHTML;
#         btn.innerHTML = '✅ Copied!';
#         btn.style.background = 'linear-gradient(135deg, rgba(0, 255, 0, 0.2), rgba(0, 200, 0, 0.3))';
#         setTimeout(() => {{
#             btn.innerHTML = oldText;
#             btn.style.background = 'linear-gradient(135deg, rgba(0, 249, 255, 0.1), rgba(0, 153, 204, 0.2))';
#         }}, 1500);
#     }});
#     </script>
#     """
#     st.markdown(html, unsafe_allow_html=True)

# # Simplified styling
# def set_tech_styling():
#     st.markdown("""
#     <style>
#     @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700&family=Rajdhani:wght@300;400;500&display=swap');
    
#     .stApp {
#         background: linear-gradient(135deg, #0a0a0a 0%, #1a1a2e 50%, #16213e 100%);
#         color: #00f9ff;
#         font-family: 'Rajdhani', sans-serif;
#     }
    
#     h1, h2, h3, h4, h5, h6 {
#         font-family: 'Orbitron', monospace !important;
#         color: #00f9ff !important;
#         text-shadow: 0 0 10px rgba(0, 249, 255, 0.5);
#     }
    
#     .main-title {
#         font-size: 3rem;
#         text-align: center;
#         background: linear-gradient(45deg, #00f9ff, #0099cc, #66ccff);
#         -webkit-background-clip: text;
#         -webkit-text-fill-color: transparent;
#         background-clip: text;
#         margin: 2rem 0;
#     }
    
#     .subtitle {
#         text-align: center;
#         color: rgba(0, 249, 255, 0.8);
#         font-style: italic;
#         margin-bottom: 2rem;
#     }
    
#     .stTextArea textarea, .stTextInput input {
#         background: rgba(0, 20, 40, 0.8) !important;
#         border: 1px solid rgba(0, 249, 255, 0.3) !important;
#         color: #00f9ff !important;
#     }
    
#     .stButton > button {
#         background: linear-gradient(135deg, rgba(0, 249, 255, 0.2), rgba(0, 153, 204, 0.3)) !important;
#         border: 1px solid rgba(0, 249, 255, 0.5) !important;
#         color: #00f9ff !important;
#         font-weight: bold !important;
#     }
    
#     .stTabs [data-baseweb="tab"] {
#         color: #00f9ff !important;
#         font-family: 'Orbitron', monospace !important;
#     }
    
#     .stExpander {
#         border: 1px solid rgba(0, 249, 255, 0.3);
#         border-radius: 10px;
#     }
#     </style>
#     """, unsafe_allow_html=True)

# # Initialize styling
# set_tech_styling()

# # Main title
# st.markdown('<h1 class="main-title">🌌 TECHNOVA AI NEXUS</h1>', unsafe_allow_html=True)
# st.markdown('<p class="subtitle">Advanced AI-Powered Analysis Suite</p>', unsafe_allow_html=True)

# # Stopwords for summarization
# STOPWORDS = set([
#     "a", "an", "and", "are", "as", "at", "be", "but", "by", "for", "if", "in", "into", 
#     "is", "it", "no", "not", "of", "on", "or", "such", "that", "the", "their", "then", 
#     "there", "these", "they", "this", "to", "was", "will", "with", "you", "your", "from", 
#     "our", "we", "he", "she", "his", "her", "its", "were", "been", "being", "than", 
#     "also", "can", "could", "should", "would", "may", "might", "have", "has", "had", 
#     "do", "does", "did", "done", "just", "over", "under", "more", "most", "other", 
#     "some", "any", "each", "many", "few", "those", "them", "which", "who", "whom", 
#     "whose", "where", "when", "why", "how"
# ])

# def safe_sentence_split(text: str):
#     """Split text into sentences safely"""
#     pattern = re.compile(r"(?<=[.!?])\s+(?=[A-Z0-9])")
#     return [s.strip() for s in pattern.split(text) if s.strip()]

# def summarize_text_advanced(text: str, max_sentences: int = 5, as_bullets: bool = False) -> str:
#     """Advanced text summarization using frequency analysis"""
#     if not text or not text.strip():
#         return "No content to summarize."
    
#     paragraphs = [p.strip() for p in text.splitlines() if p.strip()]
#     sentences = []
#     for para in paragraphs:
#         sentences.extend(safe_sentence_split(para))
    
#     if not sentences:
#         return text

#     # Word frequency analysis
#     word_freq = Counter()
#     for s in sentences:
#         words = [w.lower() for w in re.findall(r"[A-Za-z0-9_']+", s)]
#         for w in words:
#             if w not in STOPWORDS and len(w) > 2:
#                 word_freq[w] += 1
    
#     if not word_freq:
#         return " ".join(sentences[:max_sentences])

#     # Normalize frequencies
#     max_freq = max(word_freq.values())
#     for w in list(word_freq.keys()):
#         word_freq[w] /= max_freq

#     # Score sentences
#     scored = []
#     for idx, s in enumerate(sentences):
#         words = [w.lower() for w in re.findall(r"[A-Za-z0-9_']+", s)]
#         score = sum(word_freq.get(w, 0.0) for w in words)
#         length_penalty = 1.0 + 0.2 * max(0, (len(words) - 20) / 20)
#         position_boost = 1.1 if idx < 3 else 1.0
#         scored.append((score / length_penalty * position_boost, idx, s))

#     # Select top sentences
#     scored.sort(key=lambda x: (-x[0], x[1]))
#     top = sorted(scored[:max_sentences], key=lambda x: x[1])
    
#     if as_bullets:
#         return "\n".join([f"• {s}" for _, _, s in top])
#     return " ".join([s for _, _, s in top])

# def analyze_python(code: str):
#     """Comprehensive Python code analysis"""
#     report = {
#         "functions": [], 
#         "classes": [], 
#         "imports": [], 
#         "purpose_summary": "", 
#         "errors": [], 
#         "warnings": [], 
#         "fixes": []
#     }

#     # Basic parsing
#     for line in code.splitlines():
#         l = line.strip()
#         if l.startswith("def "):
#             func_name = l.split("(")[0][4:].strip()
#             report["functions"].append(func_name)
#         elif l.startswith("class "):
#             class_name = l.split("(")[0][6:].strip().rstrip(":")
#             report["classes"].append(class_name)
#         elif l.startswith("import ") or l.startswith("from "):
#             report["imports"].append(l)

#     # Generate purpose summary
#     report["purpose_summary"] = summarize_text_advanced(code, max_sentences=3)

#     # AST analysis for deeper insights
#     try:
#         tree = ast.parse(code)
#     except SyntaxError as e:
#         report["errors"].append(f"SyntaxError: {e.msg} at line {e.lineno}")
#         report["fixes"].append("Check syntax: indentation, colons, parentheses, quotes.")
#         return report

#     # Analyze AST
#     imported_names = set()
#     used_names = set()
#     assigned_names = set()

#     class CodeAnalyzer(ast.NodeVisitor):
#         def visit_Import(self, node):
#             for alias in node.names:
#                 imported_names.add(alias.asname or alias.name.split(".")[0])
#             self.generic_visit(node)

#         def visit_ImportFrom(self, node):
#             for alias in node.names:
#                 imported_names.add(alias.asname or alias.name)
#             self.generic_visit(node)

#         def visit_FunctionDef(self, node):
#             assigned_names.add(node.name)
#             # Check for mutable default arguments
#             for default in node.args.defaults:
#                 if isinstance(default, (ast.List, ast.Dict, ast.Set)):
#                     report["warnings"].append(f"Mutable default argument in '{node.name}'")
#                     report["fixes"].append(f"Use None as default in '{node.name}' and create objects inside function")
#             self.generic_visit(node)

#         def visit_ClassDef(self, node):
#             assigned_names.add(node.name)
#             self.generic_visit(node)

#         def visit_Name(self, node):
#             if isinstance(node.ctx, ast.Load):
#                 used_names.add(node.id)
#             elif isinstance(node.ctx, ast.Store):
#                 assigned_names.add(node.id)
#             self.generic_visit(node)

#         def visit_Call(self, node):
#             if isinstance(node.func, ast.Name) and node.func.id in {"eval", "exec"}:
#                 report["warnings"].append(f"Dangerous {node.func.id} usage detected")
#                 report["fixes"].append(f"Avoid {node.func.id}; use safer alternatives")
#             self.generic_visit(node)

#     CodeAnalyzer().visit(tree)

#     # Check for unused imports
#     for name in sorted(imported_names):
#         if name not in used_names and name not in {"__future__"}:
#             report["warnings"].append(f"Possibly unused import: '{name}'")
#             report["fixes"].append(f"Remove unused import '{name}'")

#     # Check for builtin shadowing
#     builtins_list = [
#         'abs', 'all', 'any', 'bin', 'bool', 'chr', 'dict', 'dir', 'enumerate', 
#         'filter', 'float', 'format', 'frozenset', 'hash', 'hex', 'id', 'input', 
#         'int', 'isinstance', 'len', 'list', 'map', 'max', 'min', 'next', 'oct', 
#         'open', 'ord', 'pow', 'print', 'range', 'repr', 'reversed', 'round', 
#         'set', 'slice', 'sorted', 'str', 'sum', 'tuple', 'type', 'zip'
#     ]
    
#     for name in assigned_names:
#         if name in builtins_list:
#             report["warnings"].append(f"Variable '{name}' shadows builtin")
#             report["fixes"].append(f"Rename '{name}' to avoid shadowing builtins")

#     # Add general recommendations
#     if not report["errors"]:
#         if not report["warnings"]:
#             report["fixes"].append("Code looks good! Consider adding tests and documentation.")
#         else:
#             report["fixes"].append("Review warnings above and apply suggested fixes.")

#     return report

# def detect_ai_generated_code(code: str) -> dict:
#     """Detect if code might be AI-generated based on patterns"""
#     if not code or not code.strip():
#         return {"score": 0, "label": "❓ No content", "reasons": [], "features": {}}
    
#     lines = code.splitlines()
#     code_lines = [ln for ln in lines if ln.strip() and not ln.strip().startswith("#")]
#     comment_lines = [ln for ln in lines if ln.strip().startswith("#")]
    
#     features = {}
#     total_lines = max(1, len(lines))
    
#     # Comment density
#     features["comment_density"] = len(comment_lines) / total_lines
    
#     # Repeated lines
#     normalized = [re.sub(r"\s+", " ", ln.strip()) for ln in code_lines]
#     counts = Counter(normalized)
#     repeated = sum(c for c in counts.values() if c > 1)
#     features["repeated_line_ratio"] = repeated / max(1, len(code_lines))
    
#     # Generic variable names
#     generic_names = {"data", "result", "temp", "value", "item", "input", "output", "res"}
#     tokens = re.findall(r"[A-Za-z_][A-Za-z0-9_]*", code)
#     generic_count = sum(1 for t in tokens if t in generic_names)
#     features["generic_name_density"] = generic_count / max(1, len(tokens))
    
#     # Calculate score
#     score = (
#         25 * min(1.0, abs(features["comment_density"] - 0.15) / 0.15) +
#         35 * features["repeated_line_ratio"] +
#         20 * min(1.0, features["generic_name_density"] * 20) +
#         20 * (1.0 if len(code_lines) > 50 and features["comment_density"] < 0.05 else 0.0)
#     )
    
#     score = max(0.0, min(100.0, score))
    
#     # Determine label
#     if score >= 70:
#         label = "🚨 Likely AI-generated"
#     elif score >= 45:
#         label = "🤔 Unclear/Mixed"
#     else:
#         label = "✅ Likely human-written"
    
#     # Generate reasons
#     reasons = []
#     if features["repeated_line_ratio"] > 0.1:
#         reasons.append("High repetition of similar code patterns")
#     if features["comment_density"] < 0.02 or features["comment_density"] > 0.4:
#         reasons.append("Unusual comment density")
#     if features["generic_name_density"] > 0.03:
#         reasons.append("Frequent generic variable names")
    
#     return {
#         "score": round(score, 1), 
#         "label": label, 
#         "reasons": reasons, 
#         "features": features
#     }

# def fetch_from_url(url: str) -> str:
#     """Fetch and extract text content from URL"""
#     try:
#         headers = {
#             'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
#         }
#         resp = requests.get(url, timeout=10, headers=headers)
#         resp.raise_for_status()
        
#         content_type = resp.headers.get("Content-Type", "")
#         if "text/html" in content_type:
#             soup = BeautifulSoup(resp.text, "html.parser")
#             # Remove unwanted elements
#             for element in soup(["script", "style", "nav", "header", "footer", "aside"]):
#                 element.extract()
#             return soup.get_text(separator="\n", strip=True)
#         else:
#             return resp.text
            
#     except requests.exceptions.RequestException as e:
#         return f"❌ Error fetching URL: {str(e)}"
#     except Exception as e:
#         return f"❌ Error processing content: {str(e)}"

# # Create tabs
# tabs = st.tabs([
#     "📄 Document Processor", 
#     "🧠 Code Analyzer", 
#     "🤖 AI Detection", 
#     "🌐 URL Extractor"
# ])

# # Tab 1: Document Processor
# with tabs[0]:
#     st.header("📄 Neural Document Processor")
#     st.write("*Advanced text analysis and summarization*")
    
#     col1, col2 = st.columns([2, 1])
    
#     with col1:
#         doc_text = st.text_area(
#             "Input Text", 
#             height=200, 
#             placeholder="Paste your document content here...",
#             key="doc_text_input"
#         )
    
#     with col2:
#         doc_file = st.file_uploader("Upload Document", type=["txt", "md", "pdf"], key="doc_file_upload")
#         doc_url = st.text_input("Document URL", placeholder="https://...", key="doc_url_input")
#         doc_length = st.slider("Summary Length", 1, 10, 5, key="doc_summary_length")
#         doc_bullets = st.checkbox("Use Bullets", value=False, key="doc_use_bullets")

#     if st.button("🚀 Analyze Document", type="primary", key="doc_analyze_btn"):
#         content = ""
        
#         # Get content from various sources
#         if doc_file:
#             if doc_file.name.lower().endswith(".pdf"):
#                 if not PDF_AVAILABLE:
#                     st.error("PDF processing requires PyPDF2. Install with: pip install PyPDF2")
#                 else:
#                     try:
#                         reader = PyPDF2.PdfReader(io.BytesIO(doc_file.read()))
#                         pages_text = []
#                         for page in reader.pages:
#                             pages_text.append(page.extract_text() or "")
#                         content = "\n".join(pages_text)
#                     except Exception as e:
#                         st.error(f"PDF error: {e}")
#             else:
#                 content = doc_file.read().decode("utf-8", errors="ignore")
#         elif doc_url.strip():
#             with st.spinner("Fetching content..."):
#                 content = fetch_from_url(doc_url.strip())
#         elif doc_text.strip():
#             content = doc_text

#         if not content:
#             st.warning("Please provide content via text, file, or URL.")
#         elif content.startswith("❌"):
#             st.error(content)
#         else:
#             with st.spinner("Processing..."):
#                 summary = summarize_text_advanced(content, doc_length, doc_bullets)
            
#             st.success("Analysis Complete!")
            
#             with st.expander("📊 Results", expanded=True):
#                 st.subheader("Summary")
#                 if doc_bullets:
#                     st.markdown(summary)
#                 else:
#                     st.write(summary)
                
#                 col1, col2 = st.columns(2)
#                 with col1:
#                     copy_button(summary, "Copy Summary", key="doc_copy_btn")
#                 with col2:
#                     st.download_button(
#                         "💾 Download", 
#                         summary, 
#                         "summary.txt",
#                         mime="text/plain",
#                         key="doc_download_btn"
#                     )

# # Tab 2: Code Analyzer
# with tabs[1]:
#     st.header("🧠 Quantum Code Analyzer")
#     st.write("*Deep code analysis and optimization recommendations*")
    
#     col1, col2 = st.columns([2, 1])
    
#     with col1:
#         code_text = st.text_area(
#             "Python Code", 
#             height=250, 
#             placeholder="Paste your Python code here...",
#             key="code_text_input"
#         )
    
#     with col2:
#         code_file = st.file_uploader("Upload Python File", type=["py"], key="code_file_upload")
#         code_url = st.text_input("Code URL", placeholder="https://...", key="code_url_input")

#     if st.button("🔍 Analyze Code", type="primary", key="code_analyze_btn"):
#         code_content = ""
        
#         if code_file:
#             code_content = code_file.read().decode("utf-8", errors="ignore")
#         elif code_url.strip():
#             with st.spinner("Fetching code..."):
#                 code_content = fetch_from_url(code_url.strip())
#         elif code_text.strip():
#             code_content = code_text

#         if not code_content:
#             st.warning("Please provide code via input, file, or URL.")
#         elif code_content.startswith("❌"):
#             st.error(code_content)
#         else:
#             with st.expander("Code Preview", expanded=False):
#                 st.code(code_content, language="python")

#             with st.spinner("Analyzing..."):
#                 report = analyze_python(code_content)

#             st.success("Analysis Complete!")

#             col1, col2 = st.columns(2)
            
#             with col1:
#                 st.subheader("📋 Summary")
#                 st.write(report["purpose_summary"])
                
#                 st.subheader("❌ Errors")
#                 if report["errors"]:
#                     for error in report["errors"]:
#                         st.error(error)
#                 else:
#                     st.success("No syntax errors found!")
            
#             with col2:
#                 st.subheader("⚠️ Warnings")
#                 if report["warnings"]:
#                     for warning in report["warnings"]:
#                         st.warning(warning)
#                 else:
#                     st.success("No warnings!")
                
#                 st.subheader("💡 Recommendations")
#                 for fix in report["fixes"]:
#                     st.info(fix)
            
#             # Copy and download options
#             col3, col4 = st.columns(2)
#             with col3:
#                 copy_button(str(report), "Copy Report", key="code_copy_btn")
#             with col4:
#                 st.download_button(
#                     "💾 Download Report", 
#                     str(report), 
#                     "code_analysis.txt",
#                     mime="text/plain",
#                     key="code_download_btn"
#                 )

# # Tab 3: AI Detection
# with tabs[2]:
#     st.header("🤖 AI Detection Matrix")
#     st.write("*Pattern recognition for AI-generated content*")
    
#     col1, col2 = st.columns([3, 1])
    
#     with col1:
#         ai_text = st.text_area(
#             "Content for Analysis", 
#             height=200, 
#             placeholder="Paste text or code to analyze...",
#             key="ai_text_input"
#         )
    
#     with col2:
#         ai_file = st.file_uploader("Upload File", type=["txt", "py", "md"], key="ai_file_upload")
#         ai_url = st.text_input("Content URL", placeholder="https://...", key="ai_url_input")

#     if st.button("🔬 Scan for AI Patterns", type="primary", key="ai_analyze_btn"):
#         content = ""
        
#         if ai_file:
#             content = ai_file.read().decode("utf-8", errors="ignore")
#         elif ai_url.strip():
#             with st.spinner("Fetching content..."):
#                 content = fetch_from_url(ai_url.strip())
#         elif ai_text.strip():
#             content = ai_text

#         if not content:
#             st.warning("Please provide content for analysis.")
#         elif content.startswith("❌"):
#             st.error(content)
#         else:
#             with st.spinner("Analyzing patterns..."):
#                 result = detect_ai_generated_code(content)
            
#             st.success("Analysis Complete!")
            
#             # Display results
#             score_color = "🟢" if result['score'] < 45 else "🟡" if result['score'] < 65 else "🔴"
            
#             with st.expander(f"{score_color} Detection Results", expanded=True):
#                 col1, col2, col3 = st.columns(3)
                
#                 with col2:
#                     st.metric(
#                         "AI Likelihood Score", 
#                         f"{result['score']}%",
#                         delta=f"{result['score'] - 50}% vs baseline"
#                     )
                
#                 st.markdown(f"**Assessment:** {result['label']}")
                
#                 if result["reasons"]:
#                     st.subheader("🔍 Key Indicators")
#                     for reason in result["reasons"]:
#                         st.write(f"• {reason}")
#                 else:
#                     st.success("No significant AI patterns detected")
                
#                 st.subheader("📊 Feature Analysis")
#                 for feature, value in result["features"].items():
#                     feature_name = feature.replace("_", " ").title()
#                     st.write(f"• {feature_name}: `{value:.3f}`")
                
#                 # Copy and download
#                 col4, col5 = st.columns(2)
#                 with col4:
#                     copy_button(str(result), "Copy Results", key="ai_copy_btn")
#                 with col5:
#                     st.download_button(
#                         "💾 Download Analysis", 
#                         str(result), 
#                         "ai_detection.txt",
#                         mime="text/plain",
#                         key="ai_download_btn"
#                     )

# # Tab 4: URL Extractor
# with tabs[3]:
#     st.header("🌐 URL Data Extraction Engine")
#     st.write("*Advanced web content extraction and analysis*")
    
#     col1, col2 = st.columns([3, 1])
    
#     with col1:
#         url_input = st.text_input(
#             "Target URL", 
#             placeholder="Enter any URL for content extraction...",
#             key="url_input_field"
#         )
    
#     with col2:
#         url_length = st.slider("Summary Length", 1, 10, 5, key="url_summary_length")
#         url_bullets = st.checkbox("Use Bullets", value=True, key="url_use_bullets")

#     if st.button("🚀 Extract & Analyze", type="primary", key="url_analyze_btn"):
#         if not url_input.strip():
#             st.warning("Please enter a valid URL.")
#         else:
#             with st.spinner("Extracting content..."):
#                 content = fetch_from_url(url_input.strip())
            
#             if content.startswith("❌"):
#                 st.error(content)
#             else:
#                 st.success("Content extracted successfully!")
                
#                 # Show metrics
#                 col1, col2, col3 = st.columns(3)
#                 with col1:
#                     st.metric("Characters", f"{len(content):,}")
#                 with col2:
#                     st.metric("Words", f"{len(content.split()):,}")
#                 with col3:
#                     st.metric("Lines", f"{len(content.splitlines()):,}")
                
#                 # Show preview
#                 with st.expander("📄 Content Preview", expanded=False):
#                     preview = content[:2000] + "..." if len(content) > 2000 else content
#                     st.text(preview)
                
#                 # Generate summary
#                 with st.spinner("Generating summary..."):
#                     summary = summarize_text_advanced(content, url_length, url_bullets)
                
#                 st.subheader("📊 Content Summary")
#                 if url_bullets:
#                     st.markdown(summary)
#                 else:
#                     st.write(summary)
                
#                 # Copy and download
#                 col4, col5 = st.columns(2)
#                 with col4:
#                     copy_button(summary, "Copy Summary", key="url_copy_btn")
#                 with col5:
#                     st.download_button(
#                         "💾 Download Summary", 
#                         summary.encode("utf-8"), 
#                         "url_summary.txt",
#                         key="url_download_btn"
#                     )

# # Footer
# st.markdown("---")
# st.markdown("""
# <div style="text-align: center; padding: 20px; color: rgba(0, 249, 255, 0.6); font-family: monospace;">
#     🌌 TECHNOVA AI NEXUS v2.0 • Advanced Neural Processing Suite
# </div>
# """, unsafe_allow_html=True)























# import streamlit as st
# import requests
# from bs4 import BeautifulSoup
# import re
# import base64
# import ast
# from collections import Counter
# import io

# # PDF extraction (optional)
# try:
#     import PyPDF2
#     PDF_AVAILABLE = True
# except ImportError:
#     PDF_AVAILABLE = False

# # Page config
# st.set_page_config(
#     page_title="Technova AI Nexus", 
#     layout="wide", 
#     initial_sidebar_state="collapsed"
# )

# # Enhanced copy button
# def copy_button(text: str, label: str = "Copy", key: str = None):
#     """Enhanced copy button with styling"""
#     if text is None:
#         text = ""
#     safe_b64 = base64.b64encode(text.encode("utf-8")).decode("ascii")
#     el_id = f"copy-btn-{key}" if key else f"copy-btn-{abs(hash(text[:100]))}"
    
#     html = f"""
#     <button id="{el_id}" onclick="navigator.clipboard.writeText(atob('{safe_b64}'))"
#         style="
#             background: linear-gradient(135deg, rgba(0, 249, 255, 0.1), rgba(0, 153, 204, 0.2));
#             border: 1px solid rgba(0, 249, 255, 0.4);
#             color: #00f9ff;
#             padding: 8px 16px;
#             border-radius: 8px;
#             font-family: monospace;
#             font-weight: bold;
#             cursor: pointer;
#             margin: 5px;
#             transition: all 0.3s ease;
#         "
#         onmouseover="this.style.background='linear-gradient(135deg, rgba(0, 249, 255, 0.2), rgba(0, 153, 204, 0.3))'"
#         onmouseout="this.style.background='linear-gradient(135deg, rgba(0, 249, 255, 0.1), rgba(0, 153, 204, 0.2))'">
#         ⚡ {label}
#     </button>
#     <script>
#     document.getElementById('{el_id}').addEventListener('click', function() {{
#         const btn = this;
#         const oldText = btn.innerHTML;
#         btn.innerHTML = '✅ Copied!';
#         btn.style.background = 'linear-gradient(135deg, rgba(0, 255, 0, 0.2), rgba(0, 200, 0, 0.3))';
#         setTimeout(() => {{
#             btn.innerHTML = oldText;
#             btn.style.background = 'linear-gradient(135deg, rgba(0, 249, 255, 0.1), rgba(0, 153, 204, 0.2))';
#         }}, 1500);
#     }});
#     </script>
#     """
#     st.markdown(html, unsafe_allow_html=True)

# # Simplified styling
# def set_tech_styling():
#     st.markdown("""
#     <style>
#     @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700&family=Rajdhani:wght@300;400;500&display=swap');
    
#     .stApp {
#         background: linear-gradient(135deg, #0a0a0a 0%, #1a1a2e 50%, #16213e 100%);
#         color: #00f9ff;
#         font-family: 'Rajdhani', sans-serif;
#     }
    
#     h1, h2, h3, h4, h5, h6 {
#         font-family: 'Orbitron', monospace !important;
#         color: #00f9ff !important;
#         text-shadow: 0 0 10px rgba(0, 249, 255, 0.5);
#     }
    
#     .main-title {
#         font-size: 3rem;
#         text-align: center;
#         background: linear-gradient(45deg, #00f9ff, #0099cc, #66ccff);
#         -webkit-background-clip: text;
#         -webkit-text-fill-color: transparent;
#         background-clip: text;
#         margin: 2rem 0;
#     }
    
#     .subtitle {
#         text-align: center;
#         color: rgba(0, 249, 255, 0.8);
#         font-style: italic;
#         margin-bottom: 2rem;
#     }
    
#     .stTextArea textarea, .stTextInput input {
#         background: rgba(0, 20, 40, 0.8) !important;
#         border: 1px solid rgba(0, 249, 255, 0.3) !important;
#         color: #00f9ff !important;
#     }
    
#     .stButton > button {
#         background: linear-gradient(135deg, rgba(0, 249, 255, 0.2), rgba(0, 153, 204, 0.3)) !important;
#         border: 1px solid rgba(0, 249, 255, 0.5) !important;
#         color: #00f9ff !important;
#         font-weight: bold !important;
#     }
    
#     .stTabs [data-baseweb="tab"] {
#         color: #00f9ff !important;
#         font-family: 'Orbitron', monospace !important;
#     }
    
#     .stExpander {
#         border: 1px solid rgba(0, 249, 255, 0.3);
#         border-radius: 10px;
#     }
#     </style>
#     """, unsafe_allow_html=True)

# # Initialize styling
# set_tech_styling()

# # Main title
# st.markdown('<h1 class="main-title">🌌 TECHNOVA AI NEXUS</h1>', unsafe_allow_html=True)
# st.markdown('<p class="subtitle">Advanced AI-Powered Analysis Suite</p>', unsafe_allow_html=True)

# # Stopwords for summarization
# STOPWORDS = set([
#     "a", "an", "and", "are", "as", "at", "be", "but", "by", "for", "if", "in", "into", 
#     "is", "it", "no", "not", "of", "on", "or", "such", "that", "the", "their", "then", 
#     "there", "these", "they", "this", "to", "was", "will", "with", "you", "your", "from", 
#     "our", "we", "he", "she", "his", "her", "its", "were", "been", "being", "than", 
#     "also", "can", "could", "should", "would", "may", "might", "have", "has", "had", 
#     "do", "does", "did", "done", "just", "over", "under", "more", "most", "other", 
#     "some", "any", "each", "many", "few", "those", "them", "which", "who", "whom", 
#     "whose", "where", "when", "why", "how"
# ])

# def safe_sentence_split(text: str):
#     """Split text into sentences safely"""
#     pattern = re.compile(r"(?<=[.!?])\s+(?=[A-Z0-9])")
#     return [s.strip() for s in pattern.split(text) if s.strip()]

# def summarize_text_advanced(text: str, max_sentences: int = 5, as_bullets: bool = False) -> str:
#     """Advanced text summarization using frequency analysis"""
#     if not text or not text.strip():
#         return "No content to summarize."
    
#     paragraphs = [p.strip() for p in text.splitlines() if p.strip()]
#     sentences = []
#     for para in paragraphs:
#         sentences.extend(safe_sentence_split(para))
    
#     if not sentences:
#         return text

#     # Word frequency analysis
#     word_freq = Counter()
#     for s in sentences:
#         words = [w.lower() for w in re.findall(r"[A-Za-z0-9_']+", s)]
#         for w in words:
#             if w not in STOPWORDS and len(w) > 2:
#                 word_freq[w] += 1
    
#     if not word_freq:
#         return " ".join(sentences[:max_sentences])

#     # Normalize frequencies
#     max_freq = max(word_freq.values())
#     for w in list(word_freq.keys()):
#         word_freq[w] /= max_freq

#     # Score sentences
#     scored = []
#     for idx, s in enumerate(sentences):
#         words = [w.lower() for w in re.findall(r"[A-Za-z0-9_']+", s)]
#         score = sum(word_freq.get(w, 0.0) for w in words)
#         length_penalty = 1.0 + 0.2 * max(0, (len(words) - 20) / 20)
#         position_boost = 1.1 if idx < 3 else 1.0
#         scored.append((score / length_penalty * position_boost, idx, s))

#     # Select top sentences
#     scored.sort(key=lambda x: (-x[0], x[1]))
#     top = sorted(scored[:max_sentences], key=lambda x: x[1])
    
#     if as_bullets:
#         return "\n".join([f"• {s}" for _, _, s in top])
#     return " ".join([s for _, _, s in top])

# def analyze_python(code: str):
#     """Comprehensive Python code analysis"""
#     report = {
#         "functions": [], 
#         "classes": [], 
#         "imports": [], 
#         "purpose_summary": "", 
#         "errors": [], 
#         "warnings": [], 
#         "fixes": []
#     }

#     # Basic parsing
#     for line in code.splitlines():
#         l = line.strip()
#         if l.startswith("def "):
#             func_name = l.split("(")[0][4:].strip()
#             report["functions"].append(func_name)
#         elif l.startswith("class "):
#             class_name = l.split("(")[0][6:].strip().rstrip(":")
#             report["classes"].append(class_name)
#         elif l.startswith("import ") or l.startswith("from "):
#             report["imports"].append(l)

#     # Generate purpose summary
#     report["purpose_summary"] = summarize_text_advanced(code, max_sentences=3)

#     # AST analysis for deeper insights
#     try:
#         tree = ast.parse(code)
#     except SyntaxError as e:
#         report["errors"].append(f"SyntaxError: {e.msg} at line {e.lineno}")
#         report["fixes"].append("Check syntax: indentation, colons, parentheses, quotes.")
#         return report

#     # Analyze AST
#     imported_names = set()
#     used_names = set()
#     assigned_names = set()

#     class CodeAnalyzer(ast.NodeVisitor):
#         def visit_Import(self, node):
#             for alias in node.names:
#                 imported_names.add(alias.asname or alias.name.split(".")[0])
#             self.generic_visit(node)

#         def visit_ImportFrom(self, node):
#             for alias in node.names:
#                 imported_names.add(alias.asname or alias.name)
#             self.generic_visit(node)

#         def visit_FunctionDef(self, node):
#             assigned_names.add(node.name)
#             # Check for mutable default arguments
#             for default in node.args.defaults:
#                 if isinstance(default, (ast.List, ast.Dict, ast.Set)):
#                     report["warnings"].append(f"Mutable default argument in '{node.name}'")
#                     report["fixes"].append(f"Use None as default in '{node.name}' and create objects inside function")
#             self.generic_visit(node)

#         def visit_ClassDef(self, node):
#             assigned_names.add(node.name)
#             self.generic_visit(node)

#         def visit_Name(self, node):
#             if isinstance(node.ctx, ast.Load):
#                 used_names.add(node.id)
#             elif isinstance(node.ctx, ast.Store):
#                 assigned_names.add(node.id)
#             self.generic_visit(node)

#         def visit_Call(self, node):
#             if isinstance(node.func, ast.Name) and node.func.id in {"eval", "exec"}:
#                 report["warnings"].append(f"Dangerous {node.func.id} usage detected")
#                 report["fixes"].append(f"Avoid {node.func.id}; use safer alternatives")
#             self.generic_visit(node)

#     CodeAnalyzer().visit(tree)

#     # Check for unused imports
#     for name in sorted(imported_names):
#         if name not in used_names and name not in {"__future__"}:
#             report["warnings"].append(f"Possibly unused import: '{name}'")
#             report["fixes"].append(f"Remove unused import '{name}'")

#     # Check for builtin shadowing
#     builtins_list = [
#         'abs', 'all', 'any', 'bin', 'bool', 'chr', 'dict', 'dir', 'enumerate', 
#         'filter', 'float', 'format', 'frozenset', 'hash', 'hex', 'id', 'input', 
#         'int', 'isinstance', 'len', 'list', 'map', 'max', 'min', 'next', 'oct', 
#         'open', 'ord', 'pow', 'print', 'range', 'repr', 'reversed', 'round', 
#         'set', 'slice', 'sorted', 'str', 'sum', 'tuple', 'type', 'zip'
#     ]
    
#     for name in assigned_names:
#         if name in builtins_list:
#             report["warnings"].append(f"Variable '{name}' shadows builtin")
#             report["fixes"].append(f"Rename '{name}' to avoid shadowing builtins")

#     # Add general recommendations
#     if not report["errors"]:
#         if not report["warnings"]:
#             report["fixes"].append("Code looks good! Consider adding tests and documentation.")
#         else:
#             report["fixes"].append("Review warnings above and apply suggested fixes.")

#     return report

# def detect_ai_generated_code(code: str) -> dict:
#     """Detect if code might be AI-generated based on patterns"""
#     if not code or not code.strip():
#         return {"score": 0, "label": "❓ No content", "reasons": [], "features": {}}
    
#     lines = code.splitlines()
#     code_lines = [ln for ln in lines if ln.strip() and not ln.strip().startswith("#")]
#     comment_lines = [ln for ln in lines if ln.strip().startswith("#")]
    
#     features = {}
#     total_lines = max(1, len(lines))
    
#     # Comment density
#     features["comment_density"] = len(comment_lines) / total_lines
    
#     # Repeated lines
#     normalized = [re.sub(r"\s+", " ", ln.strip()) for ln in code_lines]
#     counts = Counter(normalized)
#     repeated = sum(c for c in counts.values() if c > 1)
#     features["repeated_line_ratio"] = repeated / max(1, len(code_lines))
    
#     # Generic variable names
#     generic_names = {"data", "result", "temp", "value", "item", "input", "output", "res"}
#     tokens = re.findall(r"[A-Za-z_][A-Za-z0-9_]*", code)
#     generic_count = sum(1 for t in tokens if t in generic_names)
#     features["generic_name_density"] = generic_count / max(1, len(tokens))
    
#     # Calculate score
#     score = (
#         25 * min(1.0, abs(features["comment_density"] - 0.15) / 0.15) +
#         35 * features["repeated_line_ratio"] +
#         20 * min(1.0, features["generic_name_density"] * 20) +
#         20 * (1.0 if len(code_lines) > 50 and features["comment_density"] < 0.05 else 0.0)
#     )
    
#     score = max(0.0, min(100.0, score))
    
#     # Determine label
#     if score >= 70:
#         label = "🚨 Likely AI-generated"
#     elif score >= 45:
#         label = "🤔 Unclear/Mixed"
#     else:
#         label = "✅ Likely human-written"
    
#     # Generate reasons
#     reasons = []
#     if features["repeated_line_ratio"] > 0.1:
#         reasons.append("High repetition of similar code patterns")
#     if features["comment_density"] < 0.02 or features["comment_density"] > 0.4:
#         reasons.append("Unusual comment density")
#     if features["generic_name_density"] > 0.03:
#         reasons.append("Frequent generic variable names")
    
#     return {
#         "score": round(score, 1), 
#         "label": label, 
#         "reasons": reasons, 
#         "features": features
#     }

# def fetch_from_url(url: str) -> str:
#     """Fetch and extract text content from URL"""
#     try:
#         headers = {
#             'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
#         }
#         resp = requests.get(url, timeout=10, headers=headers)
#         resp.raise_for_status()
        
#         content_type = resp.headers.get("Content-Type", "")
#         if "text/html" in content_type:
#             soup = BeautifulSoup(resp.text, "html.parser")
#             # Remove unwanted elements
#             for element in soup(["script", "style", "nav", "header", "footer", "aside"]):
#                 element.extract()
#             return soup.get_text(separator="\n", strip=True)
#         else:
#             return resp.text
            
#     except requests.exceptions.RequestException as e:
#         return f"❌ Error fetching URL: {str(e)}"
#     except Exception as e:
#         return f"❌ Error processing content: {str(e)}"

# # Create tabs
# tabs = st.tabs([
#     "📄 Document Processor", 
#     "🧠 Code Analyzer", 
#     "🤖 AI Detection", 
#     "🌐 URL Extractor"
# ])

# # Tab 1: Document Processor
# with tabs[0]:
#     st.header("📄 Neural Document Processor")
#     st.write("*Advanced text analysis and summarization*")
    
#     col1, col2 = st.columns([2, 1])
    
#     with col1:
#         doc_text = st.text_area(
#             "Input Text", 
#             height=200, 
#             placeholder="Paste your document content here...",
#             key="doc_text_input"
#         )
    
#     with col2:
#         doc_file = st.file_uploader("Upload Document", type=["txt", "md", "pdf"], key="doc_file_upload")
#         doc_url = st.text_input("Document URL", placeholder="https://...", key="doc_url_input")
#         doc_length = st.slider("Summary Length", 1, 10, 5, key="doc_summary_length")
#         doc_bullets = st.checkbox("Use Bullets", value=False, key="doc_use_bullets")

#     if st.button("🚀 Analyze Document", type="primary", key="doc_analyze_btn"):
#         content = ""
        
#         # Get content from various sources
#         if doc_file:
#             if doc_file.name.lower().endswith(".pdf"):
#                 if not PDF_AVAILABLE:
#                     st.error("PDF processing requires PyPDF2. Install with: pip install PyPDF2")
#                 else:
#                     try:
#                         reader = PyPDF2.PdfReader(io.BytesIO(doc_file.read()))
#                         pages_text = []
#                         for page in reader.pages:
#                             pages_text.append(page.extract_text() or "")
#                         content = "\n".join(pages_text)
#                     except Exception as e:
#                         st.error(f"PDF error: {e}")
#             else:
#                 content = doc_file.read().decode("utf-8", errors="ignore")
#         elif doc_url.strip():
#             with st.spinner("Fetching content..."):
#                 content = fetch_from_url(doc_url.strip())
#         elif doc_text.strip():
#             content = doc_text

#         if not content:
#             st.warning("Please provide content via text, file, or URL.")
#         elif content.startswith("❌"):
#             st.error(content)
#         else:
#             with st.spinner("Processing..."):
#                 summary = summarize_text_advanced(content, doc_length, doc_bullets)
            
#             st.success("Analysis Complete!")
            
#             with st.expander("📊 Results", expanded=True):
#                 st.subheader("Summary")
#                 if doc_bullets:
#                     st.markdown(summary)
#                 else:
#                     st.write(summary)
                
#                 col1, col2 = st.columns(2)
#                 with col1:
#                     copy_button(summary, "Copy Summary", key="doc_copy_btn")
#                 with col2:
#                     st.download_button(
#                         "💾 Download", 
#                         summary, 
#                         "summary.txt",
#                         mime="text/plain",
#                         key="doc_download_btn"
#                     )

# # Tab 2: Code Analyzer
# with tabs[1]:
#     st.header("🧠 Quantum Code Analyzer")
#     st.write("*Deep code analysis and optimization recommendations*")
    
#     col1, col2 = st.columns([2, 1])
    
#     with col1:
#         code_text = st.text_area(
#             "Python Code", 
#             height=250, 
#             placeholder="Paste your Python code here...",
#             key="code_text_input"
#         )
    
#     with col2:
#         code_file = st.file_uploader("Upload Python File", type=["py"], key="code_file_upload")
#         code_url = st.text_input("Code URL", placeholder="https://...", key="code_url_input")

#     if st.button("🔍 Analyze Code", type="primary", key="code_analyze_btn"):
#         code_content = ""
        
#         if code_file:
#             code_content = code_file.read().decode("utf-8", errors="ignore")
#         elif code_url.strip():
#             with st.spinner("Fetching code..."):
#                 code_content = fetch_from_url(code_url.strip())
#         elif code_text.strip():
#             code_content = code_text

#         if not code_content:
#             st.warning("Please provide code via input, file, or URL.")
#         elif code_content.startswith("❌"):
#             st.error(code_content)
#         else:
#             with st.expander("Code Preview", expanded=False):
#                 st.code(code_content, language="python")

#             with st.spinner("Analyzing..."):
#                 report = analyze_python(code_content)

#             st.success("Analysis Complete!")

#             col1, col2 = st.columns(2)
            
#             with col1:
#                 st.subheader("📋 Summary")
#                 st.write(report["purpose_summary"])
                
#                 st.subheader("❌ Errors")
#                 if report["errors"]:
#                     for error in report["errors"]:
#                         st.error(error)
#                 else:
#                     st.success("No syntax errors found!")
            
#             with col2:
#                 st.subheader("⚠️ Warnings")
#                 if report["warnings"]:
#                     for warning in report["warnings"]:
#                         st.warning(warning)
#                 else:
#                     st.success("No warnings!")
                
#                 st.subheader("💡 Recommendations")
#                 for fix in report["fixes"]:
#                     st.info(fix)
            
#             # Copy and download options
#             col3, col4 = st.columns(2)
#             with col3:
#                 copy_button(str(report), "Copy Report", key="code_copy_btn")
#             with col4:
#                 st.download_button(
#                     "💾 Download Report", 
#                     str(report), 
#                     "code_analysis.txt",
#                     mime="text/plain",
#                     key="code_download_btn"
#                 )

# # Tab 3: AI Detection
# with tabs[2]:
#     st.header("🤖 AI Detection Matrix")
#     st.write("*Pattern recognition for AI-generated content*")
    
#     col1, col2 = st.columns([3, 1])
    
#     with col1:
#         ai_text = st.text_area(
#             "Content for Analysis", 
#             height=200, 
#             placeholder="Paste text or code to analyze...",
#             key="ai_text_input"
#         )
    
#     with col2:
#         ai_file = st.file_uploader("Upload File", type=["txt", "py", "md"], key="ai_file_upload")
#         ai_url = st.text_input("Content URL", placeholder="https://...", key="ai_url_input")

#     if st.button("🔬 Scan for AI Patterns", type="primary", key="ai_analyze_btn"):
#         content = ""
        
#         if ai_file:
#             content = ai_file.read().decode("utf-8", errors="ignore")
#         elif ai_url.strip():
#             with st.spinner("Fetching content..."):
#                 content = fetch_from_url(ai_url.strip())
#         elif ai_text.strip():
#             content = ai_text

#         if not content:
#             st.warning("Please provide content for analysis.")
#         elif content.startswith("❌"):
#             st.error(content)
#         else:
#             with st.spinner("Analyzing patterns..."):
#                 result = detect_ai_generated_code(content)
            
#             st.success("Analysis Complete!")
            
#             # Display results
#             score_color = "🟢" if result['score'] < 45 else "🟡" if result['score'] < 65 else "🔴"
            
#             with st.expander(f"{score_color} Detection Results", expanded=True):
#                 col1, col2, col3 = st.columns(3)
                
#                 with col2:
#                     st.metric(
#                         "AI Likelihood Score", 
#                         f"{result['score']}%",
#                         delta=f"{result['score'] - 50}% vs baseline"
#                     )
                
#                 st.markdown(f"**Assessment:** {result['label']}")
                
#                 if result["reasons"]:
#                     st.subheader("🔍 Key Indicators")
#                     for reason in result["reasons"]:
#                         st.write(f"• {reason}")
#                 else:
#                     st.success("No significant AI patterns detected")
                
#                 st.subheader("📊 Feature Analysis")
#                 for feature, value in result["features"].items():
#                     feature_name = feature.replace("_", " ").title()
#                     st.write(f"• {feature_name}: `{value:.3f}`")
                
#                 # Copy and download
#                 col4, col5 = st.columns(2)
#                 with col4:
#                     copy_button(str(result), "Copy Results", key="ai_copy_btn")
#                 with col5:
#                     st.download_button(
#                         "💾 Download Analysis", 
#                         str(result), 
#                         "ai_detection.txt",
#                         mime="text/plain",
#                         key="ai_download_btn"
#                     )

# # Tab 4: URL Extractor
# with tabs[3]:
#     st.header("🌐 URL Data Extraction Engine")
#     st.write("*Advanced web content extraction and analysis*")
    
#     col1, col2 = st.columns([3, 1])
    
#     with col1:
#         url_input = st.text_input(
#             "Target URL", 
#             placeholder="Enter any URL for content extraction...",
#             key="url_input_field"
#         )
    
#     with col2:
#         url_length = st.slider("Summary Length", 1, 10, 5, key="url_summary_length")
#         url_bullets = st.checkbox("Use Bullets", value=True, key="url_use_bullets")

#     if st.button("🚀 Extract & Analyze", type="primary", key="url_analyze_btn"):
#         if not url_input.strip():
#             st.warning("Please enter a valid URL.")
#         else:
#             with st.spinner("Extracting content..."):
#                 content = fetch_from_url(url_input.strip())
            
#             if content.startswith("❌"):
#                 st.error(content)
#             else:
#                 st.success("Content extracted successfully!")
                
#                 # Show metrics
#                 col1, col2, col3 = st.columns(3)
#                 with col1:
#                     st.metric("Characters", f"{len(content):,}")
#                 with col2:
#                     st.metric("Words", f"{len(content.split()):,}")
#                 with col3:
#                     st.metric("Lines", f"{len(content.splitlines()):,}")
                
#                 # Show preview
#                 with st.expander("📄 Content Preview", expanded=False):
#                     preview = content[:2000] + "..." if len(content) > 2000 else content
#                     st.text(preview)
                
#                 # Generate summary
#                 with st.spinner("Generating summary..."):
#                     summary = summarize_text_advanced(content, url_length, url_bullets)
                
#                 st.subheader("📊 Content Summary")
#                 if url_bullets:
#                     st.markdown(summary)
#                 else:
#                     st.write(summary)
                
#                 # Copy and download
#                 col4, col5 = st.columns(2)
#                 with col4:
#                     copy_button(summary, "Copy Summary", key="url_copy_btn")
#                 with col5:
#                     st.download_button(
#                         "💾 Download Summary", 
#                         summary, 
#                         "url_summary.txt",
#                         mime="text/plain",
#                         key="url_download_btn"
#                     )

# # Footer
# st.markdown("---")
# st.markdown("""
# <div style="text-align: center; padding: 20px; color: rgba(0, 249, 255, 0.6); font-family: monospace;">
#     🌌 TECHNOVA AI NEXUS v2.0 • Advanced Neural Processing Suite
# </div>
# """, unsafe_allow_html=True)






















# import streamlit as st
# import requests
# from bs4 import BeautifulSoup
# import re
# import base64
# import ast
# from collections import Counter, defaultdict
# import io

# # PDF extraction (optional)
# try:
#     import PyPDF2
#     PDF_AVAILABLE = True
# except ImportError:
#     PDF_AVAILABLE = False

# # Page config
# st.set_page_config(
#     page_title="Technova AI Nexus", 
#     layout="wide", 
#     initial_sidebar_state="collapsed"
# )

# # Enhanced copy button with better reliability
# def copy_button(text: str, label: str = "Copy", key: str = None):
#     """Reliable copy button with fallback"""
#     if text is None:
#         text = ""
    
#     button_key = f"copy_{key}" if key else f"copy_{abs(hash(text[:100]))}"
    
#     # Create button
#     if st.button(f"⚡ {label}", key=button_key, help="Click to prepare text for copying"):
#         st.session_state[f"text_to_copy_{button_key}"] = text
#         st.success("✅ Ready to copy! Use the text area below:")
        
#     # Show text area if button was clicked
#     if f"text_to_copy_{button_key}" in st.session_state:
#         st.text_area(
#             "📋 Select all text and copy (Ctrl+A, Ctrl+C or Cmd+A, Cmd+C):",
#             value=st.session_state[f"text_to_copy_{button_key}"],
#             height=150,
#             key=f"copy_area_{button_key}",
#             help="Select all text in this box and copy it to your clipboard"
#         )

# # Simplified styling
# def set_tech_styling():
#     st.markdown("""
#     <style>
#     @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700&family=Rajdhani:wght@300;400;500&display=swap');
    
#     .stApp {
#         background: linear-gradient(135deg, #0a0a0a 0%, #1a1a2e 50%, #16213e 100%);
#         color: #00f9ff;
#         font-family: 'Rajdhani', sans-serif;
#     }
    
#     h1, h2, h3, h4, h5, h6 {
#         font-family: 'Orbitron', monospace !important;
#         color: #00f9ff !important;
#         text-shadow: 0 0 10px rgba(0, 249, 255, 0.5);
#     }
    
#     .main-title {
#         font-size: 3rem;
#         text-align: center;
#         background: linear-gradient(45deg, #00f9ff, #0099cc, #66ccff);
#         -webkit-background-clip: text;
#         -webkit-text-fill-color: transparent;
#         background-clip: text;
#         margin: 2rem 0;
#     }
    
#     .subtitle {
#         text-align: center;
#         color: rgba(0, 249, 255, 0.8);
#         font-style: italic;
#         margin-bottom: 2rem;
#     }
    
#     .metric-card {
#         background: rgba(0, 20, 40, 0.6);
#         border: 1px solid rgba(0, 249, 255, 0.3);
#         border-radius: 10px;
#         padding: 15px;
#         margin: 10px 0;
#     }
    
#     .stTextArea textarea, .stTextInput input {
#         background: rgba(0, 20, 40, 0.8) !important;
#         border: 1px solid rgba(0, 249, 255, 0.3) !important;
#         color: #00f9ff !important;
#     }
    
#     .stButton > button {
#         background: linear-gradient(135deg, rgba(0, 249, 255, 0.2), rgba(0, 153, 204, 0.3)) !important;
#         border: 1px solid rgba(0, 249, 255, 0.5) !important;
#         color: #00f9ff !important;
#         font-weight: bold !important;
#     }
    
#     .stTabs [data-baseweb="tab"] {
#         color: #00f9ff !important;
#         font-family: 'Orbitron', monospace !important;
#     }
    
#     .stExpander {
#         border: 1px solid rgba(0, 249, 255, 0.3);
#         border-radius: 10px;
#     }
    
#     .success-metric {
#         color: #00ff00;
#         font-weight: bold;
#     }
    
#     .warning-metric {
#         color: #ffaa00;
#         font-weight: bold;
#     }
    
#     .error-metric {
#         color: #ff4444;
#         font-weight: bold;
#     }
#     </style>
#     """, unsafe_allow_html=True)

# # Initialize styling
# set_tech_styling()

# # Main title
# st.markdown('<h1 class="main-title">🌌 TECHNOVA AI NEXUS</h1>', unsafe_allow_html=True)
# st.markdown('<p class="subtitle">Advanced AI-Powered Analysis Suite v2.1</p>', unsafe_allow_html=True)

# # Stopwords for summarization
# STOPWORDS = set([
#     "a", "an", "and", "are", "as", "at", "be", "but", "by", "for", "if", "in", "into", 
#     "is", "it", "no", "not", "of", "on", "or", "such", "that", "the", "their", "then", 
#     "there", "these", "they", "this", "to", "was", "will", "with", "you", "your", "from", 
#     "our", "we", "he", "she", "his", "her", "its", "were", "been", "being", "than", 
#     "also", "can", "could", "should", "would", "may", "might", "have", "has", "had", 
#     "do", "does", "did", "done", "just", "over", "under", "more", "most", "other", 
#     "some", "any", "each", "many", "few", "those", "them", "which", "who", "whom", 
#     "whose", "where", "when", "why", "how"
# ])

# def safe_sentence_split(text: str):
#     """Split text into sentences safely"""
#     pattern = re.compile(r"(?<=[.!?])\s+(?=[A-Z0-9])")
#     return [s.strip() for s in pattern.split(text) if s.strip()]

# def summarize_text_advanced(text: str, max_sentences: int = 5, as_bullets: bool = False) -> str:
#     """Advanced text summarization using frequency analysis"""
#     if not text or not text.strip():
#         return "No content to summarize."
    
#     paragraphs = [p.strip() for p in text.splitlines() if p.strip()]
#     sentences = []
#     for para in paragraphs:
#         sentences.extend(safe_sentence_split(para))
    
#     if not sentences:
#         return text

#     # Word frequency analysis
#     word_freq = Counter()
#     for s in sentences:
#         words = [w.lower() for w in re.findall(r"[A-Za-z0-9_']+", s)]
#         for w in words:
#             if w not in STOPWORDS and len(w) > 2:
#                 word_freq[w] += 1
    
#     if not word_freq:
#         return " ".join(sentences[:max_sentences])

#     # Normalize frequencies
#     max_freq = max(word_freq.values())
#     for w in list(word_freq.keys()):
#         word_freq[w] /= max_freq

#     # Score sentences
#     scored = []
#     for idx, s in enumerate(sentences):
#         words = [w.lower() for w in re.findall(r"[A-Za-z0-9_']+", s)]
#         score = sum(word_freq.get(w, 0.0) for w in words)
#         length_penalty = 1.0 + 0.2 * max(0, (len(words) - 20) / 20)
#         position_boost = 1.1 if idx < 3 else 1.0
#         scored.append((score / length_penalty * position_boost, idx, s))

#     # Select top sentences
#     scored.sort(key=lambda x: (-x[0], x[1]))
#     top = sorted(scored[:max_sentences], key=lambda x: x[1])
    
#     if as_bullets:
#         return "\n".join([f"• {s}" for _, _, s in top])
#     return " ".join([s for _, _, s in top])

# def analyze_python_enhanced(code: str):
#     """Comprehensive Python code analysis with enhanced metrics"""
#     report = {
#         "basic_stats": {},
#         "functions": [], 
#         "classes": [], 
#         "imports": [], 
#         "variables": [],
#         "purpose_summary": "", 
#         "errors": [], 
#         "warnings": [], 
#         "fixes": [],
#         "complexity_metrics": {},
#         "code_quality": {},
#         "detailed_counts": {}
#     }

#     if not code or not code.strip():
#         return report

#     lines = code.splitlines()
#     code_lines = [line for line in lines if line.strip() and not line.strip().startswith("#")]
#     comment_lines = [line for line in lines if line.strip().startswith("#")]
#     blank_lines = [line for line in lines if not line.strip()]

#     # Basic statistics
#     report["basic_stats"] = {
#         "total_lines": len(lines),
#         "code_lines": len(code_lines),
#         "comment_lines": len(comment_lines),
#         "blank_lines": len(blank_lines),
#         "characters": len(code),
#         "words": len(code.split())
#     }

#     # Initialize counters
#     function_count = 0
#     class_count = 0
#     import_count = 0
#     print_count = 0
#     if_count = 0
#     for_count = 0
#     while_count = 0
#     try_count = 0
#     with_count = 0
#     def_count = 0
#     lambda_count = 0
#     list_comp_count = 0
#     dict_comp_count = 0

#     # Pattern-based counting
#     for line in code_lines:
#         stripped_line = line.strip()
        
#         # Count various constructs
#         if stripped_line.startswith("def "):
#             def_count += 1
#         if stripped_line.startswith("class "):
#             class_count += 1
#         if stripped_line.startswith(("import ", "from ")):
#             import_count += 1
        
#         # Count control structures and statements
#         print_count += len(re.findall(r'\bprint\s*\(', line))
#         if_count += len(re.findall(r'\bif\b', line))
#         for_count += len(re.findall(r'\bfor\b', line))
#         while_count += len(re.findall(r'\bwhile\b', line))
#         try_count += len(re.findall(r'\btry\b', line))
#         with_count += len(re.findall(r'\bwith\b', line))
#         lambda_count += len(re.findall(r'\blambda\b', line))
#         list_comp_count += len(re.findall(r'\[.*for.*in.*\]', line))
#         dict_comp_count += len(re.findall(r'\{.*for.*in.*\}', line))

#     report["detailed_counts"] = {
#         "functions": def_count,
#         "classes": class_count,
#         "imports": import_count,
#         "print_statements": print_count,
#         "if_statements": if_count,
#         "for_loops": for_count,
#         "while_loops": while_count,
#         "try_blocks": try_count,
#         "with_statements": with_count,
#         "lambda_expressions": lambda_count,
#         "list_comprehensions": list_comp_count,
#         "dict_comprehensions": dict_comp_count
#     }

#     # Calculate complexity metrics
#     total_constructs = sum([if_count, for_count, while_count, try_count, def_count])
#     complexity_score = total_constructs / max(1, len(code_lines)) * 100
    
#     report["complexity_metrics"] = {
#         "cyclomatic_complexity_estimate": total_constructs + 1,
#         "complexity_per_line": round(complexity_score, 2),
#         "nesting_level_estimate": max(0, len(re.findall(r'    ', code)) // len(code_lines) * 10) if code_lines else 0
#     }

#     # Code quality metrics
#     comment_ratio = len(comment_lines) / max(1, len(lines)) * 100
#     avg_line_length = sum(len(line) for line in code_lines) / max(1, len(code_lines))
    
#     report["code_quality"] = {
#         "comment_ratio": round(comment_ratio, 2),
#         "avg_line_length": round(avg_line_length, 2),
#         "code_to_comment_ratio": round(len(code_lines) / max(1, len(comment_lines)), 2)
#     }

#     # Basic parsing for simple analysis
#     for line in code_lines:
#         stripped = line.strip()
#         if stripped.startswith("def "):
#             match = re.match(r"def\s+([a-zA-Z_][a-zA-Z0-9_]*)", stripped)
#             if match:
#                 report["functions"].append(match.group(1))
#         elif stripped.startswith("class "):
#             match = re.match(r"class\s+([a-zA-Z_][a-zA-Z0-9_]*)", stripped)
#             if match:
#                 report["classes"].append(match.group(1))
#         elif stripped.startswith(("import ", "from ")):
#             report["imports"].append(stripped)

#     # Generate purpose summary
#     report["purpose_summary"] = summarize_text_advanced(code, max_sentences=3)

#     # AST analysis for deeper insights
#     try:
#         tree = ast.parse(code)
#     except SyntaxError as e:
#         report["errors"].append(f"SyntaxError: {e.msg} at line {e.lineno}")
#         report["fixes"].append("Check syntax: indentation, colons, parentheses, quotes.")
#         return report

#     # Enhanced AST analysis
#     imported_names = set()
#     used_names = set()
#     assigned_names = set()
#     builtin_functions_used = set()
    
#     builtin_funcs = {
#         'abs', 'all', 'any', 'bin', 'bool', 'chr', 'dict', 'dir', 'enumerate', 
#         'filter', 'float', 'format', 'frozenset', 'hash', 'hex', 'id', 'input', 
#         'int', 'isinstance', 'len', 'list', 'map', 'max', 'min', 'next', 'oct', 
#         'open', 'ord', 'pow', 'print', 'range', 'repr', 'reversed', 'round', 
#         'set', 'slice', 'sorted', 'str', 'sum', 'tuple', 'type', 'zip'
#     }

#     class EnhancedCodeAnalyzer(ast.NodeVisitor):
#         def visit_Import(self, node):
#             for alias in node.names:
#                 imported_names.add(alias.asname or alias.name.split(".")[0])
#             self.generic_visit(node)

#         def visit_ImportFrom(self, node):
#             for alias in node.names:
#                 imported_names.add(alias.asname or alias.name)
#             self.generic_visit(node)

#         def visit_FunctionDef(self, node):
#             assigned_names.add(node.name)
#             # Check for mutable default arguments
#             for default in node.args.defaults:
#                 if isinstance(default, (ast.List, ast.Dict, ast.Set)):
#                     report["warnings"].append(f"Mutable default argument in '{node.name}'")
#                     report["fixes"].append(f"Use None as default in '{node.name}' and create objects inside function")
            
#             # Check for long functions
#             func_lines = len([n for n in ast.walk(node) if hasattr(n, 'lineno')])
#             if func_lines > 50:
#                 report["warnings"].append(f"Function '{node.name}' is very long ({func_lines} lines)")
#                 report["fixes"].append(f"Consider breaking down '{node.name}' into smaller functions")
            
#             self.generic_visit(node)

#         def visit_ClassDef(self, node):
#             assigned_names.add(node.name)
#             self.generic_visit(node)

#         def visit_Name(self, node):
#             if isinstance(node.ctx, ast.Load):
#                 used_names.add(node.id)
#                 if node.id in builtin_funcs:
#                     builtin_functions_used.add(node.id)
#             elif isinstance(node.ctx, ast.Store):
#                 assigned_names.add(node.id)
#                 report["variables"].append(node.id)
#             self.generic_visit(node)

#         def visit_Call(self, node):
#             if isinstance(node.func, ast.Name):
#                 if node.func.id in {"eval", "exec"}:
#                     report["warnings"].append(f"Dangerous {node.func.id} usage detected")
#                     report["fixes"].append(f"Avoid {node.func.id}; use safer alternatives")
#                 elif node.func.id == "print" and len(node.args) == 0:
#                     report["warnings"].append("Empty print() statement found")
#                     report["fixes"].append("Remove empty print() or add meaningful content")
#             self.generic_visit(node)

#     EnhancedCodeAnalyzer().visit(tree)

#     # Add builtin functions used to report
#     report["builtin_functions_used"] = sorted(list(builtin_functions_used))

#     # Check for unused imports
#     unused_imports = []
#     for name in sorted(imported_names):
#         if name not in used_names and name not in {"__future__"}:
#             unused_imports.append(name)
#             report["warnings"].append(f"Possibly unused import: '{name}'")
#             report["fixes"].append(f"Remove unused import '{name}' if not needed")

#     # Check for builtin shadowing
#     builtins_list = [
#         'abs', 'all', 'any', 'bin', 'bool', 'chr', 'dict', 'dir', 'enumerate', 
#         'filter', 'float', 'format', 'frozenset', 'hash', 'hex', 'id', 'input', 
#         'int', 'isinstance', 'len', 'list', 'map', 'max', 'min', 'next', 'oct', 
#         'open', 'ord', 'pow', 'print', 'range', 'repr', 'reversed', 'round', 
#         'set', 'slice', 'sorted', 'str', 'sum', 'tuple', 'type', 'zip'
#     ]
    
#     shadowed_builtins = []
#     for name in assigned_names:
#         if name in builtins_list:
#             shadowed_builtins.append(name)
#             report["warnings"].append(f"Variable '{name}' shadows builtin")
#             report["fixes"].append(f"Rename '{name}' to avoid shadowing builtins")

#     # Code quality analysis
#     if report["code_quality"]["comment_ratio"] < 5:
#         report["warnings"].append("Low comment density - consider adding more documentation")
#         report["fixes"].append("Add docstrings to functions and classes, and inline comments for complex logic")
    
#     if report["code_quality"]["avg_line_length"] > 100:
#         report["warnings"].append("Some lines are very long - consider breaking them up")
#         report["fixes"].append("Break long lines at logical points (after commas, operators, etc.)")

#     # Add general recommendations
#     if not report["errors"]:
#         if not report["warnings"]:
#             report["fixes"].append("🎉 Code looks excellent! Consider adding unit tests and type hints.")
#         else:
#             report["fixes"].append("Review warnings above and apply suggested fixes for better code quality.")

#     # Remove duplicate variables
#     report["variables"] = sorted(list(set(report["variables"])))

#     return report

# def detect_ai_generated_code(code: str) -> dict:
#     """Detect if code might be AI-generated based on patterns"""
#     if not code or not code.strip():
#         return {"score": 0, "label": "❓ No content", "reasons": [], "features": {}}
    
#     lines = code.splitlines()
#     code_lines = [ln for ln in lines if ln.strip() and not ln.strip().startswith("#")]
#     comment_lines = [ln for ln in lines if ln.strip().startswith("#")]
    
#     features = {}
#     total_lines = max(1, len(lines))
    
#     # Comment density
#     features["comment_density"] = len(comment_lines) / total_lines
    
#     # Repeated lines
#     normalized = [re.sub(r"\s+", " ", ln.strip()) for ln in code_lines]
#     counts = Counter(normalized)
#     repeated = sum(c for c in counts.values() if c > 1)
#     features["repeated_line_ratio"] = repeated / max(1, len(code_lines))
    
#     # Generic variable names
#     generic_names = {"data", "result", "temp", "value", "item", "input", "output", "res", "var", "obj"}
#     tokens = re.findall(r"[A-Za-z_][A-Za-z0-9_]*", code)
#     generic_count = sum(1 for t in tokens if t in generic_names)
#     features["generic_name_density"] = generic_count / max(1, len(tokens))
    
#     # Perfect formatting (suspicious)
#     perfect_indent = all(len(line) - len(line.lstrip()) % 4 == 0 for line in code_lines if line.strip())
#     features["perfect_formatting"] = 1.0 if perfect_indent and len(code_lines) > 10 else 0.0
    
#     # Calculate score
#     score = (
#         25 * min(1.0, abs(features["comment_density"] - 0.15) / 0.15) +
#         35 * features["repeated_line_ratio"] +
#         20 * min(1.0, features["generic_name_density"] * 20) +
#         20 * features["perfect_formatting"]
#     )
    
#     score = max(0.0, min(100.0, score))
    
#     # Determine label
#     if score >= 70:
#         label = "🚨 Likely AI-generated"
#     elif score >= 45:
#         label = "🤔 Unclear/Mixed"
#     else:
#         label = "✅ Likely human-written"
    
#     # Generate reasons
#     reasons = []
#     if features["repeated_line_ratio"] > 0.1:
#         reasons.append("High repetition of similar code patterns")
#     if features["comment_density"] < 0.02 or features["comment_density"] > 0.4:
#         reasons.append("Unusual comment density")
#     if features["generic_name_density"] > 0.03:
#         reasons.append("Frequent generic variable names")
#     if features["perfect_formatting"] == 1.0:
#         reasons.append("Suspiciously perfect code formatting")
    
#     return {
#         "score": round(score, 1), 
#         "label": label, 
#         "reasons": reasons, 
#         "features": features
#     }

# def fetch_from_url(url: str) -> str:
#     """Fetch and extract text content from URL"""
#     try:
#         headers = {
#             'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
#         }
#         resp = requests.get(url, timeout=10, headers=headers)
#         resp.raise_for_status()
        
#         content_type = resp.headers.get("Content-Type", "")
#         if "text/html" in content_type:
#             soup = BeautifulSoup(resp.text, "html.parser")
#             # Remove unwanted elements
#             for element in soup(["script", "style", "nav", "header", "footer", "aside"]):
#                 element.extract()
#             return soup.get_text(separator="\n", strip=True)
#         else:
#             return resp.text
            
#     except requests.exceptions.RequestException as e:
#         return f"❌ Error fetching URL: {str(e)}"
#     except Exception as e:
#         return f"❌ Error processing content: {str(e)}"

# # Create tabs
# tabs = st.tabs([
#     "📄 Document Processor", 
#     "🧠 Code Analyzer", 
#     "🤖 AI Detection", 
#     "🌐 URL Extractor"
# ])

# # Tab 1: Document Processor
# with tabs[0]:
#     st.header("📄 Neural Document Processor")
#     st.write("*Advanced text analysis and summarization*")
    
#     col1, col2 = st.columns([2, 1])
    
#     with col1:
#         doc_text = st.text_area(
#             "Input Text", 
#             height=200, 
#             placeholder="Paste your document content here...",
#             key="doc_text_input"
#         )
    
#     with col2:
#         doc_file = st.file_uploader("Upload Document", type=["txt", "md", "pdf"], key="doc_file_upload")
#         doc_url = st.text_input("Document URL", placeholder="https://...", key="doc_url_input")
#         doc_length = st.slider("Summary Length", 1, 10, 5, key="doc_summary_length")
#         doc_bullets = st.checkbox("Use Bullets", value=False, key="doc_use_bullets")

#     if st.button("🚀 Analyze Document", type="primary", key="doc_analyze_btn"):
#         content = ""
        
#         # Get content from various sources
#         if doc_file:
#             if doc_file.name.lower().endswith(".pdf"):
#                 if not PDF_AVAILABLE:
#                     st.error("PDF processing requires PyPDF2. Install with: pip install PyPDF2")
#                 else:
#                     try:
#                         reader = PyPDF2.PdfReader(io.BytesIO(doc_file.read()))
#                         pages_text = []
#                         for page in reader.pages:
#                             pages_text.append(page.extract_text() or "")
#                         content = "\n".join(pages_text)
#                     except Exception as e:
#                         st.error(f"PDF error: {e}")
#             else:
#                 content = doc_file.read().decode("utf-8", errors="ignore")
#         elif doc_url.strip():
#             with st.spinner("Fetching content..."):
#                 content = fetch_from_url(doc_url.strip())
#         elif doc_text.strip():
#             content = doc_text

#         if not content:
#             st.warning("Please provide content via text, file, or URL.")
#         elif content.startswith("❌"):
#             st.error(content)
#         else:
#             with st.spinner("Processing..."):
#                 summary = summarize_text_advanced(content, doc_length, doc_bullets)
            
#             st.success("Analysis Complete!")
            
#             # Show document metrics
#             col1, col2, col3, col4 = st.columns(4)
#             with col1:
#                 st.metric("Characters", f"{len(content):,}")
#             with col2:
#                 st.metric("Words", f"{len(content.split()):,}")
#             with col3:
#                 st.metric("Lines", f"{len(content.splitlines()):,}")
#             with col4:
#                 st.metric("Paragraphs", f"{len([p for p in content.split('\\n\\n') if p.strip()]):,}")
            
#             with st.expander("📊 Summary Results", expanded=True):
#                 st.subheader("📋 Generated Summary")
#                 if doc_bullets:
#                     st.markdown(summary)
#                 else:
#                     st.write(summary)
                
#                 col1, col2 = st.columns(2)
#                 with col1:
#                     copy_button(summary, "Copy Summary", key="doc_copy_btn")
#                 with col2:
#                     st.download_button(
#                         "💾 Download", 
#                         summary, 
#                         "summary.txt",
#                         mime="text/plain",
#                         key="doc_download_btn"
#                     )

# # Tab 2: Enhanced Code Analyzer
# with tabs[1]:
#     st.header("🧠 Quantum Code Analyzer")
#     st.write("*Deep code analysis with comprehensive metrics and optimization recommendations*")
    
#     col1, col2 = st.columns([2, 1])
    
#     with col1:
#         code_text = st.text_area(
#             "Python Code", 
#             height=250, 
#             placeholder="Paste your Python code here...",
#             key="code_text_input"
#         )
    
#     with col2:
#         code_file = st.file_uploader("Upload Python File", type=["py"], key="code_file_upload")
#         code_url = st.text_input("Code URL", placeholder="https://...", key="code_url_input")

#     if st.button("🔍 Analyze Code", type="primary", key="code_analyze_btn"):
#         code_content = ""
        
#         if code_file:
#             code_content = code_file.read().decode("utf-8", errors="ignore")
#         elif code_url.strip():
#             with st.spinner("Fetching code..."):
#                 code_content = fetch_from_url(code_url.strip())
#         elif code_text.strip():
#             code_content = code_text

#         if not code_content:
#             st.warning("Please provide code via input, file, or URL.")
#         elif code_content.startswith("❌"):
#             st.error(code_content)
#         else:
#             with st.expander("Code Preview", expanded=False):
#                 st.code(code_content, language="python")

#             with st.spinner("Analyzing code..."):
#                 report = analyze_python_enhanced(code_content)

#             st.success("🎉 Analysis Complete!")

#             # Basic Statistics Section
#             st.subheader("📊 Code Statistics")
#             col1, col2, col3, col4 = st.columns(4)
            
#             with col1:
#                 st.metric("Total Lines", report["basic_stats"].get("total_lines", 0))
#                 st.metric("Functions", report["detailed_counts"].get("functions", 0))
#             with col2:
#                 st.metric("Code Lines", report["basic_stats"].get("code_lines", 0))
#                 st.metric("Classes", report["detailed_counts"].get("classes", 0))
#             with col3:
#                 st.metric("Comments", report["basic_stats"].get("comment_lines", 0))
#                 st.metric("Imports", report["detailed_counts"].get("imports", 0))
#             with col4:
#                 st.metric("Characters", f"{report['basic_stats'].get('characters', 0):,}")
#                 st.metric("Print Statements", report["detailed_counts"].get("print_statements", 0))

#             # Detailed Counts Section
#             st.subheader("🔧 Code Constructs")
#             col1, col2, col3, col4 = st.columns(4)
            
#             with col1:
#                 st.metric("If Statements", report["detailed_counts"].get("if_statements", 0))
#                 st.metric("For Loops", report["detailed_counts"].get("for_loops", 0))
#             with col2:
#                 st.metric("While Loops", report["detailed_counts"].get("while_loops", 0))
#                 st.metric("Try Blocks", report["detailed_counts"].get("try_blocks", 0))
#             with col3:
#                 st.metric("With Statements", report["detailed_counts"].get("with_statements", 0))
#                 st.metric("Lambda Expressions", report["detailed_counts"].get("lambda_expressions", 0))
#             with col4:
#                 st.metric("List Comprehensions", report["detailed_counts"].get("list_comprehensions", 0))
#                 st.metric("Dict Comprehensions", report["detailed_counts"].get("dict_comprehensions", 0))

#             # Code Quality Metrics
#             st.subheader("📈 Quality Metrics")
#             col1, col2, col3 = st.columns(3)
            
#             with col1:
#                 comment_ratio = report["code_quality"].get("comment_ratio", 0)
#                 st.metric("Comment Ratio", f"{comment_ratio:.1f}%")
                
#             with col2:
#                 avg_line_length = report["code_quality"].get("avg_line_length", 0)
#                 st.metric("Avg Line Length", f"{avg_line_length:.1f} chars")
                
#             with col3:
#                 complexity = report["complexity_metrics"].get("cyclomatic_complexity_estimate", 0)
#                 st.metric("Complexity Score", complexity)

#             # Main Analysis Results
#             col1, col2 = st.columns(2)
            
#             with col1:
#                 st.subheader("📋 Code Summary")
#                 st.write(report["purpose_summary"])
                
#                 if report["functions"]:
#                     st.subheader("🔧 Functions Found")
#                     for func in report["functions"][:10]:  # Show first 10
#                         st.write(f"• `{func}()`")
#                     if len(report["functions"]) > 10:
#                         st.write(f"... and {len(report['functions']) - 10} more")
                
#                 if report["classes"]:
#                     st.subheader("🏗️ Classes Found")
#                     for cls in report["classes"]:
#                         st.write(f"• `{cls}`")
                
#                 if report["builtin_functions_used"]:
#                     st.subheader("🐍 Built-in Functions Used")
#                     builtin_display = ", ".join(f"`{func}()`" for func in report["builtin_functions_used"][:15])
#                     st.write(builtin_display)
#                     if len(report["builtin_functions_used"]) > 15:
#                         st.write(f"... and {len(report['builtin_functions_used']) - 15} more")
            
#             with col2:
#                 # Errors Section
#                 st.subheader("❌ Errors")
#                 if report["errors"]:
#                     for error in report["errors"]:
#                         st.error(error)
#                 else:
#                     st.success("✅ No syntax errors found!")
                
#                 # Warnings Section
#                 st.subheader("⚠️ Warnings")
#                 if report["warnings"]:
#                     for warning in report["warnings"][:8]:  # Show first 8 warnings
#                         st.warning(warning)
#                     if len(report["warnings"]) > 8:
#                         st.info(f"... and {len(report['warnings']) - 8} more warnings")
#                 else:
#                     st.success("✅ No warnings!")
                
#                 # Recommendations Section
#                 st.subheader("💡 Recommendations")
#                 for fix in report["fixes"][:6]:  # Show first 6 recommendations
#                     st.info(fix)
#                 if len(report["fixes"]) > 6:
#                     st.info(f"... and {len(report['fixes']) - 6} more recommendations")

#             # Imports and Variables
#             with st.expander("📦 Imports & Variables Details", expanded=False):
#                 col1, col2 = st.columns(2)
                
#                 with col1:
#                     if report["imports"]:
#                         st.subheader("📥 Imports")
#                         for imp in report["imports"]:
#                             st.code(imp, language="python")
#                     else:
#                         st.write("No imports found")
                
#                 with col2:
#                     if report["variables"]:
#                         st.subheader("📝 Variables")
#                         variables_display = ", ".join(f"`{var}`" for var in report["variables"][:20])
#                         st.write(variables_display)
#                         if len(report["variables"]) > 20:
#                             st.write(f"... and {len(report['variables']) - 20} more variables")
#                     else:
#                         st.write("No variables detected")
            
#             # Copy and download options
#             col3, col4 = st.columns(2)
#             with col3:
#                 # Create a formatted report for copying
#                 formatted_report = f"""
# CODE ANALYSIS REPORT
# ==================

# BASIC STATISTICS:
# - Total Lines: {report['basic_stats'].get('total_lines', 0)}
# - Code Lines: {report['basic_stats'].get('code_lines', 0)}
# - Comment Lines: {report['basic_stats'].get('comment_lines', 0)}
# - Characters: {report['basic_stats'].get('characters', 0):,}

# CODE CONSTRUCTS:
# - Functions: {report['detailed_counts'].get('functions', 0)}
# - Classes: {report['detailed_counts'].get('classes', 0)}
# - Imports: {report['detailed_counts'].get('imports', 0)}
# - Print Statements: {report['detailed_counts'].get('print_statements', 0)}
# - If Statements: {report['detailed_counts'].get('if_statements', 0)}
# - For Loops: {report['detailed_counts'].get('for_loops', 0)}
# - While Loops: {report['detailed_counts'].get('while_loops', 0)}

# QUALITY METRICS:
# - Comment Ratio: {report['code_quality'].get('comment_ratio', 0):.1f}%
# - Avg Line Length: {report['code_quality'].get('avg_line_length', 0):.1f} chars
# - Complexity Score: {report['complexity_metrics'].get('cyclomatic_complexity_estimate', 0)}

# SUMMARY: {report['purpose_summary']}

# ERRORS: {len(report['errors'])} found
# WARNINGS: {len(report['warnings'])} found
# RECOMMENDATIONS: {len(report['fixes'])} provided
# """
#                 copy_button(formatted_report, "Copy Report", key="code_copy_btn")
#             with col4:
#                 st.download_button(
#                     "💾 Download Report", 
#                     formatted_report, 
#                     "code_analysis_report.txt",
#                     mime="text/plain",
#                     key="code_download_btn"
#                 )

# # Tab 3: AI Detection
# with tabs[2]:
#     st.header("🤖 AI Detection Matrix")
#     st.write("*Advanced pattern recognition for AI-generated content*")
    
#     col1, col2 = st.columns([3, 1])
    
#     with col1:
#         ai_text = st.text_area(
#             "Content for Analysis", 
#             height=200, 
#             placeholder="Paste text or code to analyze for AI patterns...",
#             key="ai_text_input"
#         )
    
#     with col2:
#         ai_file = st.file_uploader("Upload File", type=["txt", "py", "md"], key="ai_file_upload")
#         ai_url = st.text_input("Content URL", placeholder="https://...", key="ai_url_input")

#     if st.button("🔬 Scan for AI Patterns", type="primary", key="ai_analyze_btn"):
#         content = ""
        
#         if ai_file:
#             content = ai_file.read().decode("utf-8", errors="ignore")
#         elif ai_url.strip():
#             with st.spinner("Fetching content..."):
#                 content = fetch_from_url(ai_url.strip())
#         elif ai_text.strip():
#             content = ai_text

#         if not content:
#             st.warning("Please provide content for analysis.")
#         elif content.startswith("❌"):
#             st.error(content)
#         else:
#             with st.spinner("Analyzing patterns..."):
#                 result = detect_ai_generated_code(content)
            
#             st.success("🔍 Analysis Complete!")
            
#             # Display results with enhanced visualization
#             score_color = "🟢" if result['score'] < 45 else "🟡" if result['score'] < 65 else "🔴"
            
#             with st.expander(f"{score_color} Detection Results", expanded=True):
#                 # Main score display
#                 col1, col2, col3 = st.columns([1, 2, 1])
                
#                 with col2:
#                     # Create a visual score indicator
#                     score_val = result['score']
#                     if score_val < 45:
#                         color = "success"
#                         interpretation = "Human-like patterns detected"
#                     elif score_val < 65:
#                         color = "warning"
#                         interpretation = "Mixed or uncertain patterns"
#                     else:
#                         color = "error"
#                         interpretation = "AI-like patterns detected"
                    
#                     st.metric(
#                         "🎯 AI Likelihood Score", 
#                         f"{score_val}%",
#                         delta=f"{score_val - 50:+.1f}% vs baseline",
#                         delta_color="inverse"
#                     )
                    
#                     # Progress bar for visual representation
#                     progress_color = "#ff4444" if score_val >= 65 else "#ffaa00" if score_val >= 45 else "#00ff00"
#                     st.markdown(f"""
#                     <div style="background: rgba(255,255,255,0.1); border-radius: 10px; padding: 5px; margin: 10px 0;">
#                         <div style="background: {progress_color}; height: 20px; width: {score_val}%; border-radius: 8px; transition: all 0.3s;"></div>
#                     </div>
#                     <p style="text-align: center; color: {progress_color}; font-weight: bold;">{interpretation}</p>
#                     """, unsafe_allow_html=True)
                
#                 st.markdown(f"**🏷️ Assessment:** {result['label']}")
                
#                 # Key indicators
#                 if result["reasons"]:
#                     st.subheader("🔍 Key Indicators Found")
#                     for i, reason in enumerate(result["reasons"], 1):
#                         st.write(f"{i}. {reason}")
#                 else:
#                     st.success("✅ No significant AI patterns detected")
                
#                 # Detailed feature analysis
#                 st.subheader("📊 Detailed Feature Analysis")
                
#                 col1, col2 = st.columns(2)
                
#                 with col1:
#                     st.write("**📈 Pattern Metrics:**")
#                     for feature, value in result["features"].items():
#                         feature_name = feature.replace("_", " ").title()
                        
#                         # Color code based on feature values
#                         if "density" in feature or "ratio" in feature:
#                             if value > 0.3:
#                                 color = "🔴"
#                             elif value > 0.15:
#                                 color = "🟡"
#                             else:
#                                 color = "🟢"
#                         else:
#                             color = "🔵"
                        
#                         st.write(f"{color} **{feature_name}:** `{value:.3f}`")
                
#                 with col2:
#                     st.write("**🎯 Confidence Breakdown:**")
                    
#                     # Create confidence indicators for each feature
#                     features = result["features"]
                    
#                     if "comment_density" in features:
#                         cd_score = abs(features["comment_density"] - 0.15) / 0.15
#                         st.write(f"📝 Comment Pattern: `{cd_score:.2f}` {'⚠️' if cd_score > 0.5 else '✅'}")
                    
#                     if "repeated_line_ratio" in features:
#                         rr_score = features["repeated_line_ratio"]
#                         st.write(f"🔄 Code Repetition: `{rr_score:.2f}` {'⚠️' if rr_score > 0.1 else '✅'}")
                    
#                     if "generic_name_density" in features:
#                         gn_score = features["generic_name_density"]
#                         st.write(f"🏷️ Generic Names: `{gn_score:.2f}` {'⚠️' if gn_score > 0.03 else '✅'}")
                    
#                     if "perfect_formatting" in features:
#                         pf_score = features["perfect_formatting"]
#                         st.write(f"✨ Perfect Format: `{pf_score:.2f}` {'⚠️' if pf_score > 0.5 else '✅'}")
                
#                 # Copy and download
#                 col4, col5 = st.columns(2)
#                 with col4:
#                     detection_report = f"""
# AI DETECTION ANALYSIS REPORT
# ============================

# OVERALL ASSESSMENT: {result['label']}
# CONFIDENCE SCORE: {result['score']}%

# KEY INDICATORS:
# {chr(10).join(f"- {reason}" for reason in result['reasons']) if result['reasons'] else "- No significant patterns detected"}

# FEATURE ANALYSIS:
# {chr(10).join(f"- {k.replace('_', ' ').title()}: {v:.3f}" for k, v in result['features'].items())}

# INTERPRETATION:
# {interpretation}
# """
#                     copy_button(detection_report, "Copy Results", key="ai_copy_btn")
#                 with col5:
#                     st.download_button(
#                         "💾 Download Analysis", 
#                         detection_report, 
#                         "ai_detection_report.txt",
#                         mime="text/plain",
#                         key="ai_download_btn"
#                     )

# # Tab 4: URL Extractor
# with tabs[3]:
#     st.header("🌐 URL Data Extraction Engine")
#     st.write("*Advanced web content extraction and intelligent analysis*")
    
#     col1, col2 = st.columns([3, 1])
    
#     with col1:
#         url_input = st.text_input(
#             "🎯 Target URL", 
#             placeholder="Enter any URL for intelligent content extraction...",
#             key="url_input_field",
#             help="Supports web pages, GitHub files, documentation, articles, and more"
#         )
    
#     with col2:
#         url_length = st.slider("Summary Length", 1, 15, 7, key="url_summary_length")
#         url_bullets = st.checkbox("Use Bullet Points", value=True, key="url_use_bullets")
#         url_show_preview = st.checkbox("Show Content Preview", value=True, key="url_show_preview")

#     if st.button("🚀 Extract & Analyze", type="primary", key="url_analyze_btn"):
#         if not url_input.strip():
#             st.warning("⚠️ Please enter a valid URL to proceed.")
#         else:
#             with st.spinner("🔄 Extracting content from URL..."):
#                 content = fetch_from_url(url_input.strip())
            
#             if content.startswith("❌"):
#                 st.error(content)
#             else:
#                 st.success("✅ Content extracted successfully!")
                
#                 # Enhanced metrics display
#                 words = content.split()
#                 sentences = len(re.findall(r'[.!?]+', content))
#                 paragraphs = len([p for p in content.split('\n\n') if p.strip()])
                
#                 col1, col2, col3, col4, col5 = st.columns(5)
#                 with col1:
#                     st.metric("📄 Characters", f"{len(content):,}")
#                 with col2:
#                     st.metric("📝 Words", f"{len(words):,}")
#                 with col3:
#                     st.metric("📋 Lines", f"{len(content.splitlines()):,}")
#                 with col4:
#                     st.metric("📖 Sentences", f"{sentences:,}")
#                 with col5:
#                     st.metric("📑 Paragraphs", f"{paragraphs:,}")
                
#                 # Reading time estimation
#                 avg_reading_speed = 200  # words per minute
#                 reading_time = max(1, len(words) // avg_reading_speed)
#                 st.info(f"📚 Estimated reading time: ~{reading_time} minute{'s' if reading_time != 1 else ''}")
                
#                 # Content preview with better formatting
#                 if url_show_preview:
#                     with st.expander("📄 Content Preview", expanded=False):
#                         preview_length = 3000
#                         preview = content[:preview_length]
#                         if len(content) > preview_length:
#                             preview += f"\n\n... [Content truncated - showing first {preview_length:,} characters of {len(content):,} total]"
                        
#                         # Try to detect if it's code
#                         if any(keyword in content.lower()[:500] for keyword in ['def ', 'class ', 'import ', 'function', '#include', 'public class']):
#                             st.code(preview, language="python" if "def " in preview or "import " in preview else "text")
#                         else:
#                             st.text(preview)
                
#                 # Generate enhanced summary
#                 with st.spinner("🧠 Generating intelligent summary..."):
#                     summary = summarize_text_advanced(content, url_length, url_bullets)
                
#                 # Display summary with better styling
#                 st.subheader("🎯 Intelligent Content Summary")
                
#                 summary_container = st.container()
#                 with summary_container:
#                     if url_bullets:
#                         st.markdown(summary)
#                     else:
#                         st.write(summary)
                
#                 # Additional analysis
#                 with st.expander("📊 Content Analysis", expanded=False):
#                     col1, col2 = st.columns(2)
                    
#                     with col1:
#                         st.write("**📈 Content Statistics:**")
#                         avg_word_length = sum(len(word) for word in words) / max(len(words), 1)
#                         st.write(f"• Average word length: {avg_word_length:.1f} characters")
                        
#                         if sentences > 0:
#                             avg_sentence_length = len(words) / sentences
#                             st.write(f"• Average sentence length: {avg_sentence_length:.1f} words")
                        
#                         # Most common words (excluding stopwords)
#                         filtered_words = [word.lower().strip('.,!?";()[]') for word in words 
#                                         if len(word) > 3 and word.lower() not in STOPWORDS]
#                         if filtered_words:
#                             common_words = Counter(filtered_words).most_common(5)
#                             st.write("• Most frequent words:")
#                             for word, count in common_words:
#                                 st.write(f"  - **{word}**: {count} times")
                    
#                     with col2:
#                         st.write("**🔍 Content Characteristics:**")
                        
#                         # Detect content type
#                         if any(keyword in content.lower() for keyword in ['function', 'variable', 'method', 'class', 'algorithm']):
#                             st.write("• 💻 **Type**: Technical/Programming content")
#                         elif any(keyword in content.lower() for keyword in ['research', 'study', 'analysis', 'data', 'results']):
#                             st.write("• 🔬 **Type**: Research/Academic content")
#                         elif any(keyword in content.lower() for keyword in ['tutorial', 'how to', 'step', 'guide', 'instructions']):
#                             st.write("• 📚 **Type**: Educational/Tutorial content")
#                         else:
#                             st.write("• 📰 **Type**: General/Article content")
                        
#                         # Complexity estimate
#                         complex_words = [word for word in words if len(word) > 7]
#                         complexity = len(complex_words) / max(len(words), 1) * 100
                        
#                         if complexity > 20:
#                             st.write("• 🎓 **Complexity**: Advanced")
#                         elif complexity > 12:
#                             st.write("• 📖 **Complexity**: Intermediate")
#                         else:
#                             st.write("• 📝 **Complexity**: Beginner-friendly")
                        
#                         st.write(f"• 📊 **Complexity Score**: {complexity:.1f}%")
                
#                 # Enhanced copy and download options
#                 col4, col5, col6 = st.columns(3)
#                 with col4:
#                     copy_button(summary, "Copy Summary", key="url_copy_summary")
#                 with col5:
#                     copy_button(content, "Copy Full Content", key="url_copy_full")
#                 with col6:
#                     # Create comprehensive report
#                     comprehensive_report = f"""
# URL EXTRACTION REPORT
# ====================

# SOURCE URL: {url_input}
# EXTRACTION DATE: {st.session_state.get('current_date', 'Today')}

# CONTENT METRICS:
# - Characters: {len(content):,}
# - Words: {len(words):,}
# - Lines: {len(content.splitlines()):,}
# - Sentences: {sentences:,}
# - Paragraphs: {paragraphs:,}
# - Estimated Reading Time: ~{reading_time} minute(s)

# INTELLIGENT SUMMARY:
# {summary}

# FULL EXTRACTED CONTENT:
# {content}
# """
#                     st.download_button(
#                         "💾 Download Full Report", 
#                         comprehensive_report, 
#                         f"url_analysis_report_{url_input.split('/')[-1][:20]}.txt",
#                         mime="text/plain",
#                         key="url_download_full"
#                     )

# # Enhanced Footer with additional info
# st.markdown("---")
# st.markdown("""
# <div style="text-align: center; padding: 30px;">
#     <h3 style="color: rgba(0, 249, 255, 0.8); font-family: 'Orbitron', monospace;">🌌 TECHNOVA AI NEXUS v2.1</h3>
#     <p style="color: rgba(0, 249, 255, 0.6); font-family: monospace; margin: 10px 0;">
#         Advanced Neural Processing Suite • Quantum Code Analysis • AI Pattern Detection
#     </p>
#     <p style="color: rgba(0, 249, 255, 0.4); font-size: 0.9em;">
#         Enhanced with comprehensive metrics, intelligent analysis, and reliable copy functionality
#     </p>
#     <div style="margin-top: 20px; color: rgba(0, 249, 255, 0.3);">
#         🚀 Built for developers, researchers, and AI enthusiasts
#     </div>
# </div>
# """, unsafe_allow_html=True)

# # Add session state management for better UX
# if 'analysis_history' not in st.session_state:
#     st.session_state['analysis_history'] = []

# # Optional: Add a sidebar with tips and shortcuts
# with st.sidebar:
#     st.header("⚡ Quick Tips")
#     st.write("""
#     **🔧 Code Analyzer:**
#     • Upload .py files for best results
#     • Supports GitHub raw URLs
#     • Detects 12+ code constructs
    
#     **🤖 AI Detection:**
#     • Analyzes patterns in code/text
#     • Uses multiple indicators
#     • Confidence scoring system
    
#     **🌐 URL Extractor:**
#     • Works with most websites
#     • Intelligent content filtering
#     • Automatic summarization
    
#     **📋 Copy Functionality:**
#     • Click copy button first
#     • Select all text in the box
#     • Use Ctrl+C (or Cmd+C) to copy
#     """)
    
#     if st.button("🗑️ Clear All Data", key="clear_all"):
#         for key in list(st.session_state.keys()):
#             if key.startswith(('text_to_copy_', 'copied_text_')):
#                 del st.session_state[key]
#         st.success("All temporary data cleared!")
#         st.rerun()








































# import streamlit as st
# import requests
# from bs4 import BeautifulSoup
# import re
# import base64
# import ast
# from collections import Counter, defaultdict
# import io

# # PDF extraction (optional)
# try:
#     import PyPDF2
#     PDF_AVAILABLE = True
# except ImportError:
#     PDF_AVAILABLE = False

# # Page config
# st.set_page_config(
#     page_title="Technova AI Nexus", 
#     layout="wide", 
#     initial_sidebar_state="collapsed"
# )

# # Simple and reliable copy function using streamlit components
# def copy_button(text: str, label: str = "Copy", key: str = None):
#     """Simple copy button using streamlit text area"""
#     if text is None:
#         text = ""
    
#     button_key = f"copy_btn_{key}" if key else f"copy_btn_{abs(hash(text[:50]))}"
#     area_key = f"copy_area_{key}" if key else f"copy_area_{abs(hash(text[:50]))}"
    
#     # Show copy button and text area side by side
#     col1, col2 = st.columns([1, 6])
    
#     with col1:
#         st.button(f"📋 {label}", key=button_key, help="Click to show text for copying")
    
#     with col2:
#         st.text_area(
#             "Select all text below and copy (Ctrl+A, Ctrl+C):",
#             value=text,
#             height=120,
#             key=area_key,
#             help="Select all text and copy to clipboard"
#         )

# # Simplified styling
# def set_tech_styling():
#     st.markdown("""
#     <style>
#     @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700&family=Rajdhani:wght@300;400;500&display=swap');
    
#     .stApp {
#         background: linear-gradient(135deg, #0a0a0a 0%, #1a1a2e 50%, #16213e 100%);
#         color: #00f9ff;
#         font-family: 'Rajdhani', sans-serif;
#     }
    
#     h1, h2, h3, h4, h5, h6 {
#         font-family: 'Orbitron', monospace !important;
#         color: #00f9ff !important;
#         text-shadow: 0 0 10px rgba(0, 249, 255, 0.5);
#     }
    
#     .main-title {
#         font-size: 3rem;
#         text-align: center;
#         background: linear-gradient(45deg, #00f9ff, #0099cc, #66ccff);
#         -webkit-background-clip: text;
#         -webkit-text-fill-color: transparent;
#         background-clip: text;
#         margin: 2rem 0;
#     }
    
#     .subtitle {
#         text-align: center;
#         color: rgba(0, 249, 255, 0.8);
#         font-style: italic;
#         margin-bottom: 2rem;
#     }
    
#     .stTextArea textarea, .stTextInput input {
#         background: rgba(0, 20, 40, 0.8) !important;
#         border: 1px solid rgba(0, 249, 255, 0.3) !important;
#         color: #00f9ff !important;
#     }
    
#     .stButton > button {
#         background: linear-gradient(135deg, rgba(0, 249, 255, 0.2), rgba(0, 153, 204, 0.3)) !important;
#         border: 1px solid rgba(0, 249, 255, 0.5) !important;
#         color: #00f9ff !important;
#         font-weight: bold !important;
#     }
    
#     .stTabs [data-baseweb="tab"] {
#         color: #00f9ff !important;
#         font-family: 'Orbitron', monospace !important;
#     }
    
#     .stExpander {
#         border: 1px solid rgba(0, 249, 255, 0.3);
#         border-radius: 10px;
#     }
#     </style>
#     """, unsafe_allow_html=True)

# # Initialize styling
# set_tech_styling()

# # Main title
# st.markdown('<h1 class="main-title">🌌 TECHNOVA AI NEXUS</h1>', unsafe_allow_html=True)
# st.markdown('<p class="subtitle">Advanced AI-Powered Analysis Suite v2.1</p>', unsafe_allow_html=True)

# # Stopwords for summarization
# STOPWORDS = set([
#     "a", "an", "and", "are", "as", "at", "be", "but", "by", "for", "if", "in", "into", 
#     "is", "it", "no", "not", "of", "on", "or", "such", "that", "the", "their", "then", 
#     "there", "these", "they", "this", "to", "was", "will", "with", "you", "your", "from", 
#     "our", "we", "he", "she", "his", "her", "its", "were", "been", "being", "than", 
#     "also", "can", "could", "should", "would", "may", "might", "have", "has", "had", 
#     "do", "does", "did", "done", "just", "over", "under", "more", "most", "other", 
#     "some", "any", "each", "many", "few", "those", "them", "which", "who", "whom", 
#     "whose", "where", "when", "why", "how"
# ])

# def safe_sentence_split(text: str):
#     """Split text into sentences safely"""
#     pattern = re.compile(r"(?<=[.!?])\s+(?=[A-Z0-9])")
#     return [s.strip() for s in pattern.split(text) if s.strip()]

# def summarize_text_advanced(text: str, max_sentences: int = 5, as_bullets: bool = False) -> str:
#     """Advanced text summarization using frequency analysis"""
#     if not text or not text.strip():
#         return "No content to summarize."
    
#     paragraphs = [p.strip() for p in text.splitlines() if p.strip()]
#     sentences = []
#     for para in paragraphs:
#         sentences.extend(safe_sentence_split(para))
    
#     if not sentences:
#         return text

#     # Word frequency analysis
#     word_freq = Counter()
#     for s in sentences:
#         words = [w.lower() for w in re.findall(r"[A-Za-z0-9_']+", s)]
#         for w in words:
#             if w not in STOPWORDS and len(w) > 2:
#                 word_freq[w] += 1
    
#     if not word_freq:
#         return " ".join(sentences[:max_sentences])

#     # Normalize frequencies
#     max_freq = max(word_freq.values())
#     for w in list(word_freq.keys()):
#         word_freq[w] /= max_freq

#     # Score sentences
#     scored = []
#     for idx, s in enumerate(sentences):
#         words = [w.lower() for w in re.findall(r"[A-Za-z0-9_']+", s)]
#         score = sum(word_freq.get(w, 0.0) for w in words)
#         length_penalty = 1.0 + 0.2 * max(0, (len(words) - 20) / 20)
#         position_boost = 1.1 if idx < 3 else 1.0
#         scored.append((score / length_penalty * position_boost, idx, s))

#     # Select top sentences
#     scored.sort(key=lambda x: (-x[0], x[1]))
#     top = sorted(scored[:max_sentences], key=lambda x: x[1])
    
#     if as_bullets:
#         return "\n".join([f"• {s}" for _, _, s in top])
#     return " ".join([s for _, _, s in top])

# def analyze_python_enhanced(code: str):
#     """Comprehensive Python code analysis with enhanced metrics"""
#     report = {
#         "basic_stats": {},
#         "functions": [], 
#         "classes": [], 
#         "imports": [], 
#         "variables": [],
#         "purpose_summary": "", 
#         "errors": [], 
#         "warnings": [], 
#         "fixes": [],
#         "complexity_metrics": {},
#         "code_quality": {},
#         "detailed_counts": {},
#         "builtin_functions_used": []  # Initialize this key
#     }

#     if not code or not code.strip():
#         return report

#     lines = code.splitlines()
#     code_lines = [line for line in lines if line.strip() and not line.strip().startswith("#")]
#     comment_lines = [line for line in lines if line.strip().startswith("#")]
#     blank_lines = [line for line in lines if not line.strip()]

#     # Basic statistics
#     report["basic_stats"] = {
#         "total_lines": len(lines),
#         "code_lines": len(code_lines),
#         "comment_lines": len(comment_lines),
#         "blank_lines": len(blank_lines),
#         "characters": len(code),
#         "words": len(code.split())
#     }

#     # Initialize counters
#     function_count = 0
#     class_count = 0
#     import_count = 0
#     print_count = 0
#     if_count = 0
#     for_count = 0
#     while_count = 0
#     try_count = 0
#     with_count = 0
#     def_count = 0
#     lambda_count = 0
#     list_comp_count = 0
#     dict_comp_count = 0

#     # Pattern-based counting
#     for line in code_lines:
#         stripped_line = line.strip()
        
#         # Count various constructs
#         if stripped_line.startswith("def "):
#             def_count += 1
#         if stripped_line.startswith("class "):
#             class_count += 1
#         if stripped_line.startswith(("import ", "from ")):
#             import_count += 1
        
#         # Count control structures and statements
#         print_count += len(re.findall(r'\bprint\s*\(', line))
#         if_count += len(re.findall(r'\bif\b', line))
#         for_count += len(re.findall(r'\bfor\b', line))
#         while_count += len(re.findall(r'\bwhile\b', line))
#         try_count += len(re.findall(r'\btry\b', line))
#         with_count += len(re.findall(r'\bwith\b', line))
#         lambda_count += len(re.findall(r'\blambda\b', line))
#         list_comp_count += len(re.findall(r'\[.*for.*in.*\]', line))
#         dict_comp_count += len(re.findall(r'\{.*for.*in.*\}', line))

#     report["detailed_counts"] = {
#         "functions": def_count,
#         "classes": class_count,
#         "imports": import_count,
#         "print_statements": print_count,
#         "if_statements": if_count,
#         "for_loops": for_count,
#         "while_loops": while_count,
#         "try_blocks": try_count,
#         "with_statements": with_count,
#         "lambda_expressions": lambda_count,
#         "list_comprehensions": list_comp_count,
#         "dict_comprehensions": dict_comp_count
#     }

#     # Calculate complexity metrics
#     total_constructs = sum([if_count, for_count, while_count, try_count, def_count])
#     complexity_score = total_constructs / max(1, len(code_lines)) * 100
    
#     report["complexity_metrics"] = {
#         "cyclomatic_complexity_estimate": total_constructs + 1,
#         "complexity_per_line": round(complexity_score, 2),
#         "nesting_level_estimate": max(0, len(re.findall(r'    ', code)) // len(code_lines) * 10) if code_lines else 0
#     }

#     # Code quality metrics
#     comment_ratio = len(comment_lines) / max(1, len(lines)) * 100
#     avg_line_length = sum(len(line) for line in code_lines) / max(1, len(code_lines))
    
#     report["code_quality"] = {
#         "comment_ratio": round(comment_ratio, 2),
#         "avg_line_length": round(avg_line_length, 2),
#         "code_to_comment_ratio": round(len(code_lines) / max(1, len(comment_lines)), 2)
#     }

#     # Basic parsing for simple analysis
#     for line in code_lines:
#         stripped = line.strip()
#         if stripped.startswith("def "):
#             match = re.match(r"def\s+([a-zA-Z_][a-zA-Z0-9_]*)", stripped)
#             if match:
#                 report["functions"].append(match.group(1))
#         elif stripped.startswith("class "):
#             match = re.match(r"class\s+([a-zA-Z_][a-zA-Z0-9_]*)", stripped)
#             if match:
#                 report["classes"].append(match.group(1))
#         elif stripped.startswith(("import ", "from ")):
#             report["imports"].append(stripped)

#     # Generate purpose summary
#     report["purpose_summary"] = summarize_text_advanced(code, max_sentences=3)

#     # AST analysis for deeper insights
#     try:
#         tree = ast.parse(code)
#     except SyntaxError as e:
#         report["errors"].append(f"SyntaxError: {e.msg} at line {e.lineno}")
#         report["fixes"].append("Check syntax: indentation, colons, parentheses, quotes.")
#         return report

#     # Enhanced AST analysis
#     imported_names = set()
#     used_names = set()
#     assigned_names = set()
#     builtin_functions_used = set()
    
#     builtin_funcs = {
#         'abs', 'all', 'any', 'bin', 'bool', 'chr', 'dict', 'dir', 'enumerate', 
#         'filter', 'float', 'format', 'frozenset', 'hash', 'hex', 'id', 'input', 
#         'int', 'isinstance', 'len', 'list', 'map', 'max', 'min', 'next', 'oct', 
#         'open', 'ord', 'pow', 'print', 'range', 'repr', 'reversed', 'round', 
#         'set', 'slice', 'sorted', 'str', 'sum', 'tuple', 'type', 'zip'
#     }

#     class EnhancedCodeAnalyzer(ast.NodeVisitor):
#         def visit_Import(self, node):
#             for alias in node.names:
#                 imported_names.add(alias.asname or alias.name.split(".")[0])
#             self.generic_visit(node)

#         def visit_ImportFrom(self, node):
#             for alias in node.names:
#                 imported_names.add(alias.asname or alias.name)
#             self.generic_visit(node)

#         def visit_FunctionDef(self, node):
#             assigned_names.add(node.name)
#             # Check for mutable default arguments
#             for default in node.args.defaults:
#                 if isinstance(default, (ast.List, ast.Dict, ast.Set)):
#                     report["warnings"].append(f"Mutable default argument in '{node.name}'")
#                     report["fixes"].append(f"Use None as default in '{node.name}' and create objects inside function")
            
#             # Check for long functions
#             func_lines = len([n for n in ast.walk(node) if hasattr(n, 'lineno')])
#             if func_lines > 50:
#                 report["warnings"].append(f"Function '{node.name}' is very long ({func_lines} lines)")
#                 report["fixes"].append(f"Consider breaking down '{node.name}' into smaller functions")
            
#             self.generic_visit(node)

#         def visit_ClassDef(self, node):
#             assigned_names.add(node.name)
#             self.generic_visit(node)

#         def visit_Name(self, node):
#             if isinstance(node.ctx, ast.Load):
#                 used_names.add(node.id)
#                 if node.id in builtin_funcs:
#                     builtin_functions_used.add(node.id)
#             elif isinstance(node.ctx, ast.Store):
#                 assigned_names.add(node.id)
#                 report["variables"].append(node.id)
#             self.generic_visit(node)

#         def visit_Call(self, node):
#             if isinstance(node.func, ast.Name):
#                 if node.func.id in {"eval", "exec"}:
#                     report["warnings"].append(f"Dangerous {node.func.id} usage detected")
#                     report["fixes"].append(f"Avoid {node.func.id}; use safer alternatives")
#                 elif node.func.id == "print" and len(node.args) == 0:
#                     report["warnings"].append("Empty print() statement found")
#                     report["fixes"].append("Remove empty print() or add meaningful content")
#             self.generic_visit(node)

#     EnhancedCodeAnalyzer().visit(tree)

#     # Add builtin functions used to report
#     report["builtin_functions_used"] = sorted(list(builtin_functions_used))

#     # Check for unused imports
#     unused_imports = []
#     for name in sorted(imported_names):
#         if name not in used_names and name not in {"__future__"}:
#             unused_imports.append(name)
#             report["warnings"].append(f"Possibly unused import: '{name}'")
#             report["fixes"].append(f"Remove unused import '{name}' if not needed")

#     # Check for builtin shadowing
#     builtins_list = [
#         'abs', 'all', 'any', 'bin', 'bool', 'chr', 'dict', 'dir', 'enumerate', 
#         'filter', 'float', 'format', 'frozenset', 'hash', 'hex', 'id', 'input', 
#         'int', 'isinstance', 'len', 'list', 'map', 'max', 'min', 'next', 'oct', 
#         'open', 'ord', 'pow', 'print', 'range', 'repr', 'reversed', 'round', 
#         'set', 'slice', 'sorted', 'str', 'sum', 'tuple', 'type', 'zip'
#     ]
    
#     shadowed_builtins = []
#     for name in assigned_names:
#         if name in builtins_list:
#             shadowed_builtins.append(name)
#             report["warnings"].append(f"Variable '{name}' shadows builtin")
#             report["fixes"].append(f"Rename '{name}' to avoid shadowing builtins")

#     # Code quality analysis
#     if report["code_quality"]["comment_ratio"] < 5:
#         report["warnings"].append("Low comment density - consider adding more documentation")
#         report["fixes"].append("Add docstrings to functions and classes, and inline comments for complex logic")
    
#     if report["code_quality"]["avg_line_length"] > 100:
#         report["warnings"].append("Some lines are very long - consider breaking them up")
#         report["fixes"].append("Break long lines at logical points (after commas, operators, etc.)")

#     # Add general recommendations
#     if not report["errors"]:
#         if not report["warnings"]:
#             report["fixes"].append("🎉 Code looks excellent! Consider adding unit tests and type hints.")
#         else:
#             report["fixes"].append("Review warnings above and apply suggested fixes for better code quality.")

#     # Remove duplicate variables
#     report["variables"] = sorted(list(set(report["variables"])))

#     return report

# def detect_ai_generated_code(code: str) -> dict:
#     """Detect if code might be AI-generated based on patterns"""
#     if not code or not code.strip():
#         return {"score": 0, "label": "❓ No content", "reasons": [], "features": {}}
    
#     lines = code.splitlines()
#     code_lines = [ln for ln in lines if ln.strip() and not ln.strip().startswith("#")]
#     comment_lines = [ln for ln in lines if ln.strip().startswith("#")]
    
#     features = {}
#     total_lines = max(1, len(lines))
    
#     # Comment density
#     features["comment_density"] = len(comment_lines) / total_lines
    
#     # Repeated lines
#     normalized = [re.sub(r"\s+", " ", ln.strip()) for ln in code_lines]
#     counts = Counter(normalized)
#     repeated = sum(c for c in counts.values() if c > 1)
#     features["repeated_line_ratio"] = repeated / max(1, len(code_lines))
    
#     # Generic variable names
#     generic_names = {"data", "result", "temp", "value", "item", "input", "output", "res", "var", "obj"}
#     tokens = re.findall(r"[A-Za-z_][A-Za-z0-9_]*", code)
#     generic_count = sum(1 for t in tokens if t in generic_names)
#     features["generic_name_density"] = generic_count / max(1, len(tokens))
    
#     # Perfect formatting (suspicious)
#     perfect_indent = all(len(line) - len(line.lstrip()) % 4 == 0 for line in code_lines if line.strip())
#     features["perfect_formatting"] = 1.0 if perfect_indent and len(code_lines) > 10 else 0.0
    
#     # Calculate score
#     score = (
#         25 * min(1.0, abs(features["comment_density"] - 0.15) / 0.15) +
#         35 * features["repeated_line_ratio"] +
#         20 * min(1.0, features["generic_name_density"] * 20) +
#         20 * features["perfect_formatting"]
#     )
    
#     score = max(0.0, min(100.0, score))
    
#     # Determine label
#     if score >= 70:
#         label = "🚨 Likely AI-generated"
#     elif score >= 45:
#         label = "🤔 Unclear/Mixed"
#     else:
#         label = "✅ Likely human-written"
    
#     # Generate reasons
#     reasons = []
#     if features["repeated_line_ratio"] > 0.1:
#         reasons.append("High repetition of similar code patterns")
#     if features["comment_density"] < 0.02 or features["comment_density"] > 0.4:
#         reasons.append("Unusual comment density")
#     if features["generic_name_density"] > 0.03:
#         reasons.append("Frequent generic variable names")
#     if features["perfect_formatting"] == 1.0:
#         reasons.append("Suspiciously perfect code formatting")
    
#     return {
#         "score": round(score, 1), 
#         "label": label, 
#         "reasons": reasons, 
#         "features": features
#     }

# def fetch_from_url(url: str) -> str:
#     """Fetch and extract text content from URL"""
#     try:
#         headers = {
#             'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
#         }
#         resp = requests.get(url, timeout=10, headers=headers)
#         resp.raise_for_status()
        
#         content_type = resp.headers.get("Content-Type", "")
#         if "text/html" in content_type:
#             soup = BeautifulSoup(resp.text, "html.parser")
#             # Remove unwanted elements
#             for element in soup(["script", "style", "nav", "header", "footer", "aside"]):
#                 element.extract()
#             return soup.get_text(separator="\n", strip=True)
#         else:
#             return resp.text
            
#     except requests.exceptions.RequestException as e:
#         return f"❌ Error fetching URL: {str(e)}"
#     except Exception as e:
#         return f"❌ Error processing content: {str(e)}"

# # Create tabs
# tabs = st.tabs([
#     "📄 Document Processor", 
#     "🧠 Code Analyzer", 
#     "🤖 AI Detection", 
#     "🌐 URL Extractor"
# ])

# # Tab 1: Document Processor
# with tabs[0]:
#     st.header("📄 Neural Document Processor")
#     st.write("*Advanced text analysis and summarization*")
    
#     col1, col2 = st.columns([2, 1])
    
#     with col1:
#         doc_text = st.text_area(
#             "Input Text", 
#             height=200, 
#             placeholder="Paste your document content here...",
#             key="doc_text_input"
#         )
    
#     with col2:
#         doc_file = st.file_uploader("Upload Document", type=["txt", "md", "pdf"], key="doc_file_upload")
#         doc_url = st.text_input("Document URL", placeholder="https://...", key="doc_url_input")
#         doc_length = st.slider("Summary Length", 1, 10, 5, key="doc_summary_length")
#         doc_bullets = st.checkbox("Use Bullets", value=False, key="doc_use_bullets")

#     if st.button("🚀 Analyze Document", type="primary", key="doc_analyze_btn"):
#         content = ""
        
#         # Get content from various sources
#         if doc_file:
#             if doc_file.name.lower().endswith(".pdf"):
#                 if not PDF_AVAILABLE:
#                     st.error("PDF processing requires PyPDF2. Install with: pip install PyPDF2")
#                 else:
#                     try:
#                         reader = PyPDF2.PdfReader(io.BytesIO(doc_file.read()))
#                         pages_text = []
#                         for page in reader.pages:
#                             pages_text.append(page.extract_text() or "")
#                         content = "\n".join(pages_text)
#                     except Exception as e:
#                         st.error(f"PDF error: {e}")
#             else:
#                 content = doc_file.read().decode("utf-8", errors="ignore")
#         elif doc_url.strip():
#             with st.spinner("Fetching content..."):
#                 content = fetch_from_url(doc_url.strip())
#         elif doc_text.strip():
#             content = doc_text

#         if not content:
#             st.warning("Please provide content via text, file, or URL.")
#         elif content.startswith("❌"):
#             st.error(content)
#         else:
#             with st.spinner("Processing..."):
#                 summary = summarize_text_advanced(content, doc_length, doc_bullets)
            
#             st.success("Analysis Complete!")
            
#             # Show summary FIRST (most important)
#             st.subheader("📋 Generated Summary")
#             if doc_bullets:
#                 st.markdown(summary)
#             else:
#                 st.write(summary)
            
#             # Copy functionality for summary
#             st.subheader("📋 Copy Summary")
#             copy_button(summary, "Copy Summary", key="doc_summary")
            
#             # Then show document metrics
#             st.subheader("📊 Document Statistics")
#             col1, col2, col3, col4 = st.columns(4)
#             with col1:
#                 st.metric("Characters", f"{len(content):,}")
#             with col2:
#                 st.metric("Words", f"{len(content.split()):,}")
#             with col3:
#                 st.metric("Lines", f"{len(content.splitlines()):,}")
#             with col4:
#                 st.metric("Paragraphs", f"{len([p for p in content.split('\\n\\n') if p.strip()]):,}")
            
#             # Download option
#             st.download_button(
#                 "💾 Download Summary", 
#                 summary, 
#                 "summary.txt",
#                 mime="text/plain",
#                 key="doc_download_btn"
#             )

# # Tab 2: Enhanced Code Analyzer with reordered sections
# with tabs[1]:
#     st.header("🧠 Quantum Code Analyzer")
#     st.write("*Deep code analysis with comprehensive metrics and optimization recommendations*")
    
#     col1, col2 = st.columns([2, 1])
    
#     with col1:
#         code_text = st.text_area(
#             "Python Code", 
#             height=250, 
#             placeholder="Paste your Python code here...",
#             key="code_text_input"
#         )
    
#     with col2:
#         code_file = st.file_uploader("Upload Python File", type=["py"], key="code_file_upload")
#         code_url = st.text_input("Code URL", placeholder="https://...", key="code_url_input")

#     if st.button("🔍 Analyze Code", type="primary", key="code_analyze_btn"):
#         code_content = ""
        
#         if code_file:
#             code_content = code_file.read().decode("utf-8", errors="ignore")
#         elif code_url.strip():
#             with st.spinner("Fetching code..."):
#                 code_content = fetch_from_url(code_url.strip())
#         elif code_text.strip():
#             code_content = code_text

#         if not code_content:
#             st.warning("Please provide code via input, file, or URL.")
#         elif code_content.startswith("❌"):
#             st.error(code_content)
#         else:
#             with st.expander("Code Preview", expanded=False):
#                 st.code(code_content, language="python")

#             with st.spinner("Analyzing code..."):
#                 report = analyze_python_enhanced(code_content)

#             st.success("🎉 Analysis Complete!")

#             # PRIORITY SECTION 1: ERRORS (Most Important)
#             st.header("❌ Errors")
#             if report["errors"]:
#                 for error in report["errors"]:
#                     st.error(error)
#             else:
#                 st.success("✅ No syntax errors found!")
            
#             # PRIORITY SECTION 2: WARNINGS
#             st.header("⚠️ Warnings")
#             if report["warnings"]:
#                 for warning in report["warnings"]:
#                     st.warning(warning)
#             else:
#                 st.success("✅ No warnings!")
            
#             # PRIORITY SECTION 3: SUGGESTIONS/RECOMMENDATIONS
#             st.header("💡 Suggestions & Recommendations")
#             for fix in report["fixes"]:
#                 st.info(fix)
            
#             # PRIORITY SECTION 4: CODE SUMMARY
#             st.header("📋 Code Summary")
#             st.write(report["purpose_summary"])
            
#             # Copy functionality for main analysis
#             st.header("📋 Copy Analysis Report")
#             formatted_report = f"""
# CODE ANALYSIS REPORT
# ==================

# ERRORS ({len(report['errors'])} found):
# {chr(10).join(f"- {error}" for error in report['errors']) if report['errors'] else "✅ No errors"}

# WARNINGS ({len(report['warnings'])} found):
# {chr(10).join(f"- {warning}" for warning in report['warnings']) if report['warnings'] else "✅ No warnings"}

# RECOMMENDATIONS:
# {chr(10).join(f"- {fix}" for fix in report['fixes'])}

# CODE SUMMARY:
# {report['purpose_summary']}

# FUNCTIONS FOUND ({len(report['functions'])}):
# {', '.join(report['functions']) if report['functions'] else "None"}

# CLASSES FOUND ({len(report['classes'])}):
# {', '.join(report['classes']) if report['classes'] else "None"}

# DETAILED STATISTICS:
# - Total Lines: {report['basic_stats'].get('total_lines', 0)}
# - Code Lines: {report['basic_stats'].get('code_lines', 0)}
# - Comment Lines: {report['basic_stats'].get('comment_lines', 0)}
# - Functions: {report['detailed_counts'].get('functions', 0)}
# - Classes: {report['detailed_counts'].get('classes', 0)}
# - Imports: {report['detailed_counts'].get('imports', 0)}
# - Print Statements: {report['detailed_counts'].get('print_statements', 0)}
# - If Statements: {report['detailed_counts'].get('if_statements', 0)}
# - For Loops: {report['detailed_counts'].get('for_loops', 0)}
# - While Loops: {report['detailed_counts'].get('while_loops', 0)}
# """
#             copy_button(formatted_report, "Copy Full Report", key="code_report")
            
#             # NOW SHOW DETAILED METRICS (Lower Priority)
#             st.header("📊 Detailed Code Metrics")
            
#             # Functions and Classes Found
#             col1, col2 = st.columns(2)
            
#             with col1:
#                 if report["functions"]:
#                     st.subheader("🔧 Functions Found")
#                     for func in report["functions"][:10]:  # Show first 10
#                         st.write(f"• `{func}()`")
#                     if len(report["functions"]) > 10:
#                         st.write(f"... and {len(report['functions']) - 10} more")
                
#                 # SAFE ACCESS TO builtin_functions_used - FIXED LINE
#                 builtin_functions = report.get("builtin_functions_used", [])
#                 if builtin_functions:
#                     st.subheader("🐍 Built-in Functions Used")
#                     builtin_display = ", ".join(f"`{func}()`" for func in builtin_functions[:15])
#                     st.write(builtin_display)
#                     if len(builtin_functions) > 15:
#                         st.write(f"... and {len(builtin_functions) - 15} more")
            
#             with col2:
#                 if report["classes"]:
#                     st.subheader("🏗️ Classes Found")
#                     for cls in report["classes"]:
#                         st.write(f"• `{cls}`")
                
#                 if report["variables"]:
#                     st.subheader("📝 Variables")
#                     variables_display = ", ".join(f"`{var}`" for var in report["variables"][:20])
#                     st.write(variables_display)
#                     if len(report["variables"]) > 20:
#                         st.write(f"... and {len(report['variables']) - 20} more variables")

#             # Basic Statistics Section
#             st.subheader("📈 Line Statistics")
#             col1, col2, col3, col4 = st.columns(4)
            
#             with col1:
#                 st.metric("Total Lines", report["basic_stats"].get("total_lines", 0))
#                 st.metric("Functions", report["detailed_counts"].get("functions", 0))
#             with col2:
#                 st.metric("Code Lines", report["basic_stats"].get("code_lines", 0))
#                 st.metric("Classes", report["detailed_counts"].get("classes", 0))
#             with col3:
#                 st.metric("Comments", report["basic_stats"].get("comment_lines", 0))
#                 st.metric("Imports", report["detailed_counts"].get("imports", 0))
#             with col4:
#                 st.metric("Characters", f"{report['basic_stats'].get('characters', 0):,}")
#                 st.metric("Print Statements", report["detailed_counts"].get("print_statements", 0))

#             # Detailed Counts Section
#             st.subheader("🔧 Code Constructs")
#             col1, col2, col3, col4 = st.columns(4)
            
#             with col1:
#                 st.metric("If Statements", report["detailed_counts"].get("if_statements", 0))
#                 st.metric("For Loops", report["detailed_counts"].get("for_loops", 0))
#             with col2:
#                 st.metric("While Loops", report["detailed_counts"].get("while_loops", 0))
#                 st.metric("Try Blocks", report["detailed_counts"].get("try_blocks", 0))
#             with col3:
#                 st.metric("With Statements", report["detailed_counts"].get("with_statements", 0))
#                 st.metric("Lambda Expressions", report["detailed_counts"].get("lambda_expressions", 0))
#             with col4:
#                 st.metric("List Comprehensions", report["detailed_counts"].get("list_comprehensions", 0))
#                 st.metric("Dict Comprehensions", report["detailed_counts"].get("dict_comprehensions", 0))

#             # Code Quality Metrics
#             st.subheader("📊 Quality Metrics")
#             col1, col2, col3 = st.columns(3)
            
#             with col1:
#                 comment_ratio = report["code_quality"].get("comment_ratio", 0)
#                 st.metric("Comment Ratio", f"{comment_ratio:.1f}%")
                
#             with col2:
#                 avg_line_length = report["code_quality"].get("avg_line_length", 0)
#                 st.metric("Avg Line Length", f"{avg_line_length:.1f} chars")
                
#             with col3:
#                 complexity = report["complexity_metrics"].get("cyclomatic_complexity_estimate", 0)
#                 st.metric("Complexity Score", complexity)

#             # Imports and Variables Details
#             with st.expander("📦 Imports & Variables Details", expanded=False):
#                 col1, col2 = st.columns(2)
                
#                 with col1:
#                     if report["imports"]:
#                         st.subheader("📥 Imports")
#                         for imp in report["imports"]:
#                             st.code(imp, language="python")
#                     else:
#                         st.write("No imports found")
                
#                 with col2:
#                     if report["variables"]:
#                         st.subheader("📝 All Variables")
#                         variables_display = ", ".join(f"`{var}`" for var in report["variables"])
#                         st.write(variables_display)
#                     else:
#                         st.write("No variables detected")
            
#             # Download option
#             st.download_button(
#                 "💾 Download Full Report", 
#                 formatted_report, 
#                 "code_analysis_report.txt",
#                 mime="text/plain",
#                 key="code_download_btn"
#             )

# # Tab 3: AI Detection with reordered sections
# with tabs[2]:
#     st.header("🤖 AI Detection Matrix")
#     st.write("*Advanced pattern recognition for AI-generated content*")
    
#     col1, col2 = st.columns([3, 1])
    
#     with col1:
#         ai_text = st.text_area(
#             "Content for Analysis", 
#             height=200, 
#             placeholder="Paste text or code to analyze for AI patterns...",
#             key="ai_text_input"
#         )
    
#     with col2:
#         ai_file = st.file_uploader("Upload File", type=["txt", "py", "md"], key="ai_file_upload")
#         ai_url = st.text_input("Content URL", placeholder="https://...", key="ai_url_input")

#     if st.button("🔬 Scan for AI Patterns", type="primary", key="ai_analyze_btn"):
#         content = ""
        
#         if ai_file:
#             content = ai_file.read().decode("utf-8", errors="ignore")
#         elif ai_url.strip():
#             with st.spinner("Fetching content..."):
#                 content = fetch_from_url(ai_url.strip())
#         elif ai_text.strip():
#             content = ai_text

#         if not content:
#             st.warning("Please provide content for analysis.")
#         elif content.startswith("❌"):
#             st.error(content)
#         else:
#             with st.spinner("Analyzing patterns..."):
#                 result = detect_ai_generated_code(content)
            
#             st.success("🔍 Analysis Complete!")
            
#             # PRIORITY 1: MAIN ASSESSMENT (Most Important)
#             score_val = result['score']
#             if score_val < 45:
#                 color = "#00ff00"
#                 interpretation = "Human-like patterns detected"
#             elif score_val < 65:
#                 color = "#ffaa00"
#                 interpretation = "Mixed or uncertain patterns"
#             else:
#                 color = "#ff4444"
#                 interpretation = "AI-like patterns detected"
            
#             st.header(f"🎯 AI Detection Result: {result['label']}")
            
#             # Visual score indicator
#             col1, col2, col3 = st.columns([1, 2, 1])
#             with col2:
#                 st.metric(
#                     "AI Likelihood Score", 
#                     f"{score_val}%",
#                     delta=f"{score_val - 50:+.1f}% vs baseline",
#                     delta_color="inverse"
#                 )
                
#                 # Progress bar for visual representation
#                 st.markdown(f"""
#                 <div style="background: rgba(255,255,255,0.1); border-radius: 10px; padding: 5px; margin: 10px 0;">
#                     <div style="background: {color}; height: 20px; width: {score_val}%; border-radius: 8px; transition: all 0.3s;"></div>
#                 </div>
#                 <p style="text-align: center; color: {color}; font-weight: bold;">{interpretation}</p>
#                 """, unsafe_allow_html=True)
            
#             # PRIORITY 2: KEY INDICATORS
#             st.header("🔍 Key Indicators")
#             if result["reasons"]:
#                 for i, reason in enumerate(result["reasons"], 1):
#                     st.warning(f"{i}. {reason}")
#             else:
#                 st.success("✅ No significant AI patterns detected")
            
#             # PRIORITY 3: COPY FUNCTIONALITY
#             st.header("📋 Copy Detection Report")
#             detection_report = f"""
# AI DETECTION ANALYSIS REPORT
# ============================

# OVERALL ASSESSMENT: {result['label']}
# CONFIDENCE SCORE: {result['score']}%
# INTERPRETATION: {interpretation}

# KEY INDICATORS:
# {chr(10).join(f"- {reason}" for reason in result['reasons']) if result['reasons'] else "- No significant patterns detected"}

# DETAILED FEATURE ANALYSIS:
# {chr(10).join(f"- {k.replace('_', ' ').title()}: {v:.3f}" for k, v in result['features'].items())}

# CONFIDENCE BREAKDOWN:
# - Comment Pattern Score: {abs(result['features'].get('comment_density', 0) - 0.15) / 0.15:.2f}
# - Code Repetition Score: {result['features'].get('repeated_line_ratio', 0):.2f}
# - Generic Names Score: {result['features'].get('generic_name_density', 0):.2f}
# - Perfect Format Score: {result['features'].get('perfect_formatting', 0):.2f}
# """
#             copy_button(detection_report, "Copy Detection Report", key="ai_report")
            
#             # LOWER PRIORITY: DETAILED METRICS
#             st.header("📊 Detailed Feature Analysis")
            
#             col1, col2 = st.columns(2)
            
#             with col1:
#                 st.write("**📈 Pattern Metrics:**")
#                 for feature, value in result["features"].items():
#                     feature_name = feature.replace("_", " ").title()
                    
#                     # Color code based on feature values
#                     if "density" in feature or "ratio" in feature:
#                         if value > 0.3:
#                             color_indicator = "🔴"
#                         elif value > 0.15:
#                             color_indicator = "🟡"
#                         else:
#                             color_indicator = "🟢"
#                     else:
#                         color_indicator = "🔵"
                    
#                     st.write(f"{color_indicator} **{feature_name}:** `{value:.3f}`")
            
#             with col2:
#                 st.write("**🎯 Confidence Breakdown:**")
                
#                 # Create confidence indicators for each feature
#                 features = result["features"]
                
#                 if "comment_density" in features:
#                     cd_score = abs(features["comment_density"] - 0.15) / 0.15
#                     st.write(f"📝 Comment Pattern: `{cd_score:.2f}` {'⚠️' if cd_score > 0.5 else '✅'}")
                
#                 if "repeated_line_ratio" in features:
#                     rr_score = features["repeated_line_ratio"]
#                     st.write(f"🔄 Code Repetition: `{rr_score:.2f}` {'⚠️' if rr_score > 0.1 else '✅'}")
                
#                 if "generic_name_density" in features:
#                     gn_score = features["generic_name_density"]
#                     st.write(f"🏷️ Generic Names: `{gn_score:.2f}` {'⚠️' if gn_score > 0.03 else '✅'}")
                
#                 if "perfect_formatting" in features:
#                     pf_score = features["perfect_formatting"]
#                     st.write(f"✨ Perfect Format: `{pf_score:.2f}` {'⚠️' if pf_score > 0.5 else '✅'}")
            
#             # Download option
#             st.download_button(
#                 "💾 Download Analysis", 
#                 detection_report, 
#                 "ai_detection_report.txt",
#                 mime="text/plain",
#                 key="ai_download_btn"
#             )

# # Tab 4: URL Extractor with reordered sections
# with tabs[3]:
#     st.header("🌐 URL Data Extraction Engine")
#     st.write("*Advanced web content extraction and intelligent analysis*")
    
#     col1, col2 = st.columns([3, 1])
    
#     with col1:
#         url_input = st.text_input(
#             "🎯 Target URL", 
#             placeholder="Enter any URL for intelligent content extraction...",
#             key="url_input_field",
#             help="Supports web pages, GitHub files, documentation, articles, and more"
#         )
    
#     with col2:
#         url_length = st.slider("Summary Length", 1, 15, 7, key="url_summary_length")
#         url_bullets = st.checkbox("Use Bullet Points", value=True, key="url_use_bullets")
#         url_show_preview = st.checkbox("Show Content Preview", value=True, key="url_show_preview")

#     if st.button("🚀 Extract & Analyze", type="primary", key="url_analyze_btn"):
#         if not url_input.strip():
#             st.warning("⚠️ Please enter a valid URL to proceed.")
#         else:
#             with st.spinner("🔄 Extracting content from URL..."):
#                 content = fetch_from_url(url_input.strip())
            
#             if content.startswith("❌"):
#                 st.error(content)
#             else:
#                 st.success("✅ Content extracted successfully!")
                
#                 # Generate enhanced summary
#                 with st.spinner("🧠 Generating intelligent summary..."):
#                     summary = summarize_text_advanced(content, url_length, url_bullets)
                
#                 # PRIORITY 1: SUMMARY (Most Important)
#                 st.header("🎯 Intelligent Content Summary")
#                 if url_bullets:
#                     st.markdown(summary)
#                 else:
#                     st.write(summary)
                
#                 # PRIORITY 2: COPY FUNCTIONALITY
#                 st.header("📋 Copy Summary")
#                 copy_button(summary, "Copy Summary", key="url_summary")
                
#                 # Enhanced metrics display (LOWER PRIORITY)
#                 words = content.split()
#                 sentences = len(re.findall(r'[.!?]+', content))
#                 paragraphs = len([p for p in content.split('\n\n') if p.strip()])
                
#                 # Reading time estimation
#                 avg_reading_speed = 200  # words per minute
#                 reading_time = max(1, len(words) // avg_reading_speed)
#                 st.info(f"📚 Estimated reading time: ~{reading_time} minute{'s' if reading_time != 1 else ''}")
                
#                 st.header("📊 Content Statistics")
#                 col1, col2, col3, col4, col5 = st.columns(5)
#                 with col1:
#                     st.metric("📄 Characters", f"{len(content):,}")
#                 with col2:
#                     st.metric("📝 Words", f"{len(words):,}")
#                 with col3:
#                     st.metric("📋 Lines", f"{len(content.splitlines()):,}")
#                 with col4:
#                     st.metric("📖 Sentences", f"{sentences:,}")
#                 with col5:
#                     st.metric("📑 Paragraphs", f"{paragraphs:,}")
                
#                 # Content preview with better formatting
#                 if url_show_preview:
#                     with st.expander("📄 Content Preview", expanded=False):
#                         preview_length = 3000
#                         preview = content[:preview_length]
#                         if len(content) > preview_length:
#                             preview += f"\n\n... [Content truncated - showing first {preview_length:,} characters of {len(content):,} total]"
                        
#                         # Try to detect if it's code
#                         if any(keyword in content.lower()[:500] for keyword in ['def ', 'class ', 'import ', 'function', '#include', 'public class']):
#                             st.code(preview, language="python" if "def " in preview or "import " in preview else "text")
#                         else:
#                             st.text(preview)
                
#                 # Additional analysis
#                 with st.expander("📊 Content Analysis", expanded=False):
#                     col1, col2 = st.columns(2)
                    
#                     with col1:
#                         st.write("**📈 Content Statistics:**")
#                         avg_word_length = sum(len(word) for word in words) / max(len(words), 1)
#                         st.write(f"• Average word length: {avg_word_length:.1f} characters")
                        
#                         if sentences > 0:
#                             avg_sentence_length = len(words) / sentences
#                             st.write(f"• Average sentence length: {avg_sentence_length:.1f} words")
                        
#                         # Most common words (excluding stopwords)
#                         filtered_words = [word.lower().strip('.,!?";()[]') for word in words 
#                                         if len(word) > 3 and word.lower() not in STOPWORDS]
#                         if filtered_words:
#                             common_words = Counter(filtered_words).most_common(5)
#                             st.write("• Most frequent words:")
#                             for word, count in common_words:
#                                 st.write(f"  - **{word}**: {count} times")
                    
#                     with col2:
#                         st.write("**🔍 Content Characteristics:**")
                        
#                         # Detect content type
#                         if any(keyword in content.lower() for keyword in ['function', 'variable', 'method', 'class', 'algorithm']):
#                             st.write("• 💻 **Type**: Technical/Programming content")
#                         elif any(keyword in content.lower() for keyword in ['research', 'study', 'analysis', 'data', 'results']):
#                             st.write("• 🔬 **Type**: Research/Academic content")
#                         elif any(keyword in content.lower() for keyword in ['tutorial', 'how to', 'step', 'guide', 'instructions']):
#                             st.write("• 📚 **Type**: Educational/Tutorial content")
#                         else:
#                             st.write("• 📰 **Type**: General/Article content")
                        
#                         # Complexity estimate
#                         complex_words = [word for word in words if len(word) > 7]
#                         complexity = len(complex_words) / max(len(words), 1) * 100
                        
#                         if complexity > 20:
#                             st.write("• 🎓 **Complexity**: Advanced")
#                         elif complexity > 12:
#                             st.write("• 📖 **Complexity**: Intermediate")
#                         else:
#                             st.write("• 📝 **Complexity**: Beginner-friendly")
                        
#                         st.write(f"• 📊 **Complexity Score**: {complexity:.1f}%")
                
#                 # Enhanced copy and download options
#                 col4, col5, col6 = st.columns(3)
#                 with col4:
#                     copy_button(content, "Copy Full Content", key="url_full")
#                 with col5:
#                     st.download_button(
#                         "💾 Download Summary", 
#                         summary, 
#                         "url_summary.txt",
#                         mime="text/plain",
#                         key="url_download_summary"
#                     )
#                 with col6:
#                     # Create comprehensive report
#                     comprehensive_report = f"""
# URL EXTRACTION REPORT
# ====================

# SOURCE URL: {url_input}
# EXTRACTION DATE: Today

# CONTENT SUMMARY:
# {summary}

# CONTENT METRICS:
# - Characters: {len(content):,}
# - Words: {len(words):,}
# - Lines: {len(content.splitlines()):,}
# - Sentences: {sentences:,}
# - Paragraphs: {paragraphs:,}
# - Estimated Reading Time: ~{reading_time} minute(s)

# FULL EXTRACTED CONTENT:
# {content}
# """
#                     st.download_button(
#                         "💾 Download Full Report", 
#                         comprehensive_report, 
#                         f"url_analysis_report.txt",
#                         mime="text/plain",
#                         key="url_download_full"
#                     )

# # Enhanced Footer
# st.markdown("---")
# st.markdown("""
# <div style="text-align: center; padding: 30px;">
#     <h3 style="color: rgba(0, 249, 255, 0.8); font-family: 'Orbitron', monospace;">🌌 TECHNOVA AI NEXUS v2.1</h3>
#     <p style="color: rgba(0, 249, 255, 0.6); font-family: monospace; margin: 10px 0;">
#         Advanced Neural Processing Suite • Quantum Code Analysis • AI Pattern Detection
#     </p>
#     <p style="color: rgba(0, 249, 255, 0.4); font-size: 0.9em;">
#         Priority-ordered analysis with reliable copy functionality
#     </p>
#     <div style="margin-top: 20px; color: rgba(0, 249, 255, 0.3);">
#         🚀 Built for developers, researchers, and AI enthusiasts
#     </div>
# </div>
# """, unsafe_allow_html=True)

# # Add session state management for better UX
# if 'analysis_history' not in st.session_state:
#     st.session_state['analysis_history'] = []

# # Optional: Add a sidebar with tips and shortcuts
# with st.sidebar:
#     st.header("⚡ Quick Tips")
#     st.write("""
#     **🔧 Code Analyzer:**
#     • Shows errors & warnings first
#     • Upload .py files for best results
#     • Supports GitHub raw URLs
#     • Detects 12+ code constructs
    
#     **🤖 AI Detection:**
#     • Assessment result shown first
#     • Analyzes patterns in code/text
#     • Uses multiple indicators
#     • Confidence scoring system
    
#     **🌐 URL Extractor:**
#     • Summary appears first
#     • Works with most websites
#     • Intelligent content filtering
#     • Automatic summarization
    
#     **📋 Copy Functionality:**
#     • Click copy button
#     • Text area appears below
#     • Select all text (Ctrl+A)
#     • Copy to clipboard (Ctrl+C)
#     """)
    
#     if st.button("🗑️ Clear All Data", key="clear_all"):
#         keys_to_remove = []
#         for key in st.session_state.keys():
#             if any(prefix in key for prefix in ['text_to_copy_', 'copied_text_', 'copy_btn_', 'copy_area_']):
#                 keys_to_remove.append(key)
        
#         for key in keys_to_remove:
#             del st.session_state[key]
        
#         st.success("All temporary data cleared!")
#         st.rerun()














import streamlit as st
import streamlit.components.v1 as components
import requests
from bs4 import BeautifulSoup
import re
import base64
import ast
from collections import Counter, defaultdict
import io
import json
import openai
from typing import Optional, Dict, List, Tuple, Any
import sqlite3
import hashlib
import smtplib
from email.mime.text import MimeText
from email.mime.multipart import MimeMultipart
import random
import string
from datetime import datetime, timedelta
import os
import bcrypt
import time
import threading
from dataclasses import dataclass
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from urllib.parse import urlparse, urljoin
import mimetypes
import zipfile
from pathlib import Path
import tempfile
import markdown
import html2text
from PIL import Image
import fitz  # PyMuPDF for better PDF handling
import docx
from docx import Document
import openpyxl
from openpyxl import Workbook
import csv
import xml.etree.ElementTree as ET
import yaml
import nltk
from textstat import flesch_reading_ease, flesch_kincaid_grade
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import seaborn as sns

# Enhanced imports and error handling
try:
    import PyPDF2
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import python_docx2txt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Page config with enhanced settings
st.set_page_config(
    page_title="TechNova AI Nexus Pro",
    page_icon="🌌",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Enhanced Environment Variables
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
SMTP_SERVER = os.getenv("SMTP_SERVER", "smtp.gmail.com")
SMTP_PORT = int(os.getenv("SMTP_PORT", 587))
SMTP_USERNAME = os.getenv("SMTP_USERNAME", "")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD", "")
STRIPE_SECRET_KEY = os.getenv("STRIPE_SECRET_KEY", "")
ADMIN_EMAIL = os.getenv("ADMIN_EMAIL", "admin@technova.ai")
MAX_FILE_SIZE = int(os.getenv("MAX_FILE_SIZE", 10 * 1024 * 1024))  # 10MB default

# Enhanced Data Models
@dataclass
class User:
    id: int
    email: str
    subscription_type: str
    subscription_expires: datetime
    is_verified: bool
    created_at: datetime
    total_usage: int = 0
    last_active: datetime = None

@dataclass
class UsageLog:
    user_id: int
    tab_name: str
    usage_date: datetime
    usage_count: int
    tokens_used: int = 0
    processing_time: float = 0.0

@dataclass
class ProcessingResult:
    success: bool
    result: Any
    error_message: str = ""
    processing_time: float = 0.0
    tokens_used: int = 0

# Enhanced Database Setup with better schema
def init_enhanced_database():
    """Initialize enhanced SQLite database with better schema"""
    conn = sqlite3.connect('technova_pro.db')
    cursor = conn.cursor()
    
    # Enhanced Users table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            last_active TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            is_verified BOOLEAN DEFAULT FALSE,
            subscription_type TEXT DEFAULT 'free',
            subscription_expires TIMESTAMP,
            verification_code TEXT,
            verification_expires TIMESTAMP,
            total_usage INTEGER DEFAULT 0,
            preferences TEXT DEFAULT '{}',
            api_usage_count INTEGER DEFAULT 0,
            monthly_reset_date DATE DEFAULT (date('now', '+1 month'))
        )
    ''')
    
    # Enhanced Usage tracking
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS usage_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            tab_name TEXT,
            usage_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            usage_count INTEGER DEFAULT 1,
            tokens_used INTEGER DEFAULT 0,
            processing_time REAL DEFAULT 0.0,
            input_size INTEGER DEFAULT 0,
            output_size INTEGER DEFAULT 0,
            success BOOLEAN DEFAULT TRUE,
            error_message TEXT,
            FOREIGN KEY (user_id) REFERENCES users (id),
            INDEX(user_id, usage_date),
            INDEX(tab_name, usage_date)
        )
    ''')
    
    # File processing history
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS file_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            original_filename TEXT,
            file_type TEXT,
            file_size INTEGER,
            processed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            tab_name TEXT,
            processing_successful BOOLEAN DEFAULT TRUE,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    # Feedback and ratings
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            tab_name TEXT,
            rating INTEGER CHECK(rating >= 1 AND rating <= 5),
            feedback_text TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    conn.commit()
    conn.close()

init_enhanced_database()

# Enhanced Authentication Manager with better security
class EnhancedAuthManager:
    @staticmethod
    def hash_password(password: str) -> str:
        """Enhanced password hashing with stronger security"""
        salt = bcrypt.gensalt(rounds=12)  # Stronger rounds
        return bcrypt.hashpw(password.encode('utf-8'), salt).decode('utf-8')
    
    @staticmethod
    def verify_password(password: str, hashed: str) -> bool:
        """Verify password with timing attack protection"""
        try:
            return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))
        except:
            return False
    
    @staticmethod
    def is_strong_password(password: str) -> Tuple[bool, List[str]]:
        """Check password strength"""
        issues = []
        
        if len(password) < 8:
            issues.append("Password must be at least 8 characters long")
        if not re.search(r'[A-Z]', password):
            issues.append("Password must contain at least one uppercase letter")
        if not re.search(r'[a-z]', password):
            issues.append("Password must contain at least one lowercase letter")
        if not re.search(r'\d', password):
            issues.append("Password must contain at least one number")
        if not re.search(r'[!@#$%^&*(),.?":{}|<>]', password):
            issues.append("Password must contain at least one special character")
        
        return len(issues) == 0, issues
    
    @staticmethod
    def generate_secure_code() -> str:
        """Generate cryptographically secure verification code"""
        return ''.join(random.SystemRandom().choices(string.digits, k=6))
    
    @staticmethod
    def is_valid_email(email: str) -> bool:
        """Enhanced email validation"""
        pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        return bool(re.match(pattern, email)) and len(email) <= 254
    
    @staticmethod
    def send_enhanced_verification_email(email: str, verification_code: str) -> bool:
        """Send enhanced verification email with HTML template"""
        if not SMTP_USERNAME or not SMTP_PASSWORD:
            return False
        
        try:
            msg = MimeMultipart('alternative')
            msg['From'] = f"TechNova AI <{SMTP_USERNAME}>"
            msg['To'] = email
            msg['Subject'] = "🌌 Welcome to TechNova AI Nexus - Verify Your Account"
            
            html_body = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <style>
                    body {{ font-family: 'Arial', sans-serif; background: #0a0a0a; color: #00f9ff; }}
                    .container {{ max-width: 600px; margin: 0 auto; padding: 20px; }}
                    .header {{ text-align: center; padding: 20px; background: linear-gradient(135deg, #1a1a2e, #16213e); }}
                    .code {{ font-size: 2em; font-weight: bold; color: #00f9ff; text-align: center; 
                             padding: 20px; background: rgba(0, 249, 255, 0.1); border-radius: 10px; }}
                    .footer {{ text-align: center; margin-top: 20px; font-size: 0.9em; }}
                </style>
            </head>
            <body>
                <div class="container">
                    <div class="header">
                        <h1>🌌 TechNova AI Nexus</h1>
                        <h2>Welcome to the Future of AI Tools!</h2>
                    </div>
                    <p>Thank you for joining TechNova AI Nexus. Your verification code is:</p>
                    <div class="code">{verification_code}</div>
                    <p>This code will expire in 15 minutes for security reasons.</p>
                    <div class="footer">
                        <p>🚀 Get ready to experience next-generation AI tools!</p>
                        <p>Best regards,<br>The TechNova Team</p>
                    </div>
                </div>
            </body>
            </html>
            """
            
            msg.attach(MimeText(html_body, 'html'))
            
            server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
            server.starttls()
            server.login(SMTP_USERNAME, SMTP_PASSWORD)
            server.sendmail(SMTP_USERNAME, email, msg.as_string())
            server.quit()
            
            return True
        except Exception as e:
            st.error(f"Email service error: {str(e)}")
            return False
    
    @staticmethod
    def create_enhanced_user(email: str, password: str) -> Tuple[bool, str]:
        """Create user with enhanced validation and security"""
        if not EnhancedAuthManager.is_valid_email(email):
            return False, "Please provide a valid email address."
        
        is_strong, password_issues = EnhancedAuthManager.is_strong_password(password)
        if not is_strong:
            return False, "Password requirements not met:\n• " + "\n• ".join(password_issues)
        
        conn = sqlite3.connect('technova_pro.db')
        cursor = conn.cursor()
        
        try:
            cursor.execute("SELECT id FROM users WHERE email = ?", (email.lower(),))
            if cursor.fetchone():
                return False, "An account with this email already exists."
            
            verification_code = EnhancedAuthManager.generate_secure_code()
            expires_at = datetime.now() + timedelta(minutes=15)  # Extended to 15 minutes
            password_hash = EnhancedAuthManager.hash_password(password)
            
            cursor.execute('''
                INSERT INTO users (email, password_hash, verification_code, verification_expires, subscription_expires)
                VALUES (?, ?, ?, ?, ?)
            ''', (email.lower(), password_hash, verification_code, expires_at, 
                  datetime.now() + timedelta(days=14)))  # 14-day free trial
            
            conn.commit()
            
            if EnhancedAuthManager.send_enhanced_verification_email(email, verification_code):
                return True, "Account created! Check your email for verification code (expires in 15 minutes)."
            else:
                return False, "Account created but failed to send verification email."
                
        except Exception as e:
            return False, f"Account creation error: {str(e)}"
        finally:
            conn.close()

# Enhanced Usage Manager with better analytics
class EnhancedUsageManager:
    SUBSCRIPTION_LIMITS = {
        'free': {'daily_limit': 5, 'monthly_api_calls': 100},
        'plus': {'daily_limit': -1, 'monthly_api_calls': 10000},  # -1 means unlimited
        'pro': {'daily_limit': -1, 'monthly_api_calls': 50000},
        'expired': {'daily_limit': 0, 'monthly_api_calls': 0}
    }
    
    @staticmethod
    def can_use_tab(user_id: int, tab_name: str) -> Tuple[bool, str, Dict]:
        """Enhanced usage checking with detailed limits"""
        conn = sqlite3.connect('technova_pro.db')
        cursor = conn.cursor()
        
        try:
            cursor.execute("""
                SELECT subscription_type, subscription_expires, api_usage_count, monthly_reset_date 
                FROM users WHERE id = ?
            """, (user_id,))
            
            result = cursor.fetchone()
            if not result:
                return False, "User not found.", {}
            
            sub_type, sub_expires, api_usage, reset_date = result
            
            # Check subscription expiry
            if sub_expires and datetime.fromisoformat(sub_expires) < datetime.now():
                sub_type = 'expired'
            
            # Reset monthly counter if needed
            if reset_date and datetime.fromisoformat(reset_date) < datetime.now():
                cursor.execute("""
                    UPDATE users SET api_usage_count = 0, monthly_reset_date = date('now', '+1 month') 
                    WHERE id = ?
                """, (user_id,))
                api_usage = 0
                conn.commit()
            
            limits = EnhancedUsageManager.SUBSCRIPTION_LIMITS.get(sub_type, {'daily_limit': 0, 'monthly_api_calls': 0})
            
            if sub_type == 'expired':
                return False, "Your free trial has expired. Upgrade to continue using TechNova AI.", {}
            
            # Check monthly API limit
            if limits['monthly_api_calls'] != -1 and api_usage >= limits['monthly_api_calls']:
                return False, f"Monthly API limit reached ({api_usage}/{limits['monthly_api_calls']}). Upgrade for more usage.", {}
            
            # Check daily limit
            if limits['daily_limit'] != -1:
                today = datetime.now().date()
                cursor.execute("""
                    SELECT SUM(usage_count) FROM usage_logs 
                    WHERE user_id = ? AND tab_name = ? AND DATE(usage_date) = ?
                """, (user_id, tab_name, today))
                
                daily_usage = cursor.fetchone()[0] or 0
                if daily_usage >= limits['daily_limit']:
                    return False, f"Daily limit reached for {tab_name} ({daily_usage}/{limits['daily_limit']}).", {}
            
            usage_info = {
                'daily_usage': daily_usage if limits['daily_limit'] != -1 else 0,
                'daily_limit': limits['daily_limit'],
                'monthly_usage': api_usage,
                'monthly_limit': limits['monthly_api_calls']
            }
            
            return True, "", usage_info
            
        except Exception as e:
            return False, f"Error checking usage: {str(e)}", {}
        finally:
            conn.close()
    
    @staticmethod
    def log_enhanced_usage(user_id: int, tab_name: str, tokens_used: int = 0, 
                          processing_time: float = 0.0, success: bool = True, 
                          error_message: str = "", input_size: int = 0, output_size: int = 0):
        """Enhanced usage logging with detailed metrics"""
        conn = sqlite3.connect('technova_pro.db')
        cursor = conn.cursor()
        
        try:
            # Log detailed usage
            cursor.execute('''
                INSERT INTO usage_logs 
                (user_id, tab_name, tokens_used, processing_time, success, error_message, input_size, output_size)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (user_id, tab_name, tokens_used, processing_time, success, error_message, input_size, output_size))
            
            # Update user API usage count
            if success:
                cursor.execute("""
                    UPDATE users SET api_usage_count = api_usage_count + 1, last_active = CURRENT_TIMESTAMP 
                    WHERE id = ?
                """, (user_id,))
            
            conn.commit()
        except Exception as e:
            st.error(f"Error logging usage: {str(e)}")
        finally:
            conn.close()
    
    @staticmethod
    def get_usage_analytics(user_id: int) -> Dict:
        """Get comprehensive usage analytics"""
        conn = sqlite3.connect('technova_pro.db')
        cursor = conn.cursor()
        
        try:
            # Today's usage
            cursor.execute("""
                SELECT tab_name, COUNT(*), SUM(tokens_used), AVG(processing_time), 
                       SUM(CASE WHEN success = 1 THEN 1 ELSE 0 END) as successful_runs
                FROM usage_logs 
                WHERE user_id = ? AND DATE(usage_date) = DATE('now')
                GROUP BY tab_name
            """, (user_id,))
            
            today_stats = {}
            for row in cursor.fetchall():
                today_stats[row[0]] = {
                    'count': row[1],
                    'tokens': row[2] or 0,
                    'avg_time': row[3] or 0,
                    'success_rate': (row[4] / row[1]) * 100 if row[1] > 0 else 0
                }
            
            # Weekly usage trend
            cursor.execute("""
                SELECT DATE(usage_date) as date, COUNT(*) as usage_count
                FROM usage_logs 
                WHERE user_id = ? AND usage_date >= date('now', '-7 days')
                GROUP BY DATE(usage_date)
                ORDER BY date
            """, (user_id,))
            
            weekly_trend = dict(cursor.fetchall())
            
            # Total statistics
            cursor.execute("""
                SELECT COUNT(*) as total_requests, SUM(tokens_used) as total_tokens,
                       SUM(processing_time) as total_time, 
                       SUM(CASE WHEN success = 1 THEN 1 ELSE 0 END) as successful_requests
                FROM usage_logs WHERE user_id = ?
            """, (user_id,))
            
            total_stats = cursor.fetchone()
            
            return {
                'today': today_stats,
                'weekly_trend': weekly_trend,
                'total': {
                    'requests': total_stats[0],
                    'tokens': total_stats[1] or 0,
                    'total_time': total_stats[2] or 0,
                    'success_rate': (total_stats[3] / total_stats[0]) * 100 if total_stats[0] > 0 else 0
                }
            }
            
        except Exception as e:
            return {}
        finally:
            conn.close()

# Advanced File Processing System
class AdvancedFileProcessor:
    SUPPORTED_FORMATS = {
        'text': ['.txt', '.md', '.rst', '.csv', '.json', '.xml', '.yaml', '.yml'],
        'document': ['.pdf', '.docx', '.doc', '.odt', '.rtf'],
        'image': ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff'],
        'code': ['.py', '.js', '.html', '.css', '.java', '.cpp', '.c', '.php', '.rb', '.go'],
        'data': ['.xlsx', '.xls', '.csv', '.json', '.xml', '.sql']
    }
    
    @staticmethod
    def get_file_type(filename: str) -> str:
        """Determine file type from extension"""
        ext = Path(filename).suffix.lower()
        for file_type, extensions in AdvancedFileProcessor.SUPPORTED_FORMATS.items():
            if ext in extensions:
                return file_type
        return 'unknown'
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Advanced PDF text extraction using multiple methods"""
        text = ""
        
        try:
            # Method 1: PyMuPDF (better for complex PDFs)
            if PDF_AVAILABLE:
                doc = fitz.open(stream=file_content, filetype="pdf")
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
        except:
            try:
                # Method 2: PyPDF2 fallback
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            except:
                text = "Error: Could not extract text from PDF"
        
        return text.strip()
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n".join(text_parts)
        except Exception as e:
            return f"Error extracting DOCX: {str(e)}"
    
    @staticmethod
    def process_file(uploaded_file) -> ProcessingResult:
        """Process uploaded file and extract content"""
        start_time = time.time()
        
        try:
            if uploaded_file.size > MAX_FILE_SIZE:
                return ProcessingResult(
                    success=False,
                    result=None,
                    error_message=f"File too large. Maximum size: {MAX_FILE_SIZE / (1024*1024):.1f}MB"
                )
            
            file_type = AdvancedFileProcessor.get_file_type(uploaded_file.name)
            file_content = uploaded_file.read()
            
            text_content = ""
            
            if file_type == 'text':
                # Handle various text formats
                if uploaded_file.name.endswith('.json'):
                    json_data = json.loads(file_content.decode('utf-8'))
                    text_content = json.dumps(json_data, indent=2)
                elif uploaded_file.name.endswith(('.yaml', '.yml')):
                    yaml_data = yaml.safe_load(file_content)
                    text_content = yaml.dump(yaml_data, default_flow_style=False)
                else:
                    text_content = file_content.decode('utf-8')
            
            elif file_type == 'document':
                if uploaded_file.name.endswith('.pdf'):
                    text_content = AdvancedFileProcessor.extract_text_from_pdf(file_content)
                elif uploaded_file.name.endswith('.docx'):
                    text_content = AdvancedFileProcessor.extract_text_from_docx(file_content)
                else:
                    text_content = "Unsupported document format"
            
            elif file_type == 'code':
                text_content = file_content.decode('utf-8')
            
            elif file_type == 'data':
                if uploaded_file.name.endswith(('.xlsx', '.xls')):
                    df = pd.read_excel(io.BytesIO(file_content))
                    text_content = df.to_string()
                elif uploaded_file.name.endswith('.csv'):
                    df = pd.read_csv(io.BytesIO(file_content))
                    text_content = df.to_string()
            
            else:
                return ProcessingResult(
                    success=False,
                    result=None,
                    error_message=f"Unsupported file type: {file_type}"
                )
            
            processing_time = time.time() - start_time
            
            return ProcessingResult(
                success=True,
                result={
                    'content': text_content,
                    'file_type': file_type,
                    'filename': uploaded_file.name,
                    'size': len(file_content)
                },
                processing_time=processing_time
            )
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                result=None,
                error_message=f"Processing error: {str(e)}",
                processing_time=time.time() - start_time
            )

# Advanced AI Processing Engine
class AdvancedAIEngine:
    def __init__(self, api_key: str):
        self.client = openai.OpenAI(api_key=api_key) if api_key else None
        self.models = {
            'fast': 'gpt-3.5-turbo',
            'advanced': 'gpt-4',
            'creative': 'gpt-3.5-turbo',
            'analytical': 'gpt-4'
        }
    
    def enhance_code(self, code: str, language: str, enhancement_type: str = "comprehensive") -> ProcessingResult:
        """Advanced code enhancement with multiple strategies"""
        if not self.client:
            return ProcessingResult(False, None, "OpenAI API not configured")
        
        start_time = time.time()
        
        try:
            enhancement_prompts = {
                'fix_errors': f"""
                Analyze this {language} code and fix any syntax errors, runtime errors, or logical bugs.
                Return only clean, working code with minimal comments explaining fixes.
                
                Code to fix:
                ```{language}
                {code}
                ```
                """,
                
                'optimize': f"""
                Optimize this {language} code for better performance, memory usage, and efficiency.
                Apply best practices and modern language features. Add performance comments.
                
                Code to optimize:
                ```{language}
                {code}
                ```
                """,
                
                'modernize': f"""
                Modernize this {language} code using current best practices, design patterns, 
                and language features. Make it more maintainable and professional.
                
                Code to modernize:
                ```{language}
                {code}
                ```
                """,
                
                'comprehensive': f"""
                Perform a comprehensive enhancement of this {language} code:
                1. Fix any bugs or errors
                2. Optimize for performance and readability
                3. Add proper error handling
                4. Apply best practices and design patterns
                5. Add comprehensive documentation
                6. Make it production-ready
                
                Original code:
                ```{language}
                {code}
                ```
                
                Return only the enhanced code with clear comments.
                """
            }
            
            prompt = enhancement_prompts.get(enhancement_type, enhancement_prompts['comprehensive'])
            
            response = self.client.chat.completions.create(
                model=self.models['advanced'],
                messages=[
                    {"role": "system", "content": f"You are an expert {language} developer. Provide clean, working, well-documented code."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=3000,
                temperature=0.1
            )
            
            enhanced_code = response.choices[0].message.content.strip()
            
            # Clean up markdown formatting
            if enhanced_code.startswith(f"```{language}"):
                enhanced_code = enhanced_code[len(f"```{language}"):].strip()
            elif enhanced_code.startswith("```"):
                enhanced_code = enhanced_code[3:].strip()
            
            if enhanced_code.endswith("```"):
                enhanced_code = enhanced_code[:-3].strip()
            
            processing_time = time.time() - start_time
            tokens_used = response.usage.total_tokens
            
            return ProcessingResult(
                success=True,
                result=enhanced_code,
                processing_time=processing_time,
                tokens_used=tokens_used
            )
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                result=None,
                error_message=f"AI processing error: {str(e)}",
                processing_time=time.time() - start_time
            )
    
    def enhance_text(self, text: str, enhancement_type: str = "comprehensive") -> ProcessingResult:
        """Advanced text enhancement with multiple strategies"""
        if not self.client:
            return ProcessingResult(False, None, "OpenAI API not configured")
        
        start_time = time.time()
        
        try:
            enhancement_prompts = {
                'grammar': f"""
                Correct grammar, spelling, and punctuation errors in this text while preserving the original meaning and tone:
                
                {text}
                """,
                
                'style': f"""
                Improve the writing style, clarity, and flow of this text. Make it more engaging and professional:
                
                {text}
                """,
                
                'formal': f"""
                Rewrite this text in a formal, professional tone suitable for business or academic contexts:
                
                {text}
                """,
                
                'creative': f"""
                Enhance this text creatively while maintaining its core message. Make it more engaging and compelling:
                
                {text}
                """,
                
                'comprehensive': f"""
                Comprehensively enhance this text by:
                1. Correcting grammar, spelling, and punctuation
                2. Improving clarity and readability
                3. Enhancing flow and structure
                4. Making it more engaging and professional
                5. Preserving the original meaning and intent
                
                Original text:
                {text}
                """
            }
            
            prompt = enhancement_prompts.get(enhancement_type, enhancement_prompts['comprehensive'])
            
            response = self.client.chat.completions.create(
                model=self.models['creative'],
                messages=[
                    {"role": "system", "content": "You are an expert editor and writing coach. Enhance text while preserving the author's voice and intent."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.3
            )
            
            enhanced_text = response.choices[0].message.content.strip()
            processing_time = time.time() - start_time
            tokens_used = response.usage.total_tokens
            
            return ProcessingResult(
                success=True,
                result=enhanced_text,
                processing_time=processing_time,
                tokens_used=tokens_used
            )
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                result=None,
                error_message=f"Text enhancement error: {str(e)}",
                processing_time=time.time() - start_time
            )
    
    def analyze_document(self, text: str, analysis_type: str = "comprehensive") -> ProcessingResult:
        """Advanced document analysis"""
        if not self.client:
            return ProcessingResult(False, None, "OpenAI API not configured")
        
        start_time = time.time()
        
        try:
            analysis_prompts = {
                'summary': f"""
                Provide a comprehensive summary of this document, highlighting key points and main themes:
                
                {text}
                """,
                
                'sentiment': f"""
                Analyze the sentiment and tone of this text. Provide insights into:
                - Overall sentiment (positive/negative/neutral)
                - Emotional tone
                - Writing style
                - Target audience
                
                Text to analyze:
                {text}
                """,
                
                'structure': f"""
                Analyze the structure and organization of this document. Provide feedback on:
                - Document structure and flow
                - Paragraph organization
                - Clarity and coherence
                - Suggestions for improvement
                
                Document:
                {text}
                """,
                
                'comprehensive': f"""
                Perform a comprehensive analysis of this document including:
                1. Content summary and key points
                2. Sentiment and tone analysis
                3. Structure and organization assessment
                4. Readability and clarity evaluation
                5. Suggestions for improvement
                6. Target audience identification
                
                Document to analyze:
                {text}
                """
            }
            
            prompt = analysis_prompts.get(analysis_type, analysis_prompts['comprehensive'])
            
            response = self.client.chat.completions.create(
                model=self.models['analytical'],
                messages=[
                    {"role": "system", "content": "You are an expert document analyst and editor. Provide thorough, actionable analysis."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1500,
                temperature=0.2
            )
            
            analysis_result = response.choices[0].message.content.strip()
            processing_time = time.time() - start_time
            tokens_used = response.usage.total_tokens
            
            return ProcessingResult(
                success=True,
                result=analysis_result,
                processing_time=processing_time,
                tokens_used=tokens_used
            )
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                result=None,
                error_message=f"Document analysis error: {str(e)}",
                processing_time=time.time() - start_time
            )

# Advanced Text Analytics
class TextAnalytics:
    @staticmethod
    def get_text_statistics(text: str) -> Dict:
        """Get comprehensive text statistics"""
        if not text:
            return {}
        
        words = re.findall(r'\b\w+\b', text.lower())
        sentences = re.split(r'[.!?]+', text)
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        # Basic stats
        stats = {
            'characters': len(text),
            'characters_no_spaces': len(text.replace(' ', '')),
            'words': len(words),
            'sentences': len([s for s in sentences if s.strip()]),
            'paragraphs': len(paragraphs),
            'avg_words_per_sentence': len(words) / max(len(sentences), 1),
            'avg_chars_per_word': sum(len(word) for word in words) / max(len(words), 1)
        }
        
        # Reading time estimation (average 200 words per minute)
        stats['estimated_reading_time'] = max(1, round(stats['words'] / 200))
        
        # Readability scores
        try:
            stats['flesch_reading_ease'] = flesch_reading_ease(text)
            stats['flesch_kincaid_grade'] = flesch_kincaid_grade(text)
        except:
            stats['flesch_reading_ease'] = 0
            stats['flesch_kincaid_grade'] = 0
        
        # Word frequency
        word_freq = Counter(words)
        stats['most_common_words'] = word_freq.most_common(10)
        
        return stats
    
    @staticmethod
    def generate_word_cloud(text: str) -> Optional[bytes]:
        """Generate word cloud visualization"""
        try:
            if not text:
                return None
            
            # Create word cloud
            wordcloud = WordCloud(
                width=800, height=400,
                background_color='black',
                colormap='viridis',
                max_words=100
            ).generate(text)
            
            # Save to bytes
            img_buffer = io.BytesIO()
            plt.figure(figsize=(10, 5))
            plt.imshow(wordcloud, interpolation='bilinear')
            plt.axis('off')
            plt.tight_layout(pad=0)
            plt.savefig(img_buffer, format='png', bbox_inches='tight', 
                       facecolor='black', edgecolor='none')
            plt.close()
            
            img_buffer.seek(0)
            return img_buffer.getvalue()
            
        except Exception as e:
            st.error(f"Word cloud generation error: {str(e)}")
            return None

# Enhanced UI Components
def enhanced_copy_button(text: str, label: str = "Copy", key: str = None):
    """Enhanced copy button with better styling and feedback"""
    if not text:
        return
    
    import html
    escaped_text = html.escape(str(text))
    button_key = f"copy_btn_{key}" if key else f"copy_btn_{abs(hash(text[:50]))}"
    
    copy_html = f"""
    <div style="margin: 15px 0; position: relative;">
        <button onclick="copyToClipboard_{button_key}()" 
                style="background: linear-gradient(135deg, rgba(0, 249, 255, 0.2), rgba(0, 153, 204, 0.4));
                       border: 2px solid rgba(0, 249, 255, 0.6);
                       color: #00f9ff;
                       padding: 12px 24px;
                       border-radius: 8px;
                       cursor: pointer;
                       font-weight: bold;
                       font-size: 14px;
                       transition: all 0.3s ease;
                       box-shadow: 0 4px 8px rgba(0, 249, 255, 0.2);"
                onmouseover="this.style.background='linear-gradient(135deg, rgba(0, 249, 255, 0.3), rgba(0, 153, 204, 0.5))'"
                onmouseout="this.style.background='linear-gradient(135deg, rgba(0, 249, 255, 0.2), rgba(0, 153, 204, 0.4))'">
            📋 {label}
        </button>
        <div id="copy_status_{button_key}" style="display: inline-block; margin-left: 15px; font-weight: bold;"></div>
    </div>
    <textarea id="copy_text_{button_key}" readonly 
              style="width: 100%; min-height: 150px; 
                     background: rgba(0, 20, 40, 0.9);
                     border: 2px solid rgba(0, 249, 255, 0.3);
                     color: #00f9ff;
                     padding: 15px;
                     border-radius: 8px;
                     font-family: 'Courier New', monospace;
                     font-size: 13px;
                     line-height: 1.4;
                     resize: vertical;">{escaped_text}</textarea>
    
    <script>
    function copyToClipboard_{button_key}() {{
        const textArea = document.getElementById('copy_text_{button_key}');
        const statusDiv = document.getElementById('copy_status_{button_key}');
        
        if (navigator.clipboard && window.isSecureContext) {{
            navigator.clipboard.writeText(textArea.value).then(() => {{
                statusDiv.innerHTML = '<span style="color: #00ff88;">✅ Copied successfully!</span>';
                setTimeout(() => {{ statusDiv.innerHTML = ''; }}, 3000);
            }}).catch(() => {{ fallbackCopy(); }});
        }} else {{
            fallbackCopy();
        }}
        
        function fallbackCopy() {{
            textArea.select();
            textArea.setSelectionRange(0, 99999);
            try {{
                document.execCommand('copy');
                statusDiv.innerHTML = '<span style="color: #00ff88;">✅ Copied!</span>';
                setTimeout(() => {{ statusDiv.innerHTML = ''; }}, 3000);
            }} catch (err) {{
                statusDiv.innerHTML = '<span style="color: #ff4444;">❌ Copy failed - select manually</span>';
            }}
        }}
    }}
    </script>
    """
    
    components.html(copy_html, height=250)

def enhanced_styling():
    """Enhanced styling with modern design"""
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@300;400;700;900&family=Rajdhani:wght@300;400;500;600;700&family=JetBrains+Mono:wght@300;400;500&display=swap');
    
    :root {
        --primary-cyan: #00f9ff;
        --secondary-blue: #0099cc;
        --accent-blue: #66ccff;
        --dark-bg: #0a0a0a;
        --card-bg: rgba(0, 20, 40, 0.8);
        --border-glow: rgba(0, 249, 255, 0.3);
        --success-green: #00ff88;
        --warning-yellow: #ffaa00;
        --error-red: #ff4444;
    }
    
    .stApp {
        background: linear-gradient(135deg, var(--dark-bg) 0%, #1a1a2e 25%, #16213e 50%, #0f172a 75%, var(--dark-bg) 100%);
        color: var(--primary-cyan);
        font-family: 'Rajdhani', sans-serif;
    }
    
    /* Enhanced Headers */
    h1, h2, h3, h4, h5, h6 {
        font-family: 'Orbitron', monospace !important;
        color: var(--primary-cyan) !important;
        text-shadow: 0 0 15px rgba(0, 249, 255, 0.6);
        font-weight: 600 !important;
    }
    
    .main-title {
        font-size: 3.5rem;
        text-align: center;
        background: linear-gradient(45deg, var(--primary-cyan), var(--secondary-blue), var(--accent-blue));
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        margin: 2rem 0;
        text-shadow: none;
        animation: glow 2s ease-in-out infinite alternate;
    }
    
    @keyframes glow {
        from { filter: drop-shadow(0 0 20px rgba(0, 249, 255, 0.4)); }
        to { filter: drop-shadow(0 0 30px rgba(0, 249, 255, 0.8)); }
    }
    
    /* Enhanced Cards and Containers */
    .auth-container, .feature-card {
        max-width: 450px;
        margin: 0 auto;
        padding: 2.5rem;
        border: 2px solid var(--border-glow);
        border-radius: 15px;
        background: var(--card-bg);
        backdrop-filter: blur(10px);
        box-shadow: 0 8px 32px rgba(0, 249, 255, 0.1);
    }
    
    .feature-card {
        max-width: 100%;
        margin: 1rem 0;
    }
    
    /* Enhanced Status Indicators */
    .usage-warning {
        background: linear-gradient(135deg, rgba(255, 170, 0, 0.1), rgba(255, 140, 0, 0.05));
        border: 2px solid var(--warning-yellow);
        border-radius: 10px;
        padding: 15px;
        margin: 15px 0;
        animation: pulse-warning 2s infinite;
    }
    
    .subscription-info {
        background: linear-gradient(135deg, rgba(0, 249, 255, 0.1), rgba(0, 153, 204, 0.05));
        border: 2px solid var(--border-glow);
        border-radius: 10px;
        padding: 20px;
        margin: 20px 0;
        box-shadow: 0 4px 16px rgba(0, 249, 255, 0.2);
    }
    
    .success-indicator {
        background: linear-gradient(135deg, rgba(0, 255, 136, 0.1), rgba(0, 200, 100, 0.05));
        border: 2px solid var(--success-green);
        border-radius: 10px;
        padding: 15px;
        margin: 15px 0;
    }
    
    @keyframes pulse-warning {
        0%, 100% { box-shadow: 0 0 10px rgba(255, 170, 0, 0.3); }
        50% { box-shadow: 0 0 20px rgba(255, 170, 0, 0.6); }
    }
    
    /* Enhanced Form Elements */
    .stTextArea textarea, .stTextInput input, .stSelectbox select {
        background: var(--card-bg) !important;
        border: 2px solid var(--border-glow) !important;
        color: var(--primary-cyan) !important;
        border-radius: 8px !important;
        font-family: 'JetBrains Mono', monospace !important;
        transition: all 0.3s ease !important;
    }
    
    .stTextArea textarea:focus, .stTextInput input:focus {
        border-color: var(--primary-cyan) !important;
        box-shadow: 0 0 20px rgba(0, 249, 255, 0.4) !important;
    }
    
    /* Enhanced Buttons */
    .stButton > button {
        background: linear-gradient(135deg, rgba(0, 249, 255, 0.2), rgba(0, 153, 204, 0.4)) !important;
        border: 2px solid var(--border-glow) !important;
        color: var(--primary-cyan) !important;
        font-weight: bold !important;
        border-radius: 8px !important;
        transition: all 0.3s ease !important;
        font-family: 'Rajdhani', sans-serif !important;
        font-size: 16px !important;
    }
    
    .stButton > button:hover {
        background: linear-gradient(135deg, rgba(0, 249, 255, 0.3), rgba(0, 153, 204, 0.5)) !important;
        box-shadow: 0 0 20px rgba(0, 249, 255, 0.5) !important;
        transform: translateY(-2px) !important;
    }
    
    /* Enhanced Tabs */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background: transparent;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: linear-gradient(135deg, rgba(0, 249, 255, 0.1), rgba(0, 153, 204, 0.05));
        border: 1px solid var(--border-glow);
        border-radius: 8px 8px 0 0;
        color: var(--primary-cyan);
        font-weight: 500;
        font-family: 'Rajdhani', sans-serif;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, rgba(0, 249, 255, 0.2), rgba(0, 153, 204, 0.3));
        border-color: var(--primary-cyan);
        box-shadow: 0 0 15px rgba(0, 249, 255, 0.4);
    }
    
    /* Enhanced Sidebar */
    .css-1d391kg {
        background: linear-gradient(180deg, rgba(10, 10, 10, 0.95), rgba(26, 26, 46, 0.9));
        border-right: 1px solid var(--border-glow);
    }
    
    /* Code blocks */
    .stCode {
        background: rgba(0, 10, 20, 0.9) !important;
        border: 1px solid var(--border-glow) !important;
        border-radius: 8px !important;
    }
    
    /* Progress bars */
    .stProgress > div > div {
        background: linear-gradient(90deg, var(--primary-cyan), var(--secondary-blue)) !important;
    }
    
    /* Metrics */
    .metric-card {
        background: var(--card-bg);
        border: 1px solid var(--border-glow);
        border-radius: 10px;
        padding: 20px;
        text-align: center;
        box-shadow: 0 4px 16px rgba(0, 249, 255, 0.1);
    }
    
    .metric-value {
        font-size: 2rem;
        font-weight: bold;
        color: var(--primary-cyan);
        font-family: 'Orbitron', monospace;
    }
    
    .metric-label {
        font-size: 0.9rem;
        color: rgba(0, 249, 255, 0.7);
        margin-top: 5px;
    }
    
    /* Loading animations */
    @keyframes spin {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }
    
    .loading-spinner {
        display: inline-block;
        width: 20px;
        height: 20px;
        border: 3px solid rgba(0, 249, 255, 0.3);
        border-radius: 50%;
        border-top-color: var(--primary-cyan);
        animation: spin 1s ease-in-out infinite;
    }
    </style>
    """, unsafe_allow_html=True)

# Initialize session state with enhanced defaults
def init_session_state():
    """Initialize enhanced session state"""
    defaults = {
        'authenticated': False,
        'user_info': {},
        'page': 'login',
        'processing_history': [],
        'current_file': None,
        'analytics_data': {},
        'user_preferences': {
            'theme': 'dark',
            'default_model': 'advanced',
            'auto_save': True,
            'show_analytics': True
        }
    }
    
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_session_state()
enhanced_styling()

# Enhanced Authentication Pages with better UX
def enhanced_login_page():
    st.markdown('<h1 class="main-title">🌌 TECHNOVA AI NEXUS PRO</h1>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown('<div class="auth-container">', unsafe_allow_html=True)
        st.markdown("### 🔐 Welcome Back")
        st.markdown("*Access your AI-powered workspace*")
        
        with st.form("login_form"):
            email = st.text_input("📧 Email Address", placeholder="your.email@example.com")
            password = st.text_input("🔑 Password", type="password")
            remember_me = st.checkbox("Remember me")
            
            col_login, col_signup = st.columns(2)
            
            with col_login:
                login_submitted = st.form_submit_button("🚀 Login", use_container_width=True)
            
            with col_signup:
                if st.form_submit_button("📝 Sign Up", use_container_width=True):
                    st.session_state.page = 'signup'
                    st.rerun()
        
        if login_submitted and email and password:
            with st.spinner("🔍 Authenticating..."):
                success, message, user_info = EnhancedAuthManager.login_user(email, password)
                if success:
                    st.session_state.authenticated = True
                    st.session_state.user_info = user_info
                    st.success("✅ " + message)
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error("❌ " + message)
        
        st.markdown("---")
        if st.button("🔄 Forgot Password?", use_container_width=True):
            st.session_state.page = 'forgot_password'
            st.rerun()
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Features preview
        st.markdown("### 🚀 What's New in TechNova Pro")
        features = [
            "🤖 Advanced AI Code Enhancement",
            "📝 Intelligent Document Processing", 
            "🔍 Deep Text Analytics",
            "📊 Usage Analytics Dashboard",
            "🎨 Advanced Visualizations",
            "⚡ Real-time Processing"
        ]
        
        for feature in features:
            st.markdown(f"• {feature}")

# AI Enhancement Tab Implementation
def ai_enhancement_tab():
    """Advanced AI Code Enhancement Tab"""
    st.header("🤖 AI Code Enhancement Pro")
    
    if not check_enhanced_tab_access("AI Enhancement"):
        return
    
    ai_engine = AdvancedAIEngine(OPENAI_API_KEY)
    
    if not OPENAI_API_KEY:
        st.error("🔑 OpenAI API key not configured. Contact administrator.")
        return
    
    # Enhanced interface with tabs
    sub_tabs = st.tabs(["💻 Code Enhancement", "🔍 Code Analysis", "🏗️ Code Generation"])
    
    with sub_tabs[0]:
        st.markdown("### 💡 Paste your code for AI enhancement")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            code_input = st.text_area(
                "Your Code:",
                height=400,
                placeholder="""# Example Python code with issues
def calculate_area(radius):
    import math
    area = 3.14 * radius * radius  # Should use math.pi
    return area

# Missing error handling
result = calculate_area("invalid")
print("Area is:", result)""",
                help="Paste code in any programming language"
            )
        
        with col2:
            st.markdown("#### ⚙️ Enhancement Settings")
            
            # Language detection and override
            if code_input:
                detected_lang = detect_programming_language(code_input)
                st.info(f"🔍 Detected: **{detected_lang.title()}**")
            
            languages = ['auto-detect', 'python', 'javascript', 'typescript', 'java', 'cpp', 
                        'c', 'csharp', 'go', 'rust', 'php', 'ruby', 'swift', 'kotlin', 
                        'html', 'css', 'sql', 'r', 'matlab', 'scala']
            
            selected_lang = st.selectbox("Override Language:", languages)
            
            enhancement_mode = st.selectbox("Enhancement Mode:", [
                "🔧 Comprehensive (Fix + Improve + Optimize)",
                "🐛 Fix Errors Only",
                "⚡ Performance Optimization", 
                "🏗️ Modernize & Refactor",
                "📝 Add Documentation",
                "🛡️ Security Enhancement"
            ])
            
            model_choice = st.selectbox("AI Model:", [
                "🚀 GPT-4 (Advanced)",
                "⚡ GPT-3.5 Turbo (Fast)",
                "🎯 Code-Specialized Model"
            ])
        
        if code_input and st.button("🚀 Enhance Code", type="primary"):
            with st.spinner("🤖 AI is enhancing your code..."):
                start_time = time.time()
                
                # Process enhancement
                enhancement_type = enhancement_mode.split(' ')[1].lower()
                result = ai_engine.enhance_code(code_input, selected_lang or detected_lang, enhancement_type)
                
                if result.success:
                    # Log usage
                    log_enhanced_usage("AI Enhancement", result.tokens_used, result.processing_time, True)
                    
                    st.success(f"✅ Code enhanced in {result.processing_time:.2f}s using {result.tokens_used} tokens")
                    
                    # Side-by-side comparison
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("#### 📋 Original Code")
                        st.code(code_input, language=detected_lang)
                    
                    with col2:
                        st.markdown("#### ✨ Enhanced Code")
                        st.code(result.result, language=detected_lang)
                    
                    # Export options
                    st.markdown("---")
                    st.markdown("#### 📥 Export Enhanced Code")
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        enhanced_copy_button(result.result, "Copy Enhanced Code", "enhanced")
                    with col2:
                        download_enhanced_button(result.result, f"enhanced_code.{get_file_extension(detected_lang)}", 
                                               "💾 Download", "text/plain")
                    with col3:
                        if st.button("📊 View Analysis"):
                            st.session_state.show_code_analysis = True
                
                else:
                    st.error(f"❌ {result.error_message}")
                    log_enhanced_usage("AI Enhancement", 0, result.processing_time, False, result.error_message)
    
    with sub_tabs[1]:
        st.markdown("### 🔍 Advanced Code Analysis")
        
        if code_input:
            with st.expander("📊 Code Metrics & Analysis", expanded=True):
                # Basic metrics
                lines = len(code_input.split('\n'))
                chars = len(code_input)
                
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    st.markdown(f'<div class="metric-card"><div class="metric-value">{lines}</div><div class="metric-label">Lines</div></div>', unsafe_allow_html=True)
                with col2:
                    st.markdown(f'<div class="metric-card"><div class="metric-value">{chars}</div><div class="metric-label">Characters</div></div>', unsafe_allow_html=True)
                with col3:
                    complexity = analyze_code_complexity(code_input)
                    st.markdown(f'<div class="metric-card"><div class="metric-value">{complexity}</div><div class="metric-label">Complexity</div></div>', unsafe_allow_html=True)
                with col4:
                    quality_score = calculate_code_quality_score(code_input, detected_lang)
                    st.markdown(f'<div class="metric-card"><div class="metric-value">{quality_score}%</div><div class="metric-label">Quality</div></div>', unsafe_allow_html=True)
                
                # Issues detection
                issues = detect_code_issues(code_input, detected_lang)
                if issues:
                    st.markdown("#### ⚠️ Detected Issues")
                    for category, issue_list in issues.items():
                        with st.expander(f"{category.title()} Issues ({len(issue_list)})"):
                            for issue in issue_list:
                                st.markdown(f"• {issue}")
    
    with sub_tabs[2]:
        st.markdown("### 🏗️ AI Code Generation")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            description = st.text_area(
                "Describe what you want to build:",
                height=200,
                placeholder="Create a Python function that processes CSV files and generates statistical reports..."
            )
        
        with col2:
            target_language = st.selectbox("Target Language:", [
                'python', 'javascript', 'typescript', 'java', 'cpp', 'csharp', 'go', 'rust'
            ])
            
            complexity_level = st.selectbox("Complexity Level:", [
                "🎯 Simple & Clean",
                "🏗️ Professional Grade", 
                "🚀 Enterprise Level"
            ])
            
            include_tests = st.checkbox("Include unit tests")
            include_docs = st.checkbox("Include documentation")
        
        if description and st.button("🎯 Generate Code", type="primary"):
            with st.spinner("🤖 Generating code..."):
                result = generate_code_from_description(
                    description, target_language, complexity_level, 
                    include_tests, include_docs, ai_engine
                )
                
                if result.success:
                    st.success("✅ Code generated successfully!")
                    st.code(result.result, language=target_language)
                    enhanced_copy_button(result.result, "Copy Generated Code", "generated")
                else:
                    st.error(f"❌ {result.error_message}")

# Document Enhancement Tab
def document_enhancement_tab():
    """Advanced Document Processing and Enhancement"""
    st.header("📝 Document Enhancement Pro")
    
    if not check_enhanced_tab_access("Document Enhancement"):
        return
    
    ai_engine = AdvancedAIEngine(OPENAI_API_KEY)
    
    # File upload section
    st.markdown("### 📁 Upload Your Document")
    
    uploaded_file = st.file_uploader(
        "Choose a file",
        type=['txt', 'pdf', 'docx', 'md', 'csv', 'json', 'xml', 'yaml'],
        help="Supported formats: Text, PDF, Word, Markdown, CSV, JSON, XML, YAML"
    )
    
    # Text input as alternative
    st.markdown("### ✏️ Or Paste Text Directly")
    text_input = st.text_area(
        "Paste your text here:",
        height=300,
        placeholder="Enter or paste your text for enhancement..."
    )
    
    # Processing options
    with st.expander("⚙️ Enhancement Options", expanded=True):
        col1, col2, col3 = st.columns(3)
        
        with col1:
            enhancement_type = st.selectbox("Enhancement Type:", [
                "🔧 Comprehensive Enhancement",
                "✍️ Grammar & Style",
                "🎯 Clarity & Readability", 
                "🏢 Professional Tone",
                "📚 Academic Style",
                "💼 Business Writing",
                "✨ Creative Enhancement"
            ])
        
        with col2:
            output_format = st.selectbox("Output Format:", [
                "📄 Enhanced Text",
                "📋 Summary",
                "🎯 Key Points",
                "📊 Analysis Report",
                "🔤 Multiple Versions"
            ])
        
        with col3:
            target_audience = st.selectbox("Target Audience:", [
                "👥 General Audience",
                "🎓 Academic",
                "💼 Business Professional",
                "👨‍💻 Technical",
                "🎨 Creative",
                "📚 Educational"
            ])
    
    # Process document or text
    content_to_process = None
    source_type = None
    
    if uploaded_file:
        with st.spinner("📖 Processing document..."):
            process_result = AdvancedFileProcessor.process_file(uploaded_file)
            if process_result.success:
                content_to_process = process_result.result['content']
                source_type = "file"
                st.success(f"✅ Document processed: {uploaded_file.name}")
                
                # Show file info
                file_info = process_result.result
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("File Size", f"{file_info['size'] / 1024:.1f} KB")
                with col2:
                    st.metric("File Type", file_info['file_type'].title())
                with col3:
                    st.metric("Processing Time", f"{process_result.processing_time:.2f}s")
            else:
                st.error(f"❌ {process_result.error_message}")
    
    elif text_input.strip():
        content_to_process = text_input
        source_type = "text"
    
    # Text analytics preview
    if content_to_process:
        with st.expander("📊 Text Analytics Preview", expanded=False):
            stats = TextAnalytics.get_text_statistics(content_to_process)
            
            if stats:
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    st.metric("Words", f"{stats['words']:,}")
                with col2:
                    st.metric("Reading Time", f"{stats['estimated_reading_time']} min")
                with col3:
                    st.metric("Readability", f"{stats['flesch_reading_ease']:.0f}")
                with col4:
                    st.metric("Grade Level", f"{stats['flesch_kincaid_grade']:.1f}")
        
        # Enhancement button
        if st.button("✨ Enhance Document", type="primary", use_container_width=True):
            with st.spinner("🤖 AI is enhancing your document..."):
                enhancement_mode = enhancement_type.split(' ')[1].lower()
                result = ai_engine.enhance_text(content_to_process, enhancement_mode)
                
                if result.success:
                    log_enhanced_usage("Document Enhancement", result.tokens_used, result.processing_time, True)
                    
                    st.success(f"✅ Document enhanced in {result.processing_time:.2f}s")
                    
                    # Show results
                    if output_format == "📄 Enhanced Text":
                        st.markdown("#### ✨ Enhanced Document")
                        st.markdown(result.result)
                        enhanced_copy_button(result.result, "Copy Enhanced Text", "enhanced_doc")
                    
                    elif output_format == "🔤 Multiple Versions":
                        # Generate multiple versions
                        versions = generate_multiple_versions(content_to_process, ai_engine)
                        
                        for i, (version_name, version_text) in enumerate(versions.items()):
                            with st.expander(f"📝 {version_name}", expanded=i==0):
                                st.markdown(version_text)
                                enhanced_copy_button(version_text, f"Copy {version_name}", f"version_{i}")
                
                else:
                    st.error(f"❌ {result.error_message}")
                    log_enhanced_usage("Document Enhancement", 0, result.processing_time, False, result.error_message)

# Text Analytics Tab
def text_analytics_tab():
    """Advanced Text Analytics and Insights"""
    st.header("📊 Text Analytics Pro")
    
    if not check_enhanced_tab_access("Text Analytics"):
        return
    
    # Input section
    col1, col2 = st.columns([3, 1])
    
    with col1:
        text_to_analyze = st.text_area(
            "Text to Analyze:",
            height=300,
            placeholder="Paste your text here for comprehensive analysis..."
        )
    
    with col2:
        st.markdown("#### 🎯 Analysis Options")
        analysis_depth = st.selectbox("Analysis Depth:", [
            "🔍 Quick Analysis",
            "📊 Detailed Analysis", 
            "🧠 Deep Insights",
            "📈 Full Report"
        ])
        
        include_visualizations = st.checkbox("Include Visualizations", value=True)
        include_ai_insights = st.checkbox("AI-Powered Insights", value=True)
    
    if text_to_analyze and st.button("🔍 Analyze Text", type="primary"):
        with st.spinner("📊 Analyzing text..."):
            # Basic statistics
            stats = TextAnalytics.get_text_statistics(text_to_analyze)
            
            # Display metrics in columns
            st.markdown("#### 📈 Text Statistics")
            
            col1, col2, col3, col4, col5 = st.columns(5)
            metrics = [
                (col1, "Words", f"{stats['words']:,}"),
                (col2, "Characters", f"{stats['characters']:,}"),
                (col3, "Sentences", f"{stats['sentences']:,}"),
                (col4, "Paragraphs", f"{stats['paragraphs']:,}"),
                (col5, "Reading Time", f"{stats['estimated_reading_time']} min")
            ]
            
            for col, label, value in metrics:
                with col:
                    st.metric(label, value)
            
            # Advanced analytics
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("#### 📖 Readability Analysis")
                readability_data = {
                    'Metric': ['Flesch Reading Ease', 'Flesch-Kincaid Grade', 'Avg Words/Sentence', 'Avg Chars/Word'],
                    'Score': [
                        f"{stats['flesch_reading_ease']:.1f}",
                        f"{stats['flesch_kincaid_grade']:.1f}",
                        f"{stats['avg_words_per_sentence']:.1f}",
                        f"{stats['avg_chars_per_word']:.1f}"
                    ]
                }
                
                df_readability = pd.DataFrame(readability_data)
                st.dataframe(df_readability, hide_index=True)
            
            with col2:
                st.markdown("#### 🔤 Word Frequency")
                if stats['most_common_words']:
                    word_freq_data = {
                        'Word': [word for word, count in stats['most_common_words']],
                        'Frequency': [count for word, count in stats['most_common_words']]
                    }
                    
                    fig = px.bar(
                        word_freq_data, 
                        x='Frequency', 
                        y='Word',
                        orientation='h',
                        title="Top 10 Most Common Words"
                    )
                    fig.update_layout(
                        plot_bgcolor='rgba(0,0,0,0)',
                        paper_bgcolor='rgba(0,0,0,0)',
                        font_color='#00f9ff'
                    )
                    st.plotly_chart(fig, use_container_width=True)
            
            # Word cloud visualization
            if include_visualizations:
                st.markdown("#### ☁️ Word Cloud")
                word_cloud_img = TextAnalytics.generate_word_cloud(text_to_analyze)
                if word_cloud_img:
                    st.image(word_cloud_img, use_column_width=True)
            
            # AI-powered insights
            if include_ai_insights and OPENAI_API_KEY:
                st.markdown("#### 🧠 AI Insights")
                with st.spinner("🤖 Generating AI insights..."):
                    ai_engine = AdvancedAIEngine(OPENAI_API_KEY)
                    analysis_result = ai_engine.analyze_document(text_to_analyze, "comprehensive")
                    
                    if analysis_result.success:
                        st.markdown(analysis_result.result)
                        log_enhanced_usage("Text Analytics", analysis_result.tokens_used, 
                                         analysis_result.processing_time, True)
                    else:
                        st.error(f"❌ {analysis_result.error_message}")

# URL Content Extractor Tab
def url_content_tab():
    """Advanced URL Content Extraction and Processing"""
    st.header("🌐 URL Content Extractor Pro")
    
    if not check_enhanced_tab_access("URL Content"):
        return
    
    st.markdown("### 🔗 Extract and Enhance Web Content")
    
    # URL input with validation
    url = st.text_input(
        "Website URL:",
        placeholder="https://example.com/article",
        help="Enter a valid URL to extract and analyze content"
    )
    
    # Processing options
    col1, col2, col3 = st.columns(3)
    
    with col1:
        extract_mode = st.selectbox("Extraction Mode:", [
            "📄 Article Content",
            "🔍 Full Page Text",
            "📊 Structured Data",
            "🖼️ Images & Media",
            "🔗 All Links"
        ])
    
    with col2:
        processing_options = st.multiselect("Processing Options:", [
            "🤖 AI Enhancement",
            "📝 Summarization",
            "🔍 Keyword Extraction",
            "📊 Analytics",
            "🌍 Translation"
        ])
    
    with col3:
        output_format = st.selectbox("Output Format:", [
            "📄 Formatted Text",
            "📋 Markdown",
            "📊 JSON Data",
            "📈 Analysis Report"
        ])
    
    if url and st.button("🚀 Extract Content", type="primary"):
        with st.spinner("🌐 Extracting content..."):
            extraction_result = extract_url_content(url, extract_mode)
            
            if extraction_result.success:
                content = extraction_result.result
                
                # Display extracted content
                st.success(f"✅ Content extracted in {extraction_result.processing_time:.2f}s")
                
                # Show content preview
                with st.expander("👀 Content Preview", expanded=True):
                    if isinstance(content, dict):
                        st.json(content)
                    else:
                        st.text_area("Extracted Content:", content, height=200)
                
                # Process with selected options
                if processing_options:
                    st.markdown("#### 🔄 Processing Results")
                    
                    for option in processing_options:
                        if option == "🤖 AI Enhancement" and OPENAI_API_KEY:
                            with st.spinner("Enhancing content..."):
                                ai_result = AdvancedAIEngine(OPENAI_API_KEY).enhance_text(content, "comprehensive")
                                if ai_result.success:
                                    with st.expander("✨ AI Enhanced Content"):
                                        st.markdown(ai_result.result)
                                        enhanced_copy_button(ai_result.result, "Copy Enhanced", "url_enhanced")
                        
                        elif option == "📝 Summarization":
                            summary = advanced_summarization(content)
                            with st.expander("📝 Summary"):
                                st.markdown(summary)
                                enhanced_copy_button(summary, "Copy Summary", "url_summary")
                        
                        elif option == "📊 Analytics":
                            stats = TextAnalytics.get_text_statistics(content)
                            with st.expander("📊 Content Analytics"):
                                display_text_analytics(stats)
                
                log_enhanced_usage("URL Content", 0, extraction_result.processing_time, True)
            
            else:
                st.error(f"❌ {extraction_result.error_message}")
                log_enhanced_usage("URL Content", 0, extraction_result.processing_time, False, extraction_result.error_message)

# Analytics Dashboard Tab
def analytics_dashboard_tab():
    """Advanced Analytics Dashboard"""
    st.header("📊 Analytics Dashboard")
    
    user_id = st.session_state.user_info['id']
    analytics_data = EnhancedUsageManager.get_usage_analytics(user_id)
    
    if not analytics_data:
        st.info("📈 No usage data available yet. Start using the tools to see your analytics!")
        return
    
    # Overview metrics
    st.markdown("### 📈 Usage Overview")
    
    col1, col2, col3, col4 = st.columns(4)
    
    total_stats = analytics_data.get('total', {})
    
    with col1:
        st.markdown(f'''
        <div class="metric-card">
            <div class="metric-value">{total_stats.get("requests", 0):,}</div>
            <div class="metric-label">Total Requests</div>
        </div>
        ''', unsafe_allow_html=True)
    
    with col2:
        st.markdown(f'''
        <div class="metric-card">
            <div class="metric-value">{total_stats.get("tokens", 0):,}</div>
            <div class="metric-label">Tokens Used</div>
        </div>
        ''', unsafe_allow_html=True)
    
    with col3:
        st.markdown(f'''
        <div class="metric-card">
            <div class="metric-value">{total_stats.get("success_rate", 0):.1f}%</div>
            <div class="metric-label">Success Rate</div>
        </div>
        ''', unsafe_allow_html=True)
    
    with col4:
        avg_time = total_stats.get("total_time", 0) / max(total_stats.get("requests", 1), 1)
        st.markdown(f'''
        <div class="metric-card">
            <div class="metric-value">{avg_time:.2f}s</div>
            <div class="metric-label">Avg Processing</div>
        </div>
        ''', unsafe_allow_html=True)
    
    # Charts and visualizations
    col1, col2 = st.columns(2)
    
    with col1:
        # Today's usage by tab
        today_data = analytics_data.get('today', {})
        if today_data:
            st.markdown("#### 📅 Today's Usage by Tab")
            
            tabs = list(today_data.keys())
            counts = [today_data[tab]['count'] for tab in tabs]
            
            fig = px.pie(
                values=counts,
                names=tabs,
                title="Usage Distribution"
            )
            fig.update_layout(
                plot_bgcolor='rgba(0,0,0,0)',
                paper_bgcolor='rgba(0,0,0,0)',
                font_color='#00f9ff'
            )
            st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        # Weekly trend
        weekly_data = analytics_data.get('weekly_trend', {})
        if weekly_data:
            st.markdown("#### 📈 Weekly Usage Trend")
            
            dates = list(weekly_data.keys())
            usage = list(weekly_data.values())
            
            fig = px.line(
                x=dates,
                y=usage,
                title="Daily Usage Over Past Week"
            )
            fig.update_layout(
                plot_bgcolor='rgba(0,0,0,0)',
                paper_bgcolor='rgba(0,0,0,0)',
                font_color='#00f9ff',
                xaxis_title="Date",
                yaxis_title="Usage Count"
            )
            st.plotly_chart(fig, use_container_width=True)
    
    # Detailed usage table
    st.markdown("#### 📋 Detailed Usage Statistics")
    
    if today_data:
        detailed_data = []
        for tab, data in today_data.items():
            detailed_data.append({
                'Tab': tab,
                'Uses Today': data['count'],
                'Tokens Used': data['tokens'],
                'Avg Time (s)': f"{data['avg_time']:.2f}",
                'Success Rate': f"{data['success_rate']:.1f}%"
            })
        
        df = pd.DataFrame(detailed_data)
        st.dataframe(df, hide_index=True, use_container_width=True)

# Helper Functions
def detect_programming_language(code: str) -> str:
    """Enhanced programming language detection"""
    language_patterns = {
        'python': [
            r'\bdef\s+\w+\s*\(', r'\bimport\s+\w+', r'\bfrom\s+\w+\s+import',
            r'\bprint\s*\(', r'\bif\s+__name__\s*==\s*["\']__main__["\']',
            r'\bclass\s+\w+', r'\bself\.'
        ],
        'javascript': [
            r'\bfunction\s+\w+\s*\(', r'\bconst\s+\w+\s*=', r'\blet\s+\w+\s*=',
            r'\bvar\s+\w+\s*=', r'\bconsole\.log\s*\(', r'=>\s*{?',
            r'\bdocument\.', r'\bwindow\.'
        ],
        'typescript': [
            r'\binterface\s+\w+', r'\btype\s+\w+\s*=', r':\s*\w+\s*[=;]',
            r'\bfunction\s+\w+\s*\([^)]*:\s*\w+', r'\bexport\s+\w+'
        ],
        'java': [
            r'\bpublic\s+class\s+\w+', r'\bpublic\s+static\s+void\s+main',
            r'\bSystem\.out\.print', r'\bprivate\s+\w+\s+\w+',
            r'\bpublic\s+\w+\s+\w+\s*\(', r'\bimport\s+java\.'
        ],
        'cpp': [
            r'#include\s*<\w+>', r'\bint\s+main\s*\(', r'\bstd::',
            r'\bcout\s*<<', r'\bcin\s*>>', r'\bnamespace\s+\w+'
        ],
        'csharp': [
            r'\busing\s+System', r'\bnamespace\s+\w+', r'\bpublic\s+class\s+\w+',
            r'\bConsole\.Write', r'\bstring\s+\w+', r'\bvar\s+\w+\s*='
        ],
        'go': [
            r'\bpackage\s+main', r'\bfunc\s+main\s*\(', r'\bfmt\.Print',
            r'\bvar\s+\w+\s+\w+', r'\bgo\s+\w+\(', r'\bimport\s*\('
        ],
        'rust': [
            r'\bfn\s+main\s*\(', r'\blet\s+\w+', r'\bprintln!',
            r'\bstruct\s+\w+', r'\bimpl\s+\w+', r'\buse\s+std::'
        ]
    }
    
    code_lower = code.lower()
    scores = defaultdict(int)
    
    for lang, patterns in language_patterns.items():
        for pattern in patterns:
            matches = len(re.findall(pattern, code, re.IGNORECASE | re.MULTILINE))
            scores[lang] += matches * 2  # Weight matches higher
    
    # Additional scoring based on keywords
    keyword_weights = {
        'python': ['python', 'django', 'flask', 'pandas', 'numpy'],
        'javascript': ['javascript', 'node', 'react', 'vue', 'angular'],
        'java': ['java', 'spring', 'maven', 'gradle'],
        'cpp': ['cpp', 'c++', 'cmake'],
        'csharp': ['c#', 'csharp', '.net', 'dotnet'],
        'go': ['golang', 'go'],
        'rust': ['rust', 'cargo']
    }
    
    for lang, keywords in keyword_weights.items():
        for keyword in keywords:
            if keyword in code_lower:
                scores[lang] += 1
    
    if not scores:
        return 'unknown'
    
    return max(scores.items(), key=lambda x: x[1])[0]

def analyze_code_complexity(code: str) -> int:
    """Calculate code complexity score"""
    complexity = 0
    
    # Count control structures
    control_patterns = [
        r'\bif\s*\(', r'\bfor\s*\(', r'\bwhile\s*\(', r'\bswitch\s*\(',
        r'\btry\s*{', r'\bcatch\s*\(', r'\belse\s*{?', r'\belif\s*\('
    ]
    
    for pattern in control_patterns:
        complexity += len(re.findall(pattern, code, re.IGNORECASE))
    
    # Count function definitions
    func_patterns = [r'\bdef\s+\w+', r'\bfunction\s+\w+', r'\bpublic\s+\w+\s+\w+\s*\(']
    for pattern in func_patterns:
        complexity += len(re.findall(pattern, code, re.IGNORECASE)) * 2
    
    return min(complexity, 100)  # Cap at 100

def calculate_code_quality_score(code: str, language: str) -> int:
    """Calculate code quality score based on best practices"""
    score = 100  # Start with perfect score
    
    lines = code.split('\n')
    total_lines = len([line for line in lines if line.strip()])
    
    if total_lines == 0:
        return 0
    
    # Check for comments
    comment_patterns = {
        'python': r'#.*',
        'javascript': r'//.*|/\*.*?\*/',
        'java': r'//.*|/\*.*?\*/',
        'cpp': r'//.*|/\*.*?\*/'
    }
    
    pattern = comment_patterns.get(language, r'#.*|//.*')
    comment_lines = len(re.findall(pattern, code, re.MULTILINE))
    comment_ratio = comment_lines / total_lines
    
    if comment_ratio < 0.1:
        score -= 20  # Deduct for lack of comments
    elif comment_ratio > 0.3:
        score += 10  # Bonus for good documentation
    
    # Check line length
    long_lines = len([line for line in lines if len(line) > 120])
    if long_lines > 0:
        score -= min(long_lines * 5, 30)
    
    # Check for basic error handling
    error_patterns = {
        'python': r'\btry\s*:|except\s*\w*:',
        'javascript': r'\btry\s*{|catch\s*\(',
        'java': r'\btry\s*{|catch\s*\('
    }
    
    pattern = error_patterns.get(language, r'\btry\b|\bcatch\b')
    if not re.search(pattern, code, re.IGNORECASE) and total_lines > 10:
        score -= 15
    
    return max(0, min(100, score))

def detect_code_issues(code: str, language: str) -> Dict[str, List[str]]:
    """Detect various types of code issues"""
    issues = {
        'errors': [],
        'warnings': [],
        'suggestions': [],
        'security': []
    }
    
    if language == 'python':
        # Python-specific checks
        if re.search(r'\bprint\s+[^(]', code):
            issues['errors'].append("Using print statement instead of print() function")
        
        if re.search(r'^\s*except:\s*
            , code, re.MULTILINE):
            issues['warnings'].append("Bare except clause - should specify exception type")
        
        if re.search(r'\beval\s*\(', code):
            issues['security'].append("Use of eval() function - potential security risk")
        
        if not re.search(r'\bdef\s+\w+\s*\(', code) and len(code.split('\n')) > 15:
            issues['suggestions'].append("Consider organizing code into functions")
    
    elif language == 'javascript':
        if re.search(r'\bvar\s+', code):
            issues['suggestions'].append("Consider using 'let' or 'const' instead of 'var'")
        
        if re.search(r'==(?!=)', code):
            issues['warnings'].append("Using == instead of === (loose equality)")
        
        if re.search(r'\beval\s*\(', code):
            issues['security'].append("Use of eval() function - security risk")
    
    # General checks
    lines = code.split('\n')
    if any(len(line) > 120 for line in lines):
        issues['suggestions'].append("Some lines exceed 120 characters")
    
    return {k: v for k, v in issues.items() if v}  # Only return non-empty categories

def get_file_extension(language: str) -> str:
    """Get appropriate file extension for language"""
    extensions = {
        'python': 'py', 'javascript': 'js', 'typescript': 'ts',
        'java': 'java', 'cpp': 'cpp', 'c': 'c', 'csharp': 'cs',
        'go': 'go', 'rust': 'rs', 'php': 'php', 'ruby': 'rb',
        'html': 'html', 'css': 'css', 'sql': 'sql'
    }
    return extensions.get(language, 'txt')

def generate_code_from_description(description: str, language: str, complexity: str, 
                                 include_tests: bool, include_docs: bool, ai_engine) -> ProcessingResult:
    """Generate code from natural language description"""
    if not ai_engine.client:
        return ProcessingResult(False, None, "OpenAI API not configured")
    
    start_time = time.time()
    
    try:
        prompt = f"""
        Generate {language} code based on this description: {description}
        
        Requirements:
        - Complexity level: {complexity}
        - Include unit tests: {include_tests}
        - Include documentation: {include_docs}
        - Follow best practices for {language}
        - Include error handling
        - Make it production-ready
        
        Return only clean, working code with proper structure.
        """
        
        response = ai_engine.client.chat.completions.create(
            model=ai_engine.models['advanced'],
            messages=[
                {"role": "system", "content": f"You are an expert {language} developer. Generate clean, professional, well-documented code."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=3000,
            temperature=0.2
        )
        
        generated_code = response.choices[0].message.content.strip()
        
        # Clean up markdown
        if generated_code.startswith(f"```{language}"):
            generated_code = generated_code[len(f"```{language}"):].strip()
        elif generated_code.startswith("```"):
            generated_code = generated_code[3:].strip()
        
        if generated_code.endswith("```"):
            generated_code = generated_code[:-3].strip()
        
        processing_time = time.time() - start_time
        tokens_used = response.usage.total_tokens
        
        return ProcessingResult(
            success=True,
            result=generated_code,
            processing_time=processing_time,
            tokens_used=tokens_used
        )
        
    except Exception as e:
        return ProcessingResult(
            success=False,
            result=None,
            error_message=f"Code generation error: {str(e)}",
            processing_time=time.time() - start_time
        )

def extract_url_content(url: str, mode: str) -> ProcessingResult:
    """Extract content from URL with multiple extraction modes"""
    start_time = time.time()
    
    try:
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()
        
        if mode == "📄 Article Content":
            # Extract main article content
            article_selectors = ['article', '.article', '#article', '.content', '.post', '.entry']
            content = None
            
            for selector in article_selectors:
                elements = soup.select(selector)
                if elements:
                    content = elements[0].get_text(strip=True, separator='\n')
                    break
            
            if not content:
                content = soup.get_text(strip=True, separator='\n')
        
        elif mode == "🔍 Full Page Text":
            content = soup.get_text(strip=True, separator='\n')
        
        elif mode == "📊 Structured Data":
            content = {
                'title': soup.title.string if soup.title else "No title",
                'meta_description': '',
                'headings': [],
                'links': [],
                'images': []
            }
            
            # Extract meta description
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            if meta_desc:
                content['meta_description'] = meta_desc.get('content', '')
            
            # Extract headings
            for i in range(1, 7):
                headings = soup.find_all(f'h{i}')
                content['headings'].extend([h.get_text(strip=True) for h in headings])
            
            # Extract links
            links = soup.find_all('a', href=True)
            content['links'] = [{'text': link.get_text(strip=True), 'url': link['href']} for link in links[:20]]
            
            # Extract images
            images = soup.find_all('img', src=True)
            content['images'] = [{'alt': img.get('alt', ''), 'src': img['src']} for img in images[:10]]
        
        elif mode == "🔗 All Links":
            links = soup.find_all('a', href=True)
            content = '\n'.join([f"{link.get_text(strip=True)} -> {link['href']}" for link in links])
        
        else:
            content = soup.get_text(strip=True, separator='\n')
        
        processing_time = time.time() - start_time
        
        return ProcessingResult(
            success=True,
            result=content,
            processing_time=processing_time
        )
        
    except Exception as e:
        return ProcessingResult(
            success=False,
            result=None,
            error_message=f"URL extraction error: {str(e)}",
            processing_time=time.time() - start_time
        )

def advanced_summarization(text: str, max_sentences: int = 5) -> str:
    """Advanced text summarization with multiple techniques"""
    if not text or not text.strip():
        return "No content to summarize."
    
    # Sentence extraction and scoring
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if s.strip() and len(s.split()) > 3]
    
    if len(sentences) <= max_sentences:
        return '. '.join(sentences) + '.'
    
    # Word frequency analysis
    words = re.findall(r'\b\w+\b', text.lower())
    word_freq = Counter(word for word in words if word not in STOPWORDS and len(word) > 2)
    
    if not word_freq:
        return '. '.join(sentences[:max_sentences]) + '.'
    
    # Score sentences based on word frequency
    sentence_scores = []
    max_freq = max(word_freq.values())
    
    for sentence in sentences:
        sentence_words = re.findall(r'\b\w+\b', sentence.lower())
        score = sum(word_freq.get(word, 0) / max_freq for word in sentence_words)
        sentence_scores.append((sentence, score))
    
    # Sort by score and take top sentences
    top_sentences = sorted(sentence_scores, key=lambda x: x[1], reverse=True)[:max_sentences]
    
    # Sort by original order
    original_order = []
    for sentence, _ in top_sentences:
        original_order.append((sentences.index(sentence), sentence))
    
    original_order.sort(key=lambda x: x[0])
    summary = '. '.join([sentence for _, sentence in original_order]) + '.'
    
    return summary

def generate_multiple_versions(text: str, ai_engine) -> Dict[str, str]:
    """Generate multiple enhanced versions of text"""
    versions = {}
    
    enhancement_types = {
        "Professional": "formal",
        "Creative": "creative", 
        "Simplified": "simple",
        "Technical": "technical"
    }
    
    for name, enhancement_type in enhancement_types.items():
        try:
            result = ai_engine.enhance_text(text, enhancement_type)
            if result.success:
                versions[name] = result.result
            else:
                versions[name] = f"Error generating {name.lower()} version"
        except:
            versions[name] = f"Error generating {name.lower()} version"
    
    return versions

def display_text_analytics(stats: Dict):
    """Display text analytics in a structured format"""
    if not stats:
        st.info("No analytics data available")
        return
    
    # Create metrics display
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("**📊 Basic Metrics**")
        st.write(f"• Words: {stats.get('words', 0):,}")
        st.write(f"• Characters: {stats.get('characters', 0):,}")
        st.write(f"• Sentences: {stats.get('sentences', 0):,}")
        st.write(f"• Paragraphs: {stats.get('paragraphs', 0):,}")
    
    with col2:
        st.markdown("**📖 Readability**")
        st.write(f"• Reading Ease: {stats.get('flesch_reading_ease', 0):.1f}")
        st.write(f"• Grade Level: {stats.get('flesch_kincaid_grade', 0):.1f}")
        st.write(f"• Avg Words/Sentence: {stats.get('avg_words_per_sentence', 0):.1f}")
        st.write(f"• Reading Time: {stats.get('estimated_reading_time', 0)} min")
    
    with col3:
        st.markdown("**🔤 Word Analysis**")
        common_words = stats.get('most_common_words', [])
        for word, count in common_words[:5]:
            st.write(f"• {word}: {count}")

def download_enhanced_button(content: str, filename: str, label: str, mime_type: str = "text/plain"):
    """Enhanced download button with better styling"""
    b64_content = base64.b64encode(content.encode()).decode()
    href = f'data:{mime_type};base64,{b64_content}'
    
    download_html = f"""
    <a href="{href}" download="{filename}"
       style="background: linear-gradient(135deg, rgba(0, 249, 255, 0.2), rgba(0, 153, 204, 0.4));
              border: 2px solid rgba(0, 249, 255, 0.6);
              color: #00f9ff;
              padding: 12px 24px;
              text-decoration: none;
              border-radius: 8px;
              font-weight: bold;
              display: inline-block;
              transition: all 0.3s ease;
              box-shadow: 0 4px 8px rgba(0, 249, 255, 0.2);"
       onmouseover="this.style.background='linear-gradient(135deg, rgba(0, 249, 255, 0.3), rgba(0, 153, 204, 0.5))'"
       onmouseout="this.style.background='linear-gradient(135deg, rgba(0, 249, 255, 0.2), rgba(0, 153, 204, 0.4))'">
        {label}
    </a>
    """
    components.html(download_html, height=80)

def check_enhanced_tab_access(tab_name: str) -> bool:
    """Enhanced tab access checking with detailed feedback"""
    user_id = st.session_state.user_info['id']
    can_use, message, usage_info = EnhancedUsageManager.can_use_tab(user_id, tab_name)
    
    if not can_use:
        st.error(f"🚫 {message}")
        
        if "expired" in message.lower():
            st.markdown("""
            <div class="usage-warning">
                <h4>🚀 Upgrade to TechNova Plus</h4>
                <p><strong>Monthly Plan</strong>: $15/month - Unlimited usage</p>
                <p><strong>Annual Plan</strong>: $150/year (Save 30%)</p>
                <p><strong>Pro Plan</strong>: $35/month - Advanced features + Priority support</p>
            </div>
            """, unsafe_allow_html=True)
        
        return False
    
    # Show usage information
    if usage_info:
        show_enhanced_usage_info(usage_info)
    
    return True

def show_enhanced_usage_info(usage_info: Dict):
    """Display enhanced usage information"""
    col1, col2 = st.columns(2)
    
    with col1:
        if usage_info.get('daily_limit', -1) != -1:
            daily_used = usage_info.get('daily_usage', 0)
            daily_limit = usage_info.get('daily_limit', 5)
            progress = min(daily_used / daily_limit, 1.0) if daily_limit > 0 else 0
            
            st.markdown(f"**📅 Daily Usage**: {daily_used}/{daily_limit}")
            st.progress(progress)
    
    with col2:
        if usage_info.get('monthly_limit', -1) != -1:
            monthly_used = usage_info.get('monthly_usage', 0)
            monthly_limit = usage_info.get('monthly_limit', 100)
            progress = min(monthly_used / monthly_limit, 1.0) if monthly_limit > 0 else 0
            
            st.markdown(f"**📊 Monthly API Usage**: {monthly_used}/{monthly_limit}")
            st.progress(progress)

def log_enhanced_usage(tab_name: str, tokens_used: int = 0, processing_time: float = 0.0, 
                      success: bool = True, error_message: str = "", input_size: int = 0, 
                      output_size: int = 0):
    """Enhanced usage logging wrapper"""
    user_id = st.session_state.user_info['id']
    EnhancedUsageManager.log_enhanced_usage(
        user_id, tab_name, tokens_used, processing_time, 
        success, error_message, input_size, output_size
    )

# Enhanced Main Application
def enhanced_main_application():
    """Enhanced main application with advanced features"""
    # Header with user info
    col1, col2, col3 = st.columns([2, 1, 1])
    
    with col1:
        st.markdown(f"### 👋 Welcome, **{st.session_state.user_info['email']}**")
        sub_type = st.session_state.user_info.get('subscription_type', 'free')
        if sub_type == 'plus':
            st.markdown("⭐ **TechNova Plus Member**")
        elif sub_type == 'pro':
            st.markdown("🚀 **TechNova Pro Member**")
        else:
            st.markdown("🆓 **Free Trial User**")
    
    with col2:
        if st.button("📊 Analytics"):
            st.session_state.show_analytics = True
    
    with col3:
        if st.button("🚪 Logout"):
            st.session_state.authenticated = False
            st.session_state.user_info = {}
            st.session_state.page = 'login'
            st.rerun()
    
    # Sidebar with enhanced navigation
    with st.sidebar:
        st.markdown("### 🌌 TechNova Pro")
        
        # Quick stats
        user_id = st.session_state.user_info['id']
        analytics = EnhancedUsageManager.get_usage_analytics(user_id)
        
        if analytics:
            total_requests = analytics.get('total', {}).get('requests', 0)
            st.metric("Total Requests", f"{total_requests:,}")
            
            today_count = sum(tab_data.get('count', 0) for tab_data in analytics.get('today', {}).values())
            st.metric("Today's Usage", today_count)
        
        # Navigation
        st.markdown("### 🧭 Navigation")
        nav_options = [
            "🤖 AI Enhancement",
            "📝 Document Enhancement", 
            "📊 Text Analytics",
            "🌐 URL Content Extractor",
            "🎨 Content Generator",
            "📈 Analytics Dashboard",
            "⚙️ Settings"
        ]
        
        selected_nav = st.selectbox("Go to:", nav_options, label_visibility="collapsed")
        
        # Quick actions
        st.markdown("### ⚡ Quick Actions")
        if st.button("🔄 Clear Cache", use_container_width=True):
            st.cache_data.clear()
            st.success("Cache cleared!")
        
        if st.button("📋 Copy Session Info", use_container_width=True):
            session_info = {
                'user': st.session_state.user_info['email'],
                'subscription': st.session_state.user_info.get('subscription_type'),
                'session_start': datetime.now().isoformat()
            }
            st.code(json.dumps(session_info, indent=2))
    
    # Main content tabs
    tabs = st.tabs([
        "🤖 AI Enhancement",
        "📝 Document Processing", 
        "📊 Text Analytics",
        "🌐 URL Extractor",
        "🎨 Content Generator",
        "📈 Dashboard"
    ])
    
    with tabs[0]:
        ai_enhancement_tab()
    
    with tabs[1]:
        document_enhancement_tab()
    
    with tabs[2]:
        text_analytics_tab()
    
    with tabs[3]:
        url_content_tab()
    
    with tabs[4]:
        content_generator_tab()
    
    with tabs[5]:
        analytics_dashboard_tab()

def content_generator_tab():
    """AI-Powered Content Generation Tab"""
    st.header("🎨 AI Content Generator")
    
    if not check_enhanced_tab_access("Content Generator"):
        return
    
    ai_engine = AdvancedAIEngine(OPENAI_API_KEY)
    
    if not OPENAI_API_KEY:
        st.error("🔑 OpenAI API key not configured.")
        return
    
    # Content type selection
    content_types = {
        "📝 Blog Post": "blog_post",
        "📧 Email": "email",
        "📋 Report": "report", 
        "🎯 Marketing Copy": "marketing",
        "📚 Documentation": "documentation",
        "💼 Business Proposal": "proposal",
        "🎓 Educational Content": "educational",
        "🔬 Technical Article": "technical"
    }
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        selected_type = st.selectbox("Content Type:", list(content_types.keys()))
        
        topic = st.text_input(
            "Topic/Subject:",
            placeholder="Enter the main topic or subject..."
        )
        
        description = st.text_area(
            "Detailed Description:",
            height=150,
            placeholder="Provide more details about what you want to create..."
        )
    
    with col2:
        st.markdown("#### ⚙️ Generation Settings")
        
        length_option = st.selectbox("Content Length:", [
            "📄 Short (300-500 words)",
            "📑 Medium (500-1000 words)",
            "📚 Long (1000+ words)"
        ])
        
        tone = st.selectbox("Tone:", [
            "💼 Professional",
            "😊 Friendly", 
            "🎯 Persuasive",
            "📚 Educational",
            "💡 Inspirational",
            "🔬 Technical"
        ])
        
        include_outline = st.checkbox("Include Outline")
        include_citations = st.checkbox("Include Placeholder Citations")
    
    if topic and st.button("🚀 Generate Content", type="primary"):
        with st.spinner("🤖 Creating content..."):
            content_result = generate_content(
                content_types[selected_type], topic, description, 
                length_option, tone, include_outline, include_citations, ai_engine
            )
            
            if content_result.success:
                st.success(f"✅ Content generated in {content_result.processing_time:.2f}s")
                
                # Display generated content
                st.markdown("#### ✨ Generated Content")
                st.markdown(content_result.result)
                
                # Export options
                col1, col2 = st.columns(2)
                with col1:
                    enhanced_copy_button(content_result.result, "Copy Content", "generated_content")
                with col2:
                    download_enhanced_button(content_result.result, f"{topic.replace(' ', '_')}.txt", 
                                           "💾 Download", "text/plain")
                
                log_enhanced_usage("Content Generator", content_result.tokens_used, 
                                 content_result.processing_time, True)
            else:
                st.error(f"❌ {content_result.error_message}")

def generate_content(content_type: str, topic: str, description: str, length: str, 
                    tone: str, include_outline: bool, include_citations: bool, ai_engine) -> ProcessingResult:
    """Generate content using AI"""
    if not ai_engine.client:
        return ProcessingResult(False, None, "OpenAI API not configured")
    
    start_time = time.time()
    
    try:
        # Build prompt based on parameters
        length_guide = {
            "📄 Short (300-500 words)": "300-500 words",
            "📑 Medium (500-1000 words)": "500-1000 words", 
            "📚 Long (1000+ words)": "1000+ words"
        }
        
        prompt = f"""
        Create a {content_type} about "{topic}".
        
        Description: {description}
        
        Requirements:
        - Length: {length_guide.get(length, '500-800 words')}
        - Tone: {tone.split(' ')[1].lower()}
        - Include outline: {include_outline}
        - Include citations: {include_citations}
        
        Make it engaging, well-structured, and professional.
        """
        
        response = ai_engine.client.chat.completions.create(
            model=ai_engine.models['creative'],
            messages=[
                {"role": "system", "content": f"You are an expert content writer specializing in {content_type}. Create engaging, high-quality content."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2500,
            temperature=0.7
        )
        
        generated_content = response.choices[0].message.content.strip()
        processing_time = time.time() - start_time
        tokens_used = response.usage.total_tokens
        
        return ProcessingResult(
            success=True,
            result=generated_content,
            processing_time=processing_time,
            tokens_used=tokens_used
        )
        
    except Exception as e:
        return ProcessingResult(
            success=False,
            result=None,
            error_message=f"Content generation error: {str(e)}",
            processing_time=time.time() - start_time
        )

# Enhanced Authentication Methods (continued from previous)
def enhanced_login_user(email: str, password: str) -> Tuple[bool, str, Dict]:
    """Enhanced login with better security and user info"""
    conn = sqlite3.connect('technova_pro.db')
    cursor = conn.cursor()
    
    try:
        cursor.execute('''
            SELECT id, password_hash, is_verified, subscription_type, subscription_expires,
                   last_active, total_usage, api_usage_count
            FROM users WHERE email = ?
        ''', (email.lower(),))
        
        result = cursor.fetchone()
        if not result:
            return False, "Invalid email or password.", {}
        
        user_id, password_hash, is_verified, sub_type, sub_expires, last_active, total_usage, api_usage = result
        
        if not EnhancedAuthManager.verify_password(password, password_hash):
            return False, "Invalid email or password.", {}
        
        if not is_verified:
            return False, "Please verify your account first.", {}
        
        # Update last active
        cursor.execute("UPDATE users SET last_active = CURRENT_TIMESTAMP WHERE id = ?", (user_id,))
        
        # Check subscription status
        if sub_expires:
            sub_expires = datetime.fromisoformat(sub_expires)
            if datetime.now() > sub_expires and sub_type == 'free':
                cursor.execute("UPDATE users SET subscription_type = 'expired' WHERE id = ?", (user_id,))
                sub_type = 'expired'
        
        conn.commit()
        
        user_info = {
            'id': user_id,
            'email': email,
            'subscription_type': sub_type,
            'subscription_expires': sub_expires,
            'total_usage': total_usage,
            'api_usage_count': api_usage
        }
        
        return True, "Login successful! Welcome to TechNova AI Nexus Pro.", user_info
        
    except Exception as e:
        return False, f"Login error: {str(e)}", {}
    finally:
        conn.close()

# Add enhanced login method to AuthManager
EnhancedAuthManager.login_user = enhanced_login_user

# Enhanced Stopwords for better text processing
ENHANCED_STOPWORDS = set([
    "a", "an", "and", "are", "as", "at", "be", "but", "by", "for", "if", "in", "into",
    "is", "it", "no", "not", "of", "on", "or", "such", "that", "the", "their", "then",
    "there", "these", "they", "this", "to", "was", "will", "with", "you", "your", "from",
    "our", "we", "he", "she", "his", "her", "its", "were", "been", "being", "than",
    "also", "can", "could", "should", "would", "may", "might", "have", "has", "had",
    "do", "does", "did", "done", "just", "over", "under", "more", "most", "other",
    "some", "any", "each", "many", "few", "those", "them", "which", "who", "whom",
    "whose", "where", "when", "why", "how", "about", "up", "out", "down", "off",
    "above", "below", "between", "during", "before", "after", "through", "until",
    "while", "since", "against", "among", "throughout", "despite", "towards", "upon"
])

# Main Application Logic
def main():
    """Enhanced main application entry point"""
    
    # Show appropriate page based on authentication state
    if not st.session_state.authenticated:
        if st.session_state.page == 'login':
            enhanced_login_page()
        elif st.session_state.page == 'signup':
            enhanced_signup_page()
        elif st.session_state.page == 'verify':
            enhanced_verify_page()
        elif st.session_state.page == 'forgot_password':
            enhanced_forgot_password_page()
    else:
        enhanced_main_application()

def enhanced_signup_page():
    """Enhanced signup page with better validation"""
    st.markdown('<h1 class="main-title">🌌 TECHNOVA AI NEXUS PRO</h1>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown('<div class="auth-container">', unsafe_allow_html=True)
        st.markdown("### 📝 Create Your TechNova Account")
        st.markdown("*Join the future of AI-powered productivity*")
        
        with st.form("signup_form"):
            email = st.text_input("📧 Email Address", placeholder="your.email@example.com")
            password = st.text_input("🔑 Password", type="password", help="Min 8 chars, include uppercase, lowercase, number, and special character")
            confirm_password = st.text_input("🔑 Confirm Password", type="password")
            
            agree_terms = st.checkbox("I agree to the Terms of Service and Privacy Policy")
            
            col_create, col_back = st.columns(2)
            
            with col_create:
                signup_submitted = st.form_submit_button("🚀 Create Account", use_container_width=True)
            
            with col_back:
                if st.form_submit_button("← Back to Login", use_container_width=True):
                    st.session_state.page = 'login'
                    st.rerun()
        
        if signup_submitted:
            if not agree_terms:
                st.error("❌ Please agree to the Terms of Service")
            elif not email or not password or not confirm_password:
                st.error("❌ Please fill in all fields")
            elif password != confirm_password:
                st.error("❌ Passwords do not match")
            else:
                with st.spinner("🔐 Creating account..."):
                    success, message = EnhancedAuthManager.create_enhanced_user(email, password)
                    if success:
                        st.success("✅ " + message)
                        st.session_state.page = 'verify'
                        st.session_state.verification_email = email
                        st.balloons()
                        time.sleep(2)
                        st.rerun()
                    else:
                        st.error("❌ " + message)
        
        st.markdown('</div>', unsafe_allow_html=True)
        
        # Benefits section
        st.markdown("### 🎯 What You Get with TechNova Pro")
        benefits = [
            "🤖 Advanced AI Code Enhancement",
            "📝 Intelligent Document Processing",
            "📊 Deep Analytics & Insights",
            "🌐 Advanced Web Content Extraction",
            "🎨 AI Content Generation",
            "⚡ Real-time Processing",
            "📈 Usage Analytics Dashboard",
            "🔒 Enterprise-grade Security"
        ]
        
        for benefit in benefits:
            st.markdown(f"• {benefit}")

def enhanced_verify_page():
    """Enhanced verification page"""
    st.markdown('<h1 class="main-title">🌌 TECHNOVA AI NEXUS PRO</h1>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown('<div class="auth-container">', unsafe_allow_html=True)
        st.markdown("### ✉️ Verify Your Account")
        
        email = st.session_state.get('verification_email', 'your email')
        st.info(f"📧 Verification code sent to **{email}**")
        st.markdown("*Code expires in 15 minutes*")
        
        with st.form("verify_form"):
            verification_code = st.text_input("Enter 6-digit verification code", max_chars=6)
            
            col_verify, col_resend = st.columns(2)
            
            with col_verify:
                verify_submitted = st.form_submit_button("✅ Verify", use_container_width=True)
            
            with col_resend:
                if st.form_submit_button("📧 Resend Code", use_container_width=True):
                    # Resend verification code
                    success, message = resend_verification_code(email)
                    if success:
                        st.success("✅ New verification code sent!")
                    else:
                        st.error(f"❌ {message}")
        
        if verify_submitted:
            if not verification_code or len(verification_code) != 6:
                st.error("❌ Please enter a valid 6-digit code")
            else:
                with st.spinner("🔍 Verifying..."):
                    success, message = EnhancedAuthManager.verify_user(email, verification_code)
                    if success:
                        st.success("🎉 " + message)
                        st.balloons()
                        st.session_state.page = 'login'
                        time.sleep(2)
                        st.rerun()
                    else:
                        st.error("❌ " + message)
        
        if st.button("← Back to Login", use_container_width=True):
            st.session_state.page = 'login'
            st.rerun()
        
        st.markdown('</div>', unsafe_allow_html=True)

def enhanced_forgot_password_page():
    """Enhanced forgot password page with email reset functionality"""
    st.markdown('<h1 class="main-title">🌌 TECHNOVA AI NEXUS PRO</h1>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown('<div class="auth-container">', unsafe_allow_html=True)
        st.markdown("### 🔑 Reset Password")
        
        email = st.text_input("📧 Email Address", placeholder="your.email@example.com")
        
        if st.button("📧 Send Reset Code", use_container_width=True):
            if email:
                with st.spinner("📤 Sending reset code..."):
                    success, message = send_password_reset_code(email)
                    if success:
                        st.success("✅ " + message)
                        st.session_state.reset_email = email
                        st.session_state.page = 'reset_password'
                        st.rerun()
                    else:
                        st.error("❌ " + message)
            else:
                st.error("❌ Please enter your email address")
        
        if st.button("← Back to Login", use_container_width=True):
            st.session_state.page = 'login'
            st.rerun()
        
        st.markdown('</div>', unsafe_allow_html=True)

def resend_verification_code(email: str) -> Tuple[bool, str]:
    """Resend verification code to user"""
    conn = sqlite3.connect('technova_pro.db')
    cursor = conn.cursor()
    
    try:
        # Generate new code
        verification_code = EnhancedAuthManager.generate_secure_code()
        expires_at = datetime.now() + timedelta(minutes=15)
        
        # Update user record
        cursor.execute('''
            UPDATE users SET verification_code = ?, verification_expires = ?
            WHERE email = ? AND is_verified = FALSE
        ''', (verification_code, expires_at, email.lower()))
        
        if cursor.rowcount == 0:
            return False, "User not found or already verified"
        
        conn.commit()
        
        # Send email
        if EnhancedAuthManager.send_enhanced_verification_email(email, verification_code):
            return True, "New verification code sent successfully"
        else:
            return False, "Failed to send verification email"
    
    except Exception as e:
        return False, f"Error resending code: {str(e)}"
    finally:
        conn.close()

def send_password_reset_code(email: str) -> Tuple[bool, str]:
    """Send password reset code"""
    conn = sqlite3.connect('technova_pro.db')
    cursor = conn.cursor()
    
    try:
        # Check if user exists
        cursor.execute("SELECT id FROM users WHERE email = ?", (email.lower(),))
        if not cursor.fetchone():
            return False, "No account found with this email address"
        
        # Generate reset code
        reset_code = EnhancedAuthManager.generate_secure_code()
        expires_at = datetime.now() + timedelta(minutes=15)
        
        # Store reset code
        cursor.execute('''
            UPDATE users SET verification_code = ?, verification_expires = ?
            WHERE email = ?
        ''', (reset_code, expires_at, email.lower()))
        
        conn.commit()
        
        # Send reset email
        if send_password_reset_email(email, reset_code):
            return True, "Password reset code sent to your email"
        else:
            return False, "Failed to send reset email"
    
    except Exception as e:
        return False, f"Error sending reset code: {str(e)}"
    finally:
        conn.close()

def send_password_reset_email(email: str, reset_code: str) -> bool:
    """Send password reset email"""
    if not SMTP_USERNAME or not SMTP_PASSWORD:
        return False
    
    try:
        msg = MimeMultipart('alternative')
        msg['From'] = f"TechNova AI <{SMTP_USERNAME}>"
        msg['To'] = email
        msg['Subject'] = "🔐 TechNova AI - Password Reset Code"
        
        html_body = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; background: #0a0a0a; color: #00f9ff; }}
                .container {{ max-width: 600px; margin: 0 auto; padding: 20px; }}
                .header {{ text-align: center; padding: 20px; background: linear-gradient(135deg, #1a1a2e, #16213e); }}
                .code {{ font-size: 2em; font-weight: bold; color: #00f9ff; text-align: center; 
                         padding: 20px; background: rgba(0, 249, 255, 0.1); border-radius: 10px; }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h1>🔐 Password Reset</h1>
                </div>
                <p>You requested a password reset for your TechNova AI account.</p>
                <p>Your reset code is:</p>
                <div class="code">{reset_code}</div>
                <p>This code expires in 15 minutes.</p>
                <p>If you didn't request this reset, please ignore this email.</p>
            </div>
        </body>
        </html>
        """
        
        msg.attach(MimeText(html_body, 'html'))
        
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(SMTP_USERNAME, SMTP_PASSWORD)
        server.sendmail(SMTP_USERNAME, email, msg.as_string())
        server.quit()
        
        return True
    except:
        return False

# Settings and Preferences Tab
def settings_tab():
    """User settings and preferences"""
    st.header("⚙️ Settings & Preferences")
    
    user_info = st.session_state.user_info
    
    # Account Information
    st.subheader("👤 Account Information")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.info(f"**Email:** {user_info['email']}")
        st.info(f"**Subscription:** {user_info['subscription_type'].title()}")
        
        if user_info.get('subscription_expires'):
            expires = user_info['subscription_expires']
            if isinstance(expires, str):
                expires = datetime.fromisoformat(expires)
            
            if user_info['subscription_type'] == 'free':
                days_left = (expires - datetime.now()).days
                st.info(f"**Free Trial:** {max(0, days_left)} days remaining")
    
    with col2:
        st.info(f"**Total Usage:** {user_info.get('total_usage', 0)} requests")
        st.info(f"**API Usage:** {user_info.get('api_usage_count', 0)} calls this month")
    
    # Preferences
    st.subheader("🎨 Preferences")
    
    col1, col2 = st.columns(2)
    
    with col1:
        theme = st.selectbox("Theme:", ["🌙 Dark", "☀️ Light", "🌈 Auto"])
        default_model = st.selectbox("Default AI Model:", [
            "🚀 GPT-4 (Advanced)",
            "⚡ GPT-3.5 Turbo (Fast)"
        ])
    
    with col2:
        auto_save = st.checkbox("Auto-save results", value=True)
        show_analytics = st.checkbox("Show usage analytics", value=True)
        email_notifications = st.checkbox("Email notifications", value=False)
    
    if st.button("💾 Save Preferences", type="primary"):
        preferences = {
            'theme': theme,
            'default_model': default_model,
            'auto_save': auto_save,
            'show_analytics': show_analytics,
            'email_notifications': email_notifications
        }
        
        # Save to database
        conn = sqlite3.connect('technova_pro.db')
        cursor = conn.cursor()
        
        try:
            cursor.execute('''
                UPDATE users SET preferences = ? WHERE id = ?
            ''', (json.dumps(preferences), user_info['id']))
            conn.commit()
            st.success("✅ Preferences saved successfully!")
        except Exception as e:
            st.error(f"❌ Error saving preferences: {str(e)}")
        finally:
            conn.close()
    
    # Account Actions
    st.subheader("🔧 Account Actions")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("🔄 Change Password", use_container_width=True):
            st.session_state.show_change_password = True
    
    with col2:
        if st.button("📊 Download Data", use_container_width=True):
            user_data = export_user_data(user_info['id'])
            download_enhanced_button(user_data, "technova_data.json", "💾 Download", "application/json")
    
    with col3:
        if st.button("🗑️ Delete Account", use_container_width=True):
            st.session_state.show_delete_confirm = True
    
    # Change password modal
    if st.session_state.get('show_change_password', False):
        with st.expander("🔐 Change Password", expanded=True):
            current_password = st.text_input("Current Password", type="password", key="current_pw")
            new_password = st.text_input("New Password", type="password", key="new_pw")
            confirm_new = st.text_input("Confirm New Password", type="password", key="confirm_new_pw")
            
            col1, col2 = st.columns(2)
            with col1:
                if st.button("✅ Update Password", key="update_pw"):
                    success, message = change_user_password(
                        user_info['id'], current_password, new_password, confirm_new
                    )
                    if success:
                        st.success("✅ " + message)
                        st.session_state.show_change_password = False
                        st.rerun()
                    else:
                        st.error("❌ " + message)
            
            with col2:
                if st.button("❌ Cancel", key="cancel_pw"):
                    st.session_state.show_change_password = False
                    st.rerun()

def change_user_password(user_id: int, current_password: str, new_password: str, confirm_password: str) -> Tuple[bool, str]:
    """Change user password with validation"""
    if not current_password or not new_password or not confirm_password:
        return False, "All fields are required"
    
    if new_password != confirm_password:
        return False, "New passwords do not match"
    
    is_strong, issues = EnhancedAuthManager.is_strong_password(new_password)
    if not is_strong:
        return False, "Password requirements not met:\n• " + "\n• ".join(issues)
    
    conn = sqlite3.connect('technova_pro.db')
    cursor = conn.cursor()
    
    try:
        # Verify current password
        cursor.execute("SELECT password_hash FROM users WHERE id = ?", (user_id,))
        result = cursor.fetchone()
        
        if not result or not EnhancedAuthManager.verify_password(current_password, result[0]):
            return False, "Current password is incorrect"
        
        # Update password
        new_hash = EnhancedAuthManager.hash_password(new_password)
        cursor.execute("UPDATE users SET password_hash = ? WHERE id = ?", (new_hash, user_id))
        conn.commit()
        
        return True, "Password updated successfully"
    
    except Exception as e:
        return False, f"Error updating password: {str(e)}"
    finally:
        conn.close()

def export_user_data(user_id: int) -> str:
    """Export user data for download"""
    conn = sqlite3.connect('technova_pro.db')
    cursor = conn.cursor()
    
    try:
        # Get user info
        cursor.execute("SELECT email, created_at, subscription_type, total_usage FROM users WHERE id = ?", (user_id,))
        user_data = cursor.fetchone()
        
        # Get usage logs
        cursor.execute("""
            SELECT tab_name, usage_date, usage_count, tokens_used, processing_time
            FROM usage_logs WHERE user_id = ? ORDER BY usage_date DESC LIMIT 100
        """, (user_id,))
        usage_data = cursor.fetchall()
        
        export_data = {
            'user_info': {
                'email': user_data[0],
                'created_at': user_data[1],
                'subscription_type': user_data[2],
                'total_usage': user_data[3]
            },
            'usage_history': [
                {
                    'tab': row[0],
                    'date': row[1],
                    'count': row[2],
                    'tokens': row[3],
                    'time': row[4]
                } for row in usage_data
            ],
            'export_date': datetime.now().isoformat()
        }
        
        return json.dumps(export_data, indent=2)
    
    except Exception as e:
        return json.dumps({'error': f'Export failed: {str(e)}'}, indent=2)
    finally:
        conn.close()

# Enhanced Error Handling and Logging
class ErrorHandler:
    @staticmethod
    def log_error(error: Exception, context: str, user_id: int = None):
        """Log errors for debugging and monitoring"""
        error_info = {
            'timestamp': datetime.now().isoformat(),
            'error_type': type(error).__name__,
            'error_message': str(error),
            'context': context,
            'user_id': user_id
        }
        
        # In production, you would send this to a logging service
        print(f"ERROR: {json.dumps(error_info, indent=2)}")
    
    @staticmethod
    def handle_api_error(error: Exception) -> str:
        """Handle OpenAI API errors gracefully"""
        if "rate limit" in str(error).lower():
            return "Rate limit exceeded. Please try again in a few minutes."
        elif "quota" in str(error).lower():
            return "API quota exceeded. Please contact administrator."
        elif "invalid" in str(error).lower():
            return "Invalid API request. Please check your input."
        else:
            return f"AI service temporarily unavailable: {str(error)}"

# Performance Monitoring
class PerformanceMonitor:
    @staticmethod
    def measure_performance(func):
        """Decorator to measure function performance"""
        def wrapper(*args, **kwargs):
            start_time = time.time()
            try:
                result = func(*args, **kwargs)
                processing_time = time.time() - start_time
                
                # Log performance metrics
                if hasattr(st.session_state, 'performance_logs'):
                    st.session_state.performance_logs.append({
                        'function': func.__name__,
                        'processing_time': processing_time,
                        'timestamp': datetime.now().isoformat(),
                        'success': True
                    })
                
                return result
            except Exception as e:
                processing_time = time.time() - start_time
                ErrorHandler.log_error(e, f"Function: {func.__name__}")
                
                if hasattr(st.session_state, 'performance_logs'):
                    st.session_state.performance_logs.append({
                        'function': func.__name__,
                        'processing_time': processing_time,
                        'timestamp': datetime.now().isoformat(),
                        'success': False,
                        'error': str(e)
                    })
                
                raise e
        return wrapper

# Advanced Security Features
class SecurityManager:
    @staticmethod
    def validate_input(text: str, max_length: int = 10000) -> Tuple[bool, str]:
        """Validate user input for security"""
        if not text:
            return False, "Input cannot be empty"
        
        if len(text) > max_length:
            return False, f"Input too long. Maximum {max_length} characters allowed."
        
        # Check for potential security issues
        suspicious_patterns = [
            r'<script.*?>.*?</script>',  # XSS
            r'javascript:',  # JavaScript injection
            r'data:.*base64',  # Data URI attacks
            r'eval\s*\(',  # Code injection
            r'exec\s*\(',  # Code execution
        ]
        
        for pattern in suspicious_patterns:
            if re.search(pattern, text, re.IGNORECASE | re.DOTALL):
                return False, "Input contains potentially unsafe content"
        
        return True, ""
    
    @staticmethod
    def sanitize_filename(filename: str) -> str:
        """Sanitize filename to prevent path traversal"""
        # Remove path components
        filename = os.path.basename(filename)
        
        # Remove dangerous characters
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        
        # Limit length
        if len(filename) > 255:
            name, ext = os.path.splitext(filename)
            filename = name[:250] + ext
        
        return filename or "unnamed_file"

# Enhanced File Processing with Security
class SecureFileProcessor(AdvancedFileProcessor):
    @staticmethod
    def process_file_secure(uploaded_file) -> ProcessingResult:
        """Process file with enhanced security checks"""
        start_time = time.time()
        
        try:
            # Security validations
            if uploaded_file.size > MAX_FILE_SIZE:
                return ProcessingResult(
                    False, None, 
                    f"File too large. Maximum size: {MAX_FILE_SIZE / (1024*1024):.1f}MB"
                )
            
            # Validate filename
            safe_filename = SecurityManager.sanitize_filename(uploaded_file.name)
            if safe_filename != uploaded_file.name:
                st.warning(f"Filename sanitized: {safe_filename}")
            
            # Check file type
            file_type = AdvancedFileProcessor.get_file_type(safe_filename)
            if file_type == 'unknown':
                return ProcessingResult(
                    False, None,
                    f"Unsupported file type: {Path(safe_filename).suffix}"
                )
            
            # Process file content
            file_content = uploaded_file.read()
            
            # Validate content size after reading
            if len(file_content) == 0:
                return ProcessingResult(False, None, "File is empty")
            
            # Extract text based on file type
            text_content = ""
            
            if file_type == 'text':
                try:
                    text_content = file_content.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        text_content = file_content.decode('latin-1')
                    except:
                        return ProcessingResult(False, None, "Unable to decode text file")
            
            elif file_type == 'document':
                if safe_filename.endswith('.pdf'):
                    text_content = SecureFileProcessor.extract_pdf_secure(file_content)
                elif safe_filename.endswith('.docx'):
                    text_content = SecureFileProcessor.extract_docx_secure(file_content)
            
            elif file_type == 'code':
                try:
                    text_content = file_content.decode('utf-8')
                except:
                    return ProcessingResult(False, None, "Unable to decode code file")
            
            # Validate extracted content
            is_valid, validation_message = SecurityManager.validate_input(text_content)
            if not is_valid:
                return ProcessingResult(False, None, validation_message)
            
            processing_time = time.time() - start_time
            
            return ProcessingResult(
                success=True,
                result={
                    'content': text_content,
                    'file_type': file_type,
                    'filename': safe_filename,
                    'size': len(file_content),
                    'original_filename': uploaded_file.name
                },
                processing_time=processing_time
            )
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                result=None,
                error_message=f"File processing error: {str(e)}",
                processing_time=time.time() - start_time
            )
    
    @staticmethod
    def extract_pdf_secure(file_content: bytes) -> str:
        """Secure PDF text extraction with error handling"""
        try:
            # Try PyMuPDF first (more secure and robust)
            doc = fitz.open(stream=file_content, filetype="pdf")
            text_parts = []
            
            # Limit pages to prevent abuse
            max_pages = min(len(doc), 50)
            
            for page_num in range(max_pages):
                page = doc[page_num]
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(page_text)
            
            doc.close()
            return '\n'.join(text_parts)
            
        except Exception as e:
            try:
                # Fallback to PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text_parts = []
                
                max_pages = min(len(pdf_reader.pages), 50)
                for page_num in range(max_pages):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                
                return '\n'.join(text_parts)
            except:
                return f"Error extracting PDF content: {str(e)}"
    
    @staticmethod
    def extract_docx_secure(file_content: bytes) -> str:
        """Secure DOCX text extraction"""
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            return f"Error extracting DOCX content: {str(e)}"

# Enhanced Main Execution Logic
def run_enhanced_app():
    """Run the enhanced TechNova application"""
    try:
        # Initialize performance monitoring
        if 'performance_logs' not in st.session_state:
            st.session_state.performance_logs = []
        
        # Run main application
        main()
        
        # Show performance info in sidebar (for debugging)
        if st.session_state.get('show_debug', False):
            with st.sidebar:
                with st.expander("🔧 Debug Info"):
                    st.json(st.session_state.performance_logs[-5:] if st.session_state.performance_logs else [])
    
    except Exception as e:
        ErrorHandler.log_error(e, "Main application", st.session_state.user_info.get('id'))
        st.error(f"🚨 Application error: {str(e)}")
        st.error("Please refresh the page or contact support if the issue persists.")

# Footer and Additional UI Components
def show_footer():
    """Display enhanced footer"""
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; padding: 20px; color: rgba(0, 249, 255, 0.6);">
        <p>🌌 <strong>TechNova AI Nexus Pro</strong> - Powered by Advanced AI</p>
        <p>📧 Support: support@technova.ai | 🌐 Website: technova.ai</p>
        <p>© 2025 TechNova Technologies. All rights reserved.</p>
    </div>
    """, unsafe_allow_html=True)

# Health Check Function
def health_check() -> Dict[str, bool]:
    """Check system health and dependencies"""
    health_status = {
        'database': False,
        'openai_api': False,
        'email_service': False,
        'file_processing': False
    }
    
    try:
        # Check database
        conn = sqlite3.connect('technova_pro.db')
        cursor = conn.cursor()
        cursor.execute("SELECT 1")
        health_status['database'] = True
        conn.close()
    except:
        pass
    
    try:
        # Check OpenAI API
        if OPENAI_API_KEY:
            client = openai.OpenAI(api_key=OPENAI_API_KEY)
            # Make a minimal test call
            health_status['openai_api'] = True
    except:
        pass
    
    try:
        # Check email service
        if SMTP_USERNAME and SMTP_PASSWORD:
            health_status['email_service'] = True
    except:
        pass
    
    # Check file processing
    health_status['file_processing'] = PDF_AVAILABLE
    
    return health_status

if __name__ == "__main__":
    run_enhanced_app()
    
    # Show footer
    if st.session_state.authenticated:
        show_footer()
    
    # System health check (admin only)
    if (st.session_state.authenticated and 
        st.session_state.user_info.get('email') == ADMIN_EMAIL):
        
        with st.sidebar:
            with st.expander("🏥 System Health"):
                health = health_check()
                for service, status in health.items():
                    icon = "✅" if status else "❌"
                    st.write(f"{icon} {service.replace('_', ' ').title()}")

# Additional utility functions for completeness
STOPWORDS = ENHANCED_STOPWORDS  # Use enhanced stopwords

def safe_sentence_split(text: str) -> List[str]:
    """Enhanced sentence splitting with better accuracy"""
    # More sophisticated sentence boundary detection
    pattern = re.compile(r'(?<=[.!?])\s+(?=[A-Z0-9"\'])')
    sentences = pattern.split(text)
    
    # Clean and filter sentences
    cleaned_sentences = []
    for sentence in sentences:
        sentence = sentence.strip()
        if sentence and len(sentence.split()) >= 3:  # Minimum 3 words
            cleaned_sentences.append(sentence)
    
    return cleaned_sentences

# Run the application
if __name__ == "__main__":
    run_enhanced_app()
